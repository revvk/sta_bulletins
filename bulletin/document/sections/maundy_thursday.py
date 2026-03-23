"""
Maundy Thursday section module.

Composes the Maundy Thursday liturgy from shared building blocks:
  1. Word of God (festal acclamation, no Nicene Creed)
  2. Foot Washing (unique to Maundy Thursday)
  3. Prayers of the People + Confession + Peace
  4. Holy Communion (no blessing, no dismissal)
  5. Post-Communion Hymn
  6. Stripping of the Altar + Psalm 22
"""

from docx import Document

from bulletin.config import CROSS_SYMBOL, FONT_BODY_BOLD
from bulletin.data.loader import load_common_prayers, load_eucharistic_prayers
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line,
)
from bulletin.document.sections.word_of_god import (
    add_standard_opening, add_reading, add_psalm, add_gospel,
    add_confession, add_pop, add_song_smart, add_body_with_amen,
)
from bulletin.document.sections.holy_communion import (
    add_prayer_a_or_b, add_prayer_c,
    add_short_offertory_rubric, add_agnus_dei_spoken,
)
from bulletin.logic.rules import SeasonalRules


def add_maundy_thursday(doc: Document, rules: SeasonalRules,
                         wog_data: dict, hc_data: dict, mt_data: dict):
    """Add the complete Maundy Thursday liturgy.

    Args:
        doc: The Document.
        rules: SeasonalRules (with festal acclamation overrides).
        wog_data: Standard Word of God data dict.
        hc_data: Standard Holy Communion data dict.
        mt_data: Maundy Thursday specific data:
            - foot_washing_songs: list of song data dicts
            - foot_washing: dict with invitation, rubric
            - anthem: dict with rubric, sections (list of {text, bold})
            - post_communion_hymn: song data dict | None
            - stripping: dict with rubric, psalm_ref, psalm_rubric,
                        closing_rubric
            - psalm_22_text: list of str (psalm verse lines)
            - watch: dict with text
    """
    prayers = load_common_prayers()

    # ==========================================
    # Part 1: Word of God
    # ==========================================
    _add_word_of_god(doc, rules, wog_data, prayers)

    # ==========================================
    # Part 2: Foot Washing
    # ==========================================
    _add_foot_washing(doc, mt_data)

    # ==========================================
    # Part 3: Prayers, Confession, Peace
    # ==========================================
    _add_prayers_and_peace(doc, rules, wog_data, prayers)

    # ==========================================
    # Part 4: Holy Communion (modified ending)
    # ==========================================
    _add_holy_communion(doc, rules, hc_data, mt_data)

    # ==========================================
    # Part 5: Stripping of the Altar + Psalm 22
    # ==========================================
    _add_stripping(doc, mt_data)


def _add_word_of_god(doc: Document, rules: SeasonalRules,
                      data: dict, prayers: dict):
    """Word of God section — festal opening, no Creed."""
    # Standard opening (rules already has festal acclamation)
    add_standard_opening(doc, rules, data, prayers)

    # Collect of the Day
    add_spacer(doc)
    add_heading2(doc, "Collect of the Day")
    add_celebrant_line(doc, "Celebrant", "The Lord be with you.")
    add_people_line(doc, "People", "And also with you.")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, data["collect_text"])

    # Be seated for readings
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # First Reading
    add_reading(doc, data["reading_1_ref"], data["reading_1_text"])

    # Psalm
    add_spacer(doc)
    add_psalm(doc, data["psalm_ref"], data.get("psalm_rubric", ""),
               data["psalm_text"])

    # Sequence Hymn
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Sequence Hymn")
    add_song_smart(doc, data.get("sequence_hymn"))

    # Gospel
    add_gospel(doc, data["gospel_ref"], data["gospel_book"],
                data["gospel_text"])

    # The Homily (not "Sermon" — Maundy Thursday tradition)
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "The Homily")
    add_body(doc, data.get("preacher", ""))

    # NO Nicene Creed on Maundy Thursday


def _add_foot_washing(doc: Document, mt_data: dict):
    """Add the Foot Washing section unique to Maundy Thursday."""
    add_spacer(doc)
    add_heading2(doc, "Foot Washing")

    # Invitation text from the celebrant
    foot_washing = mt_data.get("foot_washing", {})
    invitation = foot_washing.get("invitation", "")
    if invitation:
        add_body(doc, invitation)
        add_spacer(doc)

    # Rubric about washing stations
    rubric = foot_washing.get("rubric", "")
    if rubric:
        add_introductory_rubric(doc, rubric)

    # Music during foot washing
    songs = mt_data.get("foot_washing_songs", [])
    for i, song in enumerate(songs):
        if i > 0:
            add_spacer(doc)
        add_song_smart(doc, song)

    # Anthem
    anthem = mt_data.get("anthem", {})
    anthem_rubric = anthem.get("rubric", "")
    if anthem_rubric:
        add_spacer(doc)
        add_introductory_rubric(doc, anthem_rubric)

    sections = anthem.get("sections", [])
    if sections:
        # Render alternating normal/bold sections
        for section in sections:
            text = section.get("text", "")
            is_bold = section.get("bold", False)
            if is_bold:
                p = doc.add_paragraph(style="Body")
                run = p.add_run(text)
                run.bold = True
                run.font.name = FONT_BODY_BOLD
            else:
                add_body(doc, text)


def _add_prayers_and_peace(doc: Document, rules: SeasonalRules,
                             data: dict, prayers: dict):
    """Add Prayers of the People, Confession, and Peace."""
    # Prayers of the People
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Prayers of the People")
    add_pop(doc, data.get("pop_elements", []))
    add_spacer(doc)
    add_rubric(doc, data.get("pop_concluding_rubric",
                              "The Celebrant concludes with a suitable Collect."))

    # Confession & Absolution
    add_spacer(doc)
    add_introductory_rubric(doc, "Please kneel or remain standing.")
    add_confession(doc, prayers)

    # The Peace
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Peace")
    add_celebrant_line(doc, "Celebrant",
                       "The peace of the Lord be always with you.")
    add_people_line(doc, "People", "And also with you.")


def _add_holy_communion(doc: Document, rules: SeasonalRules,
                          hc_data: dict, mt_data: dict):
    """Modified Holy Communion — no blessing, no dismissal."""
    prayers = load_common_prayers()
    ep_data = load_eucharistic_prayers()

    add_heading(doc, "The Holy Communion")

    # --- Offertory ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "Offertory")
    add_short_offertory_rubric(doc)
    add_spacer(doc)

    # Offering Music
    add_heading2(doc, "Offering Music")
    add_song_smart(doc, hc_data.get("offertory_song"))

    # Doxology
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Doxology")
    _DOXOLOGY = [
        "Praise God, from Whom all blessings flow;",
        "Praise Him, all creatures here below;",
        "Praise Him above, ye heavenly host:",
        "Praise Father, Son, and Holy Ghost.",
    ]
    for line in _DOXOLOGY:
        doc.add_paragraph(line, style="Body - Lyrics")

    # --- The Great Thanksgiving ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "The Great Thanksgiving")

    # Sursum Corda
    for exchange in ep_data["sursum_corda"]:
        add_celebrant_line(doc, "Celebrant", exchange["celebrant"])
        add_people_line(doc, "People", exchange["people"])
    add_spacer(doc)

    # Eucharistic Prayer
    prayer_key = hc_data.get("eucharistic_prayer", "A").upper()
    if prayer_key == "C":
        add_prayer_c(doc, ep_data, hc_data, prayers)
    else:
        add_prayer_a_or_b(doc, ep_data, hc_data, prayers, prayer_key)

    # --- Lord's Prayer ---
    add_spacer(doc)
    add_rubric(doc, "Please stand and, as you are comfortable, join hands "
               "with those around you.")
    add_body(doc, prayers["lords_prayer_intro"]["option_1"])
    text = " ".join(line.strip() for line in prayers["lords_prayer"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]

    # --- Breaking of the Bread ---
    add_spacer(doc)
    add_heading2(doc, "Breaking of the Bread")
    if rules.use_fraction_anthem:
        fraction_song = hc_data.get("fraction_song")
        if fraction_song:
            add_song_smart(doc, fraction_song, force_single_column=True)
        else:
            add_agnus_dei_spoken(doc)
    else:
        add_celebrant_line(doc, "Celebrant", rules.fraction_celebrant)
        add_people_line(doc, "People", rules.fraction_people)

    # --- Invitation ---
    add_spacer(doc)
    add_heading2(doc, "Invitation")
    add_rubric(doc, "Facing the people, the Celebrant says the "
               "following Invitation")
    add_body(doc, prayers["invitation_to_communion"] + " " +
             prayers["invitation_addition"])

    # --- Holy Communion ---
    add_spacer(doc)
    add_heading2(doc, "Holy Communion")
    add_rubric(doc, "All baptized Christians are invited to the altar rail to "
               "receive Holy Communion. If you are not going to receive, you "
               "are still welcome at the rail for a blessing; just cross your "
               "arms over your chest.")

    # --- Communion Music ---
    add_spacer(doc)
    add_heading2(doc, "Communion Music")
    comm_songs = hc_data.get("communion_songs", [])
    for i, song in enumerate(comm_songs):
        if i > 0:
            add_spacer(doc)
        add_song_smart(doc, song)

    # --- Closing Prayer ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Closing Prayer")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")

    prayer_choice = hc_data.get("post_communion_prayer", "a")
    if prayer_choice == "b":
        text = " ".join(prayers["post_communion_prayer_b"])
    else:
        text = " ".join(prayers["post_communion_prayer_a"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]

    # --- Post-Communion Hymn (replaces closing hymn) ---
    post_comm_hymn = mt_data.get("post_communion_hymn")
    if post_comm_hymn:
        add_spacer(doc)
        add_introductory_rubric(doc, "Remain standing.")
        add_heading2(doc, "Post-Communion Hymn")
        add_rubric(doc, "As the reserved elements are taken to the Altar of "
                   "Repose for The Watch, the Cantor will lead the "
                   "congregation in the following hymn")
        add_song_smart(doc, post_comm_hymn)

    # NO blessing
    # NO dismissal


def _add_stripping(doc: Document, mt_data: dict):
    """Add the Stripping of the Altar with Psalm 22."""
    stripping = mt_data.get("stripping", {})

    add_spacer(doc)
    add_heading2(doc, "Stripping the Altar")

    rubric = stripping.get("rubric", "")
    if rubric:
        add_rubric(doc, rubric)

    # Psalm 22
    psalm_ref = stripping.get("psalm_ref", "Psalm 22")
    psalm_rubric = stripping.get("psalm_rubric", "Read in unison.")
    psalm_text = mt_data.get("psalm_22_text", [])

    if psalm_text:
        add_spacer(doc)
        add_psalm(doc, psalm_ref, psalm_rubric, psalm_text)

    # Closing rubric
    closing_rubric = stripping.get("closing_rubric", "")
    if closing_rubric:
        add_spacer(doc)
        add_rubric(doc, closing_rubric)

    # Watch announcement
    watch = mt_data.get("watch", {})
    watch_text = watch.get("text", "")
    if watch_text:
        add_spacer(doc)
        add_heading2(doc, "The Watch Begins")
        add_introductory_rubric(doc, watch_text)
