"""
Cover page section of the bulletin.

Generates the first page with:
  - Date and service time
  - "Invite | Involve | Instruct | Inspire" tagline
  - Welcome text for newcomers (static)
  - Online giving info
  - Liturgical day title
"""

from docx import Document

from bulletin.document.formatting import add_spacer


# Static cover text — same every week
_WELCOME_LINES = [
    ("New Here? Wonderful! And welcome!", True, False),
    ("Fill out a Connection Card. If you do…", False, False),
    ("We will pray for you by name this week.", False, True),
    ("You'll get a handwritten welcome note from the clergy.", False, False),
    ("Grab a What Is Going On Here? card from the tables outside the door. "
     "It explains much of what happens during worship and why we do it.",
     False, True),
    ("Clergy and lay volunteers are happy to answer any questions you might have.",
     False, False),
    ("Holy Communion: Share in the Body and Blood of our Lord Jesus Christ.",
     True, False),
    ("Bread is offered first; wine second", False, False),
    ("The first chalice is the common cup (sip directly). The second chalice "
     "is for intinction (the bread will be dipped).", False, False),
    ("Let the clergy know if you need a Gluten Free wafer", False, False),
    ("Prayer Partners can be found at the back of the church during holy "
     "communion to personally pray for whatever is on your heart.", True, False),
    ("Assisted Hearing Devices are available for those that would like one. "
     "Please ask an usher.", True, False),
    ("Crosses in the Bulletin? The crosses (\u2720) are placed in the text "
     "of the liturgy at times when it is customary to make the sign of "
     "the cross.", True, True),
    ("Giving: Thank you for your gifts!", True, True),
    ("Offering plates are passed for in-person giving (cash and checks).",
     False, False),
    ("Online giving is available and easy to do using the link and QR code "
     "below (cards).", False, False),
]


def add_cover(doc: Document, date_str: str, service_time: str,
              liturgical_title: str, giving_url: str = "standrewsmckinney.org/give"):
    """Add the cover page to the document.

    Args:
        doc: The Document to add to.
        date_str: Formatted date string (e.g., "March 1, 2026")
        service_time: e.g., "9 am"
        liturgical_title: e.g., "The Third Sunday in Lent"
        giving_url: URL for online giving
    """
    # Date and time line
    p = doc.add_paragraph(style="Body")
    run = p.add_run(f"{date_str}  |  {service_time}")
    run.bold = True

    # Tagline
    doc.add_paragraph("Invite \tInvolve \tInstruct \tInspire", style="Body")

    # Welcome text
    for text, bold, italic in _WELCOME_LINES:
        p = doc.add_paragraph(style="Cover Note")
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True

    # Online giving heading
    doc.add_paragraph("Online Giving", style="Heading 2")
    doc.add_paragraph(giving_url, style="Heading 2")

    # Welcome label (italic)
    p = doc.add_paragraph(style="Body")
    run = p.add_run("Welcome")
    run.italic = True

    # Liturgical day title
    doc.add_paragraph(liturgical_title, style="Heading")

    # Prelude note
    p = doc.add_paragraph(style="Body - Introductory Rubric")
    p.add_run(
        "Once the Prelude begins, please refrain from further visiting and "
        "conversation as we prepare our hearts and thoughts for worship."
    )
    add_spacer(doc)

    # Prelude heading
    doc.add_paragraph("Prelude", style="Heading 2")
    add_spacer(doc)
