"""
The Word of God section of the bulletin.

Covers everything from the processional hymn through The Peace:
  - Processional hymn
  - Opening Acclamation
  - Collect for Purity (standard) or Penitential Order (Lent)
  - Song of Praise / Kyrie / Advent Wreath
  - Collect of the Day
  - Children's Sermon
  - Scripture readings (Epistle/OT)
  - Psalm / Canticle
  - Sequence Hymn
  - Gospel
  - Sermon
  - Nicene Creed
  - Prayers of the People
  - Confession & Absolution (if not in Penitential Order)
  - The Peace
"""

from docx import Document
from docx.shared import Pt

from bulletin.config import CROSS_SYMBOL, FONT_BODY_BOLD
from bulletin.data.loader import load_common_prayers
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line, add_dialogue, add_cross_symbol,
    add_hymn_header, add_song, add_song_two_column,
    add_scripture_text, _add_text_runs,
)
from bulletin.logic.rules import SeasonalRules


def add_word_of_god(doc: Document, rules: SeasonalRules, data: dict):
    """Add the entire Word of God section.

    Args:
        doc: The Document.
        rules: SeasonalRules for this service.
        data: Dict with keys:
            - processional: song data dict (or None)
            - song_of_praise: song data dict (or None)
            - sequence_hymn: song data dict (or None)
            - collect_text: str (the collect of the day text)
            - reading_1_ref: str (e.g., "1 Corinthians 10:1-13")
            - reading_1_text: ScriptureReading
            - psalm_ref: str (e.g., "Psalm 63:1-8")
            - psalm_text: list of str lines
            - psalm_rubric: str (e.g., "Read in unison.")
            - gospel_ref: str (e.g., "Luke 13:1-9")
            - gospel_book: str (e.g., "Luke")
            - gospel_text: ScriptureReading
            - preacher: str (e.g., "The Rev. Andrew Van Kirk")
            - pop_elements: list of POP element dicts
            - pop_concluding_rubric: str
            - advent_wreath_verse: str | None (for Advent)
            - advent_hymnal_ref: str | None
            - penitential_sentence: str | None (for Lent 2-5)
    """
    prayers = load_common_prayers()

    # --- Penitential Order (Lent) or standard Word of God ---
    if rules.use_penitential_order:
        add_penitential_order(doc, rules, data, prayers)
    else:
        add_standard_opening(doc, rules, data, prayers)

    # --- Collect of the Day ---
    add_spacer(doc)
    add_heading2(doc, "Collect of the Day")
    add_celebrant_line(doc, "Celebrant", "The Lord be with you.")
    add_people_line(doc, "People", "And also with you.")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, data["collect_text"])

    # --- Be seated for readings / Children ---
    add_spacer(doc)
    service_time = data.get("service_time", "9 am")
    if service_time == "11 am":
        # Children's Church rubric (replaces Children's Sermon heading)
        add_introductory_rubric(
            doc,
            "Elementary-aged children are invited to attend Children\u2019s Church. "
            "They will rejoin the service during The Peace."
        )
        add_spacer(doc)
        add_introductory_rubric(doc, "Be seated.")
    elif service_time == "8 am":
        # No children's section at 8am
        add_introductory_rubric(doc, "Be seated.")
    else:
        add_introductory_rubric(doc, "Be seated.")
        # --- Children's Sermon ---
        add_heading2(doc, "Children\u2019s Sermon")
        add_spacer(doc)

    # --- First Reading ---
    add_reading(doc, data["reading_1_ref"], data["reading_1_text"])

    # --- Psalm ---
    add_spacer(doc)
    add_psalm(doc, data["psalm_ref"], data.get("psalm_rubric", ""),
               data["psalm_text"])

    # --- Sequence Hymn ---
    add_spacer(doc)
    if service_time != "8 am":
        add_introductory_rubric(doc, "Please stand.")
        add_heading2(doc, "Sequence Hymn")
        add_song_smart(doc, data.get("sequence_hymn"))

    # --- Gospel ---
    add_gospel(doc, data["gospel_ref"], data["gospel_book"],
                data["gospel_text"], service_time=service_time)

    # --- Sermon ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "Sermon")
    add_body(doc, data.get("preacher", ""))

    # --- Nicene Creed ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Nicene Creed")
    add_nicene_creed(doc, prayers)

    # --- Prayers of the People ---
    add_spacer(doc)
    add_heading2(doc, "Prayers of the People")
    add_pop(doc, data.get("pop_elements", []))
    add_spacer(doc)
    add_rubric(doc, data.get("pop_concluding_rubric",
                              "The Celebrant concludes with a suitable Collect."))

    # --- Confession & Absolution (if not already done in Penitential Order) ---
    if not rules.no_confession_after_pop:
        add_spacer(doc)
        add_confession(doc, prayers)

    # --- The Peace ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Peace")
    add_celebrant_line(doc, "Celebrant", "The peace of the Lord be always with you.")
    add_people_line(doc, "People", "And also with you.")


def add_standard_opening(doc: Document, rules: SeasonalRules,
                          data: dict, prayers: dict):
    """Standard (non-Lent) opening: Word of God → Processional → Acclamation → Collect for Purity → Song of Praise."""
    service_time = data.get("service_time", "9 am")

    add_heading(doc, "The Word of God")
    add_spacer(doc)

    # Processional
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Processional")
    if service_time != "8 am":
        add_song_smart(doc, data.get("processional"))

    # Opening Acclamation
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Opening Acclamation")
    cel_text = rules.acclamation_celebrant.replace("{cross}", CROSS_SYMBOL)
    add_celebrant_with_cross(doc, "Celebrant", cel_text)
    add_people_line(doc, "People", rules.acclamation_people)
    add_spacer(doc)

    # Collect for Purity
    if rules.include_collect_for_purity:
        add_body_with_amen(doc, prayers["collect_for_purity"])

    # Song of Praise / Kyrie / Advent Wreath
    add_spacer(doc)
    add_heading2(doc, rules.song_of_praise_label)
    if rules.is_advent:
        _add_advent_wreath(doc, data)
    elif service_time == "8 am":
        _add_gloria_spoken(doc, prayers)
    else:
        add_song_smart(doc, data.get("song_of_praise"))


def add_penitential_order(doc: Document, rules: SeasonalRules,
                           data: dict, prayers: dict):
    """Lenten opening: Penitential Order → Processional → Acclamation → [Decalogue/Sentence] → Confession → Word of God → Kyrie."""
    service_time = data.get("service_time", "9 am")

    add_heading(doc, "A Penitential Order")
    add_spacer(doc)

    # Processional
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Processional")
    if service_time != "8 am":
        add_song_smart(doc, data.get("processional"))

    # Opening Acclamation
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Opening Acclamation")
    cel_text = rules.acclamation_celebrant.replace("{cross}", CROSS_SYMBOL)
    add_celebrant_with_cross(doc, "Celebrant", cel_text)
    add_people_line(doc, "People", rules.acclamation_people)
    add_spacer(doc)

    # Decalogue (Lent 1) or Scripture Sentence (Lent 2-5)
    if rules.use_decalogue:
        _add_decalogue(doc, prayers)
    else:
        sentence = data.get("penitential_sentence", "")
        if sentence:
            sentence_ref = data.get("penitential_sentence_ref", "")
            p = doc.add_paragraph(style="Body")
            p.add_run(sentence)
            if sentence_ref:
                run = p.add_run(f" ({sentence_ref})")
                run.italic = True

    # Confession
    add_spacer(doc)
    add_introductory_rubric(doc, "Please kneel or remain standing.")
    add_confession(doc, prayers)

    # Then "The Word of God" section heading
    add_heading(doc, "The Word of God")

    # Kyrie
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Kyrie")
    if service_time == "8 am":
        _add_kyrie_spoken(doc)
    else:
        add_song_smart(doc, data.get("song_of_praise"))


def add_confession(doc: Document, prayers: dict):
    """Add the Confession of Sin and Absolution."""
    add_heading2(doc, "Confession of Sin")
    add_rubric(doc, "The Deacon or a Priest says")
    add_celebrant_line(doc, "", prayers["confession_invitation"])

    # Confession text (bold via People character style, recited by people)
    text = " ".join(prayers["confession"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]
    add_spacer(doc)

    # Absolution
    add_rubric(doc, "The Priest says the Absolution.")
    abs_text = prayers["absolution"]
    p = doc.add_paragraph(style="Body")
    # Insert cross at "forgive you all your sins" and bold the final "Amen."
    if "forgive you" in abs_text:
        before, after = abs_text.split("forgive you", 1)
        p.add_run(before)
        add_cross_symbol(p)
        # Split off final "Amen." to bold it
        remainder = " forgive you" + after
        if remainder.rstrip().endswith("Amen."):
            body_part = remainder.rstrip()[:-5]
            p.add_run(body_part)
            run = p.add_run("Amen.")
            run.bold = True
            run.font.name = FONT_BODY_BOLD
        else:
            p.add_run(remainder)
    else:
        p.add_run(abs_text)


def _add_decalogue(doc: Document, prayers: dict):
    """Add the Decalogue (Ten Commandments) for Lent 1."""
    add_body(doc, prayers["decalogue_intro"])
    for item in prayers["decalogue"]:
        add_celebrant_line(doc, "", item["commandment"])
        p = doc.add_paragraph(style="Body - Dialogue")
        run = p.add_run(item["response"])
        run.style = doc.styles["People"]
    add_spacer(doc)


def _add_reading_text(doc: Document, reading):
    """Render the text of a scripture reading (prose, poetry, or interleaved).

    Supports three formats:
    - segments: list of {type: 'prose'/'poetry'} dicts for interleaved content
    - paragraphs + poetry_lines: prose first, then poetry appended
    - plain string fallback
    """
    if hasattr(reading, "segments") and reading.segments:
        for seg in reading.segments:
            if seg["type"] == "prose":
                add_scripture_text(doc, seg["text"])
            elif seg["type"] == "poetry":
                for line in seg["lines"]:
                    # Lines can be plain strings (legacy) or dicts with indent
                    if isinstance(line, dict):
                        indent_level = line.get("indent", 0)
                        text = line["text"]
                    else:
                        indent_level = 0
                        text = line
                    if indent_level == 0:
                        style = "Reading (Poetry)"
                    elif indent_level == 1:
                        style = "Reading (Poetry Indent 1)"
                    else:
                        style = "Reading (Poetry Indent 2)"
                    add_scripture_text(doc, text, style=style)
    elif hasattr(reading, "paragraphs"):
        for i, para in enumerate(reading.paragraphs):
            add_scripture_text(doc, para, indent=(i > 0))
        if reading.has_poetry:
            for line in reading.poetry_lines:
                add_scripture_text(doc, line, style="Reading (Poetry)")
    else:
        add_scripture_text(doc, str(reading))


def add_reading(doc: Document, reference: str, reading):
    """Add a scripture reading with responses."""
    add_heading2(doc, f"The Scriptures: {reference}")

    _add_reading_text(doc, reading)

    add_spacer(doc)
    add_rubric(doc, "After the reading, the Reader will say")
    add_celebrant_line(doc, "", "The Word of the Lord.")
    add_people_line(doc, "People", "Thanks be to God.")


def add_psalm(doc: Document, reference: str, rubric: str, lines):
    """Add a psalm or canticle.

    Each entry in *lines* is a single verse formatted as:
      "first half text\\n\\tsecond half line 1\\n\\tsecond half line 2"
    We render each verse as one "Psalm"-styled paragraph, using
    line breaks (not new paragraphs) within a verse so the hanging
    indent applies to the second-half lines.

    When the psalm is read in unison, all verses are bold (the whole
    congregation reads together).  When read responsively, antiphonally,
    or alternating, verses alternate between regular and bold — the
    congregation reads the bold verses.
    """
    add_heading2(doc, reference)

    if rubric:
        add_rubric(doc, rubric)

    # Determine bold pattern from the rubric
    rubric_lower = rubric.lower() if rubric else ""
    alternating = any(kw in rubric_lower for kw in
                      ("responsiv", "antiphon", "alternating", "men and women",
                       "half verse"))
    # Unison = all bold (everyone reads together)
    all_bold = "unison" in rubric_lower

    if isinstance(lines, list):
        for verse_idx, verse_text in enumerate(lines):
            if all_bold:
                bold_verse = True
            elif alternating:
                bold_verse = (verse_idx % 2 == 1)
            else:
                bold_verse = False
            p = doc.add_paragraph(style="Psalm")
            if "\n" in verse_text:
                sub_lines = verse_text.split("\n")
                for i, sub in enumerate(sub_lines):
                    if sub.startswith("\v"):
                        # First-half continuation: new paragraph so it
                        # starts at the left like the first line
                        p.paragraph_format.space_after = Pt(0)
                        p = doc.add_paragraph(style="Psalm")
                        p.paragraph_format.space_before = Pt(0)
                        _add_text_runs(p, sub[1:], bold=bold_verse)
                    elif i > 0:
                        run = p.add_run()
                        run.add_break()
                        _add_text_runs(p, sub, bold=bold_verse)
                    else:
                        _add_text_runs(p, sub, bold=bold_verse)
            else:
                _add_text_runs(p, verse_text, bold=bold_verse)
    elif hasattr(lines, "paragraphs"):
        for verse_idx, para in enumerate(lines.paragraphs):
            if all_bold:
                bold_verse = True
            elif alternating:
                bold_verse = (verse_idx % 2 == 1)
            else:
                bold_verse = False
            p = doc.add_paragraph(style="Psalm")
            _add_text_runs(p, para, bold=bold_verse)


def add_gospel(doc: Document, reference: str, book: str, reading,
                service_time: str = "9 am"):
    """Add the Gospel reading with announcement and response."""
    add_spacer(doc)
    if service_time == "8 am":
        add_introductory_rubric(doc, "Please stand.")
    else:
        add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, f"The Gospel: {reference}")

    add_rubric(doc, "The Deacon or a Priest reads the Gospel, first saying")

    # Announcement
    add_celebrant_line(doc, "",
                       f"The Holy Gospel of our Lord Jesus Christ according to {book}.")
    add_people_line(doc, "People", "Glory to you, Lord Christ.")
    add_spacer(doc)

    # Gospel text
    if hasattr(reading, "paragraphs"):
        for i, para in enumerate(reading.paragraphs):
            add_scripture_text(doc, para, indent=(i > 0))
    else:
        add_scripture_text(doc, str(reading))

    add_spacer(doc)
    add_rubric(doc, "After the Gospel, the Reader will say")
    add_celebrant_line(doc, "", "The Gospel of the Lord.")
    add_people_line(doc, "People", "Praise to you, Lord Christ.")


def add_nicene_creed(doc: Document, prayers: dict):
    """Add the Nicene Creed in the three-article format."""
    creed_lines = prayers["nicene_creed"]
    # Group into three articles (separated by blank lines)
    articles = []
    current = []
    for line in creed_lines:
        if line == "":
            if current:
                articles.append(" ".join(current))
                current = []
        else:
            current.append(line.strip())
    if current:
        articles.append(" ".join(current))

    for article in articles:
        p = doc.add_paragraph(style="Body - People Recitation (Creed)")
        if CROSS_SYMBOL in article:
            parts = article.split(CROSS_SYMBOL)
            for i, part in enumerate(parts):
                if i > 0:
                    add_cross_symbol(p)
                if part:
                    run = p.add_run(part)
                    run.style = doc.styles["People"]
        else:
            run = p.add_run(article)
            run.style = doc.styles["People"]


def add_pop(doc: Document, elements: list[dict]):
    """Add Prayers of the People elements."""
    for i, elem in enumerate(elements):
        etype = elem.get("type", "leader")
        text = elem.get("text", "")

        if etype == "leader":
            add_body(doc, text)
            # Spacer after intro line (leader followed by another leader)
            next_type = elements[i + 1].get("type") if i + 1 < len(elements) else None
            if next_type == "leader":
                add_spacer(doc)
        elif etype == "people":
            # Support split leader/people text on the same line
            leader_text = elem.get("leader_text", "")
            people_text = elem.get("people_text", text)
            p = doc.add_paragraph(style="Body")
            if leader_text:
                p.add_run(leader_text + " ")
            run = p.add_run(people_text)
            run.bold = True
            run.font.name = FONT_BODY_BOLD
        elif etype == "rubric":
            add_rubric(doc, text)
        elif etype == "both":
            add_body(doc, elem.get("leader_text", ""))
            p = doc.add_paragraph(style="Body")
            run = p.add_run(elem.get("people_text", ""))
            run.bold = True
            run.font.name = FONT_BODY_BOLD

        # Add spacer between petitions (but not after the last one)
        if etype in ("people", "both") and i + 1 < len(elements):
            add_spacer(doc)


def _add_advent_wreath(doc: Document, data: dict):
    """Add the Advent wreath lighting with O Come Emmanuel."""
    hymn_ref = data.get("advent_hymnal_ref", "#56 (Hymnal 1982)")
    p = doc.add_paragraph(style="Body")
    p.add_run("O come, O come, Emmanuel")
    run = p.add_run(f"\t{hymn_ref}")
    run.italic = True

    # Verse 1 (always)
    advent_v1 = [
        "O come, O come, Emmanuel,",
        "and ransom captive Israel,",
        "that mourns in lonely exile here",
        "until the Son of God appear.",
    ]
    for line in advent_v1:
        doc.add_paragraph(line, style="Body - Lyrics")
    add_spacer(doc)

    # Chorus
    chorus = ["Rejoice! Rejoice! Emmanuel", "shall come to thee, O Israel!"]
    for line in chorus:
        p = doc.add_paragraph(style="Body - Lyrics")
        run = p.add_run(line)
        run.italic = True

    # Additional O Antiphon verse if provided
    advent_verse = data.get("advent_wreath_verse")
    if advent_verse and isinstance(advent_verse, list):
        add_spacer(doc)
        for line in advent_verse:
            doc.add_paragraph(line, style="Body - Lyrics")
        add_spacer(doc)
        for line in chorus:
            p = doc.add_paragraph(style="Body - Lyrics")
            run = p.add_run(line)
            run.italic = True


def add_song_smart(doc: Document, song_data: dict | None,
                    force_single_column: bool = False):
    """Add a song, choosing two-column layout when it saves space.

    If the song has no sections (hymnal-only), renders just the header
    line (title + hymnal reference) without wrapping in a lyric table.

    Args:
        force_single_column: If True, always use single-column layout
            with one row per section (multi-row table).  Use for
            large-print bulletins and liturgical service music.
    """
    if not song_data:
        add_body(doc, "[Song lyrics not found]")
        return

    sections = song_data.get("sections", [])

    # Hymnal-only: just the header line, no lyrics
    if not sections:
        add_hymn_header(
            doc,
            song_data["title"],
            song_data.get("tune_name"),
            song_data.get("hymnal_number"),
            song_data.get("hymnal_name"),
        )
        return

    max_line_len = max(
        (len(line) for s in sections for line in s["lines"]),
        default=0
    )

    # Two-column if 2+ sections and lines fit half-width
    if not force_single_column and len(sections) >= 2 and max_line_len <= 52:
        add_song_two_column(doc, song_data)
    else:
        add_song(doc, song_data, multi_row=force_single_column)


def add_body_with_amen(doc: Document, text: str):
    """Add body text, making the final 'Amen.' bold."""
    if text.rstrip().endswith("Amen."):
        body = text.rstrip()[:-5]
        p = doc.add_paragraph(style="Body")
        p.add_run(body)
        run = p.add_run("Amen.")
        run.bold = True
        run.font.name = FONT_BODY_BOLD
    else:
        add_body(doc, text)


def add_celebrant_with_cross(doc: Document, label: str, text: str):
    """Add a celebrant line, rendering {cross} / ✠ as a bold-italic cross symbol."""
    p = doc.add_paragraph(style="Body - Dialogue")
    if label:
        p.add_run(label)
        p.add_run("\t")

    if CROSS_SYMBOL in text:
        parts = text.split(CROSS_SYMBOL)
        p.add_run(parts[0])
        add_cross_symbol(p)
        if len(parts) > 1:
            p.add_run(parts[1])
    else:
        p.add_run(text)


def _add_gloria_spoken(doc: Document, prayers: dict):
    """Add the Gloria as spoken bold text (People recitation) for 8am.

    Renders each section as a separate bold paragraph with line breaks
    within, so the poetic structure is visible while remaining bold.
    """
    gloria = prayers["gloria"]
    # Support both old format (list of strings) and new (dict with sections)
    if isinstance(gloria, dict):
        sections = gloria.get("sections", [])
    else:
        # Legacy: flat list → one section
        sections = [gloria]

    for section_lines in sections:
        text = "\n".join(line.strip() for line in section_lines)
        p = doc.add_paragraph(style="Body - People Recitation")
        # Use add_break() for line breaks within the paragraph
        for i, line in enumerate(section_lines):
            if i > 0:
                run = p.runs[-1] if p.runs else p.add_run("")
                run.add_break()
            run = p.add_run(line.strip())
            run.style = doc.styles["People"]


def _add_kyrie_spoken(doc: Document):
    """Add the Kyrie as spoken Celebrant/People dialogue for 8am.

    Renders the simplified three-line form:
        Celebrant   Lord, have mercy.
        People      Christ, have mercy.   (bold)
        Celebrant   Lord, have mercy.
    """
    add_celebrant_line(doc, "Celebrant", "Lord, have mercy.")
    add_people_line(doc, "People", "Christ, have mercy.")
    add_celebrant_line(doc, "Celebrant", "Lord, have mercy.")
