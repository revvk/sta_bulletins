"""
Burial / funeral / memorial section module.

Composes the funeral liturgy from a per-service ``FuneralData``
dataclass, the ``funeral_texts.yaml`` BCP texts (both rites), the
existing scripture cache, and the existing song-lookup pipeline.

The top-level ``add_burial_service`` walks the
HC × Commendation × Committal matrix from
docs/plans/funeral_bulletins.md and emits the appropriate sequence
of section blocks. Most section helpers are private (``_add_*``).

Pronoun + name placeholders ({N}, {he}, {him}, {his}, {brother}) are
substituted via ``FuneralData.substitutions`` whenever the renderer
emits text drawn from funeral_texts.yaml.
"""

from __future__ import annotations

from docx import Document

from bulletin.data.loader import (
    load_common_prayers,
    load_eucharistic_prayers,
    load_funeral_texts,
    load_proper_prefaces,
    load_special_prayers,
)
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line, add_dialogue, add_hymn_header,
    add_lyric_verse, add_lyric_chorus,
)
from bulletin.document.sections.word_of_god import (
    add_psalm, add_song_smart, add_body_with_amen,
    _add_reading_text,
)
from bulletin.document.sections.holy_communion import (
    add_prayer_a_or_b, add_prayer_c, add_sanctus_text,
    add_doxology_amen,
)
from bulletin.sources.funeral_data import FuneralData


# =====================================================================
# Top-level entry point
# =====================================================================

def add_burial_service(doc: Document, fd: FuneralData,
                        scripture_readings: dict, song_lookup_fn) -> None:
    """Render the body of the funeral bulletin into ``doc``.

    ``doc`` already has the cover applied (placeholders substituted),
    page setup configured, and styles registered. This function adds
    everything from the start of the Liturgy of the Word through the
    end-of-service block (commendation / committal / blessing close,
    plus interment notice / reception / participants page).
    """
    rite_texts = load_funeral_texts()[f"rite_{fd.rite}"]
    common = load_common_prayers()

    # Each section helper handles its own internal spacing; we insert
    # a separating spacer between major sections here so the bulletin
    # has visible breathing room at every section boundary (matches
    # the Sunday bulletin conventions).

    # ----- Liturgy of the Word ----------------------------------------
    _add_seating_and_prelude(doc)
    add_spacer(doc)
    _add_opening(doc, fd, rite_texts, song_lookup_fn)
    add_spacer(doc)
    _add_collect(doc, fd, rite_texts)
    add_spacer(doc)
    _add_readings_block(doc, fd, rite_texts, scripture_readings,
                         song_lookup_fn)
    add_spacer(doc)
    _add_homily_block(doc, fd, song_lookup_fn)
    add_spacer(doc)
    _add_apostles_creed(doc, fd, common)
    add_spacer(doc)
    _add_prayers_for_departed(doc, fd, rite_texts)
    add_spacer(doc)
    _add_special_prayers_block(doc, fd)
    add_spacer(doc)
    _add_peace(doc, fd)

    # ----- End-of-service: HC × Commendation × Committal matrix -------
    if fd.holy_eucharist["enabled"]:
        add_spacer(doc)
        _add_holy_communion(doc, fd, rite_texts, common, song_lookup_fn)

    if fd.include_commendation:
        add_spacer(doc)
        _add_commendation(doc, fd, rite_texts)

    if fd.include_committal:
        add_spacer(doc)
        _add_procession_hymn_if_present(doc, fd, song_lookup_fn)
        add_spacer(doc)
        _add_committal(doc, fd, rite_texts, common)
    else:
        # No committal in this bulletin — close with the standard
        # blessing/closing-hymn/dismissal block.
        add_spacer(doc)
        _add_blessing_close(doc, fd, rite_texts, song_lookup_fn)

    _add_postlude(doc)

    # ----- Optional notices + participants ----------------------------
    if fd.interment_notice:
        _add_interment_notice(doc, fd.interment_notice)
    if fd.reception.get("shown"):
        _add_reception_notice(doc, fd.reception.get("text", ""))
    # Participants are rendered on the dedicated funeral back-cover
    # template (templates/back_cover_funeral.docx, see
    # funeral_builder._inject_participants), not inline in the body.


# =====================================================================
# Helpers — substitution + style
# =====================================================================

def _sub(text: str, fd: FuneralData) -> str:
    """Substitute {N}/{he}/{him}/{his}/{brother} placeholders. The
    {NN} placeholder (bereaved family-member names) is filled
    separately in the for-those-who-mourn collect helper."""
    if not text:
        return text
    out = text
    for k, v in fd.substitutions.items():
        out = out.replace(k, v)
    return out


def _flow(text: str) -> str:
    """Collapse a multi-line block scalar into single-paragraph text
    (newlines → spaces). Used for prayer paragraphs that should
    render as one flowing paragraph in the bulletin."""
    if not text:
        return ""
    return " ".join(text.split())


def _add_unison_prayer(doc: Document, text: str) -> None:
    """Render a prayer said in unison by Celebrant + People — bold, in
    the Body - People Recitation style. Matches the Sunday pattern
    used for the Lord's Prayer and Post-communion Prayer.
    """
    if not text:
        return
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]


def _join_preface(opening: str, preface_text: str) -> str:
    """Join the BCP preface_opening with a Proper Preface insert into
    a single grammatical paragraph — matches the Sunday code's logic
    in add_prayer_a_or_b. Most prefaces start with connecting words
    ("Through", "For", "Because") and join with a comma; otherwise
    keep the opening's period and start a new sentence.
    """
    if not preface_text:
        return opening
    _CONNECTING_WORDS = {"through", "for", "because", "but", "who"}
    first_word = preface_text.split()[0].lower().rstrip(".,;")
    opening = opening.rstrip()
    if first_word in _CONNECTING_WORDS:
        opening_no_period = opening.rstrip(".")
        joined = preface_text[0].lower() + preface_text[1:]
        return f"{opening_no_period}, {joined}"
    return f"{opening} {preface_text}"


def _add_creed_lines(doc: Document, lines: list[str], doc_obj=None) -> None:
    """Render a creed as a single paragraph with line breaks. Empty
    strings in `lines` become paragraph breaks. The Apostles' Creed
    is said in unison, so every run is bold via the People char
    style (matches the Lord's Prayer / Postcommunion treatment).
    """
    people_style = doc.styles["People"]
    paragraph = doc.add_paragraph(style="Body - People Recitation (Creed)")
    first = True
    for line in lines:
        if line == "":
            paragraph = doc.add_paragraph(style="Body - People Recitation (Creed)")
            first = True
            continue
        if not first:
            run = paragraph.add_run()
            run.add_break()
        # The docx Body - People Recitation (Creed) style provides the
        # hanging indent; leading spaces in `line` visually nudge
        # nested lines further. Bolding via People style for unison.
        run = paragraph.add_run(line)
        run.style = people_style
        first = False


# =====================================================================
# Section: seating + prelude
# =====================================================================

def _add_seating_and_prelude(doc: Document) -> None:
    """The 'Seating of the Family' line + Prelude heading. Both Cox and
    Owens use this exact lead-in; Stuhler skips it (no body / family
    procession)."""
    add_heading2(doc, "Seating of the Family")
    add_spacer(doc)
    add_heading2(doc, "Prelude")
    add_spacer(doc)


# =====================================================================
# Section: opening anthem + opening hymn
# =====================================================================

def _add_opening(doc: Document, fd: FuneralData, rite_texts: dict,
                 song_lookup_fn) -> None:
    """Render the Opening Anthem (BCP text or sung opening) and the
    Opening Hymn."""
    add_introductory_rubric(doc, "Please stand.")

    anthem_choice = fd.music.get("opening_anthem", "from_bcp")
    if anthem_choice == "from_bcp" or anthem_choice is None:
        # BCP-text opening anthem from funeral_texts.yaml.
        anthems = rite_texts["opening_anthems"]
        # Default to whichever anthem entry comes first; the YAML
        # convention puts the standard "I am Resurrection / I am the
        # resurrection" anthem first in both rites.
        first_key = next(iter(anthems))
        anthem = anthems[first_key]
        add_heading2(doc, anthem.get("label", "Opening Anthem"))
        for stanza in anthem.get("stanzas", []):
            for line in stanza.strip().split("\n"):
                add_body(doc, line)
            add_spacer(doc)
    else:
        # Sung opening — look up via the song catalog.
        add_heading2(doc, "Opening Anthem")
        song = song_lookup_fn(anthem_choice, "11 am")
        add_song_smart(doc, song)

    # Opening Hymn
    if fd.music.get("opening_hymn"):
        add_heading2(doc, "Opening Hymn")
        song = song_lookup_fn(fd.music["opening_hymn"], "11 am")
        add_song_smart(doc, song)


# =====================================================================
# Section: The Collect
# =====================================================================

def _add_collect(doc: Document, fd: FuneralData, rite_texts: dict) -> None:
    """Render the collect dialogue, the chosen collect, and (Rite II
    only) the optional for-those-who-mourn addendum."""
    add_heading2(doc, "The Collect")

    # Dialogue ("The Lord be with you. / And [also/with thy spirit]…")
    for entry in rite_texts["collect_dialogue"]:
        if "people" in entry:
            add_celebrant_line(doc, "Celebrant", entry["celebrant"])
            add_people_line(doc, "People", entry["people"])
        else:
            add_celebrant_line(doc, "Celebrant", entry["celebrant"])
    add_spacer(doc)

    collects = rite_texts["collects"]

    # Rite II: three Adult options (per fd.collects.adult.choice) plus
    # one Child collect plus an optional for-those-who-mourn addendum.
    # Rite I: a single Adult collect and a single Child collect, no
    # addendum, no choice field.
    if fd.rite == "II":
        choice = (fd.collects.get("adult") or {}).get("choice") or "option_3"
        adult_entry = collects["adult"][choice]
        text = _sub(_flow(adult_entry["text"]), fd)
        add_body_with_amen(doc, text + " " + adult_entry["amen"])

        if fd.collects.get("add_for_those_who_mourn"):
            mourn = collects["for_those_who_mourn"]
            text = _flow(mourn["text"])
            # {NN} is the bereaved family-member names. Schema doesn't
            # define a dedicated field yet; fall back to the deceased's
            # full name so the renderer never emits the literal
            # placeholder.
            text = text.replace("{NN}", fd.deceased.get("full_name", ""))
            add_spacer(doc)
            add_body_with_amen(doc, text + " " + mourn["amen"])
    else:
        # Rite I: only one Adult option in the YAML; no nested choice.
        adult_entry = collects["adult"]
        text = _sub(_flow(adult_entry["text"]), fd)
        add_body_with_amen(doc, text + " " + adult_entry["amen"])


# =====================================================================
# Section: readings + psalm + sequence + gospel
# =====================================================================

def _add_readings_block(doc: Document, fd: FuneralData, rite_texts: dict,
                         scripture: dict, song_lookup_fn) -> None:
    """Render the readings + psalm + sequence hymn + gospel block."""
    add_introductory_rubric(doc, "Be seated.")

    # Whether this rite uses Rite I labels ("The First Reading: …",
    # "Here endeth the Reading.") or Rite II labels ("First Reading: …",
    # "The Word of the Lord. / Thanks be to God.").
    rite_i = fd.rite == "I"
    label_prefix = "The " if rite_i else ""

    # First Reading
    first_ref = fd.readings.get("first")
    if first_ref:
        _add_funeral_reading(doc, scripture, first_ref,
                              label_prefix + "First Reading", rite_i)

    # Psalm
    psalm_ref = fd.readings.get("psalm")
    if psalm_ref:
        add_spacer(doc)
        psalm = scripture.get(psalm_ref)
        rubric = _psalm_rubric(fd.readings.get("psalm_mode") or "unison")
        if psalm is not None:
            add_psalm(doc, psalm_ref, rubric, psalm)
        else:
            add_heading2(doc, psalm_ref)
            add_rubric(doc, rubric)
            add_body(doc, "[Psalm text not fetched]")

    # Second Reading (optional)
    second_ref = fd.readings.get("second")
    if second_ref:
        add_spacer(doc)
        _add_funeral_reading(doc, scripture, second_ref,
                              label_prefix + "Second Reading", rite_i)

    # Sequence Hymn
    seq = fd.music.get("sequence_hymn")
    if seq:
        add_spacer(doc)
        add_introductory_rubric(doc, "Please stand.")
        add_heading2(doc, "Sequence Hymn")
        song = song_lookup_fn(seq, "11 am")
        add_song_smart(doc, song)

    # Gospel
    gospel_ref = fd.readings.get("gospel")
    if gospel_ref:
        add_spacer(doc)
        _add_funeral_gospel(doc, scripture, gospel_ref, rite_i)


def _add_funeral_reading(doc: Document, scripture: dict, ref: str,
                          heading_label: str, rite_i: bool) -> None:
    """Render a non-gospel scripture reading with rite-appropriate
    response."""
    reading = scripture.get(ref)
    add_heading2(doc, f"{heading_label}: {ref}")
    if reading is not None:
        _add_reading_text(doc, reading)
    else:
        add_body(doc, f"[Reading text for {ref} not fetched]")
    add_spacer(doc)
    add_rubric(doc, "After the reading, the Reader will say")
    if rite_i:
        # Rite I: just "Here endeth the Reading." — no people response.
        add_celebrant_line(doc, "", "Here endeth the Reading.")
    else:
        add_celebrant_line(doc, "", "The Word of the Lord.")
        add_people_line(doc, "People", "Thanks be to God.")


def _add_funeral_gospel(doc: Document, scripture: dict, ref: str,
                         rite_i: bool) -> None:
    """Render the Gospel reading with rite-appropriate announcement
    and response."""
    book = ref.split()[0]
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, f"Gospel: {ref}")
    add_rubric(doc, "The Deacon or a Priest reads the Gospel, first saying")

    add_celebrant_line(doc, "", f"The Holy Gospel of our Lord Jesus Christ according to {book}.")
    if rite_i:
        add_people_line(doc, "People", "Glory be to thee, O Lord.")
    else:
        add_people_line(doc, "People", "Glory to you, Lord Christ.")

    reading = scripture.get(ref)
    if reading is not None:
        _add_reading_text(doc, reading)
    else:
        add_body(doc, f"[Gospel text for {ref} not fetched]")

    add_spacer(doc)
    add_rubric(doc, "After the Gospel, the Reader will say")
    add_celebrant_line(doc, "", "The Gospel of the Lord.")
    if rite_i:
        add_people_line(doc, "People", "Praise be to thee, O Christ.")
    else:
        add_people_line(doc, "People", "Praise to you, Lord Christ.")


def _psalm_rubric(mode: str) -> str:
    """Map psalm_mode to the printable rubric line."""
    return {
        "unison":        "Read in unison.",
        "responsive":    "Read responsively by whole verse.",
        "antiphonal":    "Read antiphonally.",
        "by_half_verse": "Read responsively by half verse.",
    }.get(mode, "Read in unison.")


# =====================================================================
# Section: Words of Remembrance + post-WoR / post-homily anthems +
# Homily
# =====================================================================

def _add_homily_block(doc: Document, fd: FuneralData, song_lookup_fn) -> None:
    """Render Words of Remembrance (if any), the post-WoR hymn/anthem
    (if any), the Homily, and the post-Homily hymn/anthem (if any)."""
    add_introductory_rubric(doc, "Be seated.")

    # Words of Remembrance — one entry per speaker.
    speakers = fd.participants.get("words_of_remembrance") or []
    if speakers:
        add_heading2(doc, "Words of Remembrance")
        for name in speakers:
            add_body(doc, name)
        add_spacer(doc)

    # Post-Words-of-Remembrance music (hymn OR anthem; both optional).
    rh = fd.music.get("remembrance_hymn")
    ra = fd.music.get("remembrance_anthem")
    if rh:
        add_spacer(doc)
        add_heading2(doc, "Hymn")
        song = song_lookup_fn(rh, "11 am")
        add_song_smart(doc, song)
    elif ra:
        add_spacer(doc)
        _add_anthem_block(doc, "Anthem", ra)

    # Homily
    add_spacer(doc)
    add_heading2(doc, "Homily")
    homilist = fd.participants.get("celebrant", "")
    if homilist:
        add_body(doc, _format_clergy_title(homilist))
    add_spacer(doc)

    # Post-Homily music (hymn OR anthem; both optional).
    hh = fd.music.get("homily_hymn")
    ha = fd.music.get("homily_anthem")
    if hh:
        add_spacer(doc)
        add_heading2(doc, "Hymn")
        song = song_lookup_fn(hh, "11 am")
        add_song_smart(doc, song)
    elif ha:
        add_spacer(doc)
        _add_anthem_block(doc, "Anthem", ha)


def _add_anthem_block(doc: Document, heading: str, anthem: dict) -> None:
    """Render an anthem (title + composer/arranger + soloist credit).

    Anthems are not congregational — only the title line is printed,
    with the composer right-aligned in small caps. If a soloist is
    listed, their name appears italic on the line below.
    """
    add_heading2(doc, heading)
    title = anthem.get("title", "")
    arranger = anthem.get("arranger") or ""
    p = doc.add_paragraph(style="Body")
    p.add_run(title)
    if arranger:
        p.add_run("\t")
        run = p.add_run(arranger)
        # Small-caps via the existing Body style's run formatting; if a
        # dedicated style is later added we'll switch to it.
        run.font.small_caps = True
    soloist = anthem.get("soloist")
    if soloist:
        p2 = doc.add_paragraph(style="Body")
        run = p2.add_run(soloist)
        run.italic = True


def _format_clergy_title(name: str) -> str:
    """Add 'The Rev.' prefix when missing — matches the convention
    used in the Sunday bulletins for the homilist line."""
    n = name.strip()
    if n.lower().startswith(("the rev", "the most rev", "the rt.", "the rt rev",
                              "fr.", "father ", "mtr.", "mother ")):
        return n
    return f"The Rev. {n}"


# =====================================================================
# Section: Apostles' Creed
# =====================================================================

def _add_apostles_creed(doc: Document, fd: FuneralData, common: dict) -> None:
    """Render the Apostles' Creed appropriate to the rite.

    The "In the assurance of eternal life…" line is the Celebrant's
    introduction to the Creed (BCP p. 496) — set as a celebrant
    line (regular body weight, hanging indent), NOT as an italic
    rubric. A blank line separates it from the Creed text itself
    (matching the Cox bulletin's typesetting).
    """
    key = "apostles_creed_rite_i" if fd.rite == "I" else "apostles_creed_rite_ii"
    creed = common[key]
    add_heading2(doc, "Apostles' Creed")
    add_introductory_rubric(doc, "All stand, the Celebrant introduces the Creed, saying")
    # Empty label — the intro is unattributed in the bulletin; the
    # rubric above identifies it as the Celebrant's. The Body -
    # Dialogue style's hanging indent + tab gives the right visual.
    add_celebrant_line(doc, "", _flow(creed["intro"]))
    add_spacer(doc)
    _add_creed_lines(doc, creed["lines"])


# =====================================================================
# Section: Prayers for the Departed
# =====================================================================

def _add_prayers_for_departed(doc: Document, fd: FuneralData,
                                rite_texts: dict) -> None:
    """Render the Prayers for the Departed in the rite-appropriate
    shape."""
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Prayers for the Departed")

    pop = rite_texts["prayers_for_the_departed"]
    flags = fd.prayers_for_the_departed or {}

    if fd.rite == "II":
        _add_prayers_for_departed_rite_ii(doc, fd, pop, flags)
    else:
        _add_prayers_for_departed_rite_i(doc, fd, pop, flags)


def _add_prayers_for_departed_rite_ii(doc: Document, fd: FuneralData,
                                        pop: dict, flags: dict) -> None:
    """Render Rite II's 'Hear us, Lord' litany."""
    add_body(doc, _sub(_flow(pop["bidding"]), fd))
    add_spacer(doc)

    # Optional petitions are flagged in YAML by `omit_when` markers.
    inc_baptism = flags.get("include_baptism_petition", True)
    inc_communion = flags.get("include_communion_petition",
                                fd.holy_eucharist.get("enabled", False))

    for petition in pop["petitions"]:
        if petition.get("optional"):
            omit = petition.get("omit_when") or []
            if "unbaptized" in omit and not inc_baptism:
                continue
            if "no_communion" in omit and not inc_communion:
                continue
        add_body(doc, _sub(_flow(petition["text"]), fd))
        if petition.get("response"):
            # The response ("Hear us, Lord.") is said by the
            # congregation. Use the People char style so the run
            # uses the actual Adobe Garamond Pro Bold font — setting
            # `run.bold = True` on a regular Adobe Garamond run
            # gives Word's faux-bold synthesis, which looks subtly
            # different from the real bold face.
            p = doc.add_paragraph(style="Body")
            run = p.add_run(petition["response"])
            run.style = doc.styles["People"]
        add_spacer(doc)

    add_rubric(doc, pop["silence_rubric"])
    add_spacer(doc)
    # The renderer always prints one of the BCP's concluding prayers
    # (controlled by the `conclusion` flag in the per-service YAML),
    # so the rubric immediately above it can be definitive:
    # "with the following prayer:" rather than the BCP's
    # "with one of the following or some other prayer".
    add_rubric(doc, "The Celebrant concludes with the following prayer:")

    conclusion_key = flags.get("conclusion") or "commend"
    conclusion = pop["conclusions"][conclusion_key]
    add_body_with_amen(doc, _sub(_flow(conclusion["text"]), fd) + " " + conclusion["amen"])


def _add_prayers_for_departed_rite_i(doc: Document, fd: FuneralData,
                                       pop: dict, flags: dict) -> None:
    """Render Rite I's long-form intercessions, each ending in Amen.

    The two preamble rubrics ("The People respond to every petition
    with Amen." + "The Deacon or other leader says") are printed on
    a single line — the BCP prints them as two separate rubric
    paragraphs, but Andrew flagged the two-line treatment as
    unnecessarily wordy. A single italic line "The People respond
    to every petition with Amen. The Deacon or other leader says"
    captures both rubrics without the extra vertical space.
    """
    response_rubric = pop["response_rubric"].rstrip(".")
    bidding_rubric = pop["bidding_rubric"]
    add_rubric(doc, f"{response_rubric}. {bidding_rubric}")
    add_body(doc, pop["bidding"])
    add_spacer(doc)

    include_optional = set(flags.get("include_optional_petitions") or [6, 7, 8, 9])
    patron_phrase = flags.get("patron_phrase") or ""

    petitions_to_emit = [
        (idx, p) for idx, p in enumerate(pop["petitions"], start=1)
        if not (p.get("optional") and idx not in include_optional)
    ]
    for i, (idx, petition) in enumerate(petitions_to_emit):
        text = _flow(petition["text"])
        text = _sub(text, fd)
        text = text.replace("{patron_phrase}", patron_phrase)
        add_body_with_amen(doc, text + " " + petition["amen"])
        # Spacer between petitions only — no trailing spacer after
        # the concluding petition (which is itself a collect), so
        # there isn't an extra blank line stacked on top of The
        # Peace heading's space-before.
        if i < len(petitions_to_emit) - 1:
            add_spacer(doc)


# =====================================================================
# Section: Special prayers (e.g. Daughters of the King)
# =====================================================================

def _add_special_prayers_block(doc: Document, fd: FuneralData) -> None:
    """Render any opt-in special prayers (currently all anchored at
    `after_prayers_for_departed`)."""
    if not fd.special_prayers:
        return
    library = load_special_prayers()
    for i, entry in enumerate(fd.special_prayers):
        key = entry.get("key")
        if not key or key not in library:
            continue
        if i > 0:
            add_spacer(doc)
        prayer = library[key]
        add_heading2(doc, prayer["title"])
        # Reader credit (italic rubric line)
        if entry.get("reader") and prayer.get("credit_format"):
            credit = prayer["credit_format"].format(
                reader=entry.get("reader", ""),
                reader_role=entry.get("reader_role", ""),
            )
            add_rubric(doc, credit)
        text = _sub(_flow(prayer["text"]), fd)
        add_body_with_amen(doc, text + " " + prayer.get("amen", "Amen."))


# =====================================================================
# Section: The Peace
# =====================================================================

def _add_peace(doc: Document, fd: FuneralData) -> None:
    """Render the Peace exchange in the rite-appropriate form."""
    add_heading2(doc, "The Peace")
    add_rubric(doc, "The Celebrant says to the people")
    add_celebrant_line(doc, "Celebrant", "The peace of the Lord be always with you.")
    if fd.rite == "I":
        add_people_line(doc, "People", "And with thy spirit.")
    else:
        add_people_line(doc, "People", "And also with you.")


# =====================================================================
# Section: Holy Communion (when fd.holy_eucharist.enabled)
# =====================================================================

def _add_holy_communion(doc: Document, fd: FuneralData, rite_texts: dict,
                         common: dict, song_lookup_fn) -> None:
    """Render the HC section: offertory anthem, Great Thanksgiving,
    Lord's Prayer, breaking, invitation, communion music, postcommunion
    prayer."""
    eucharistic_prayers = load_eucharistic_prayers()

    add_spacer(doc)
    add_heading(doc, "The Holy Communion")
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # Offertory music — either a soloist/choir anthem (small-caps
    # composer credit, optional soloist) OR a congregational hymn
    # (title + tab + hymnal reference, full lyrics if any). Mirrors
    # the remembrance_hymn / remembrance_anthem pair: at most one
    # of the two fields should be set per service.
    offertory_hymn = fd.music.get("offertory_hymn")
    offertory_anthem = fd.music.get("offertory_anthem") or {}
    if offertory_hymn:
        add_heading2(doc, "Offertory Hymn")
        song = song_lookup_fn(offertory_hymn, "11 am")
        add_song_smart(doc, song)
    elif offertory_anthem.get("title"):
        _add_anthem_block(doc, "Offertory Anthem As the Altar is Prepared",
                          offertory_anthem)

    # Great Thanksgiving
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "The Great Thanksgiving")

    prayer_key = fd.holy_eucharist.get("prayer", "A")
    prayer_yaml_key = "prayer_" + (prayer_key.lower())
    prayer = eucharistic_prayers.get(prayer_yaml_key, {})

    # The Proper Preface inserted between the opening and the Sanctus
    # transition for funerals is "Of the Commemoration of the Dead"
    # (BCP p. 382) for both rites. Cached here for the helpers below.
    prefaces = load_proper_prefaces()
    commemoration_preface = _flow(
        (prefaces.get("commemoration_of_the_dead") or {}).get("text") or ""
    )

    # ----- Rite II Prayer A / B / C: defer to the Sunday code path,
    # which already handles the preface-joining grammar (connecting
    # words like "Through" join to the preface_opening with a comma)
    # and inserts the right spacers between every paragraph of the
    # prayer body. We synthesize the small `data` dict the Sunday
    # functions expect.
    if fd.rite == "II" and prayer_key.upper() in ("A", "B", "C"):
        ep_data_funeral = {
            "preface_opening":     eucharistic_prayers.get("preface_opening", ""),
            "sanctus_transition":  eucharistic_prayers.get("sanctus_transition", ""),
            "prayer_a":            eucharistic_prayers.get("prayer_a", {}),
            "prayer_b":            eucharistic_prayers.get("prayer_b", {}),
            "prayer_c":            eucharistic_prayers.get("prayer_c", {}),
        }
        # Sursum Corda dialogue
        for entry in eucharistic_prayers.get("sursum_corda", []):
            add_celebrant_line(doc, "Celebrant", entry["celebrant"])
            add_people_line(doc, "People", entry["people"])
        add_spacer(doc)

        sunday_data = {
            "service_time":         "11 am",
            "proper_preface_text":  commemoration_preface,
            "sanctus_song":         None,    # use the spoken text Sanctus
        }
        add_rubric(doc, "Then, facing the Holy Table, the Celebrant proceeds")
        if prayer_key.upper() == "C":
            add_prayer_c(doc, ep_data_funeral, sunday_data, common)
        else:
            add_prayer_a_or_b(doc, ep_data_funeral, sunday_data, common,
                              prayer_key.upper())
    else:
        # ----- Rite II Prayer D and Rite I Prayers I / II: the Sunday
        # helpers don't handle these shapes; we render them inline,
        # matching the structure of each prayer's eucharistic_prayers.yaml
        # entry. Spacers go between every paragraph for parity with the
        # Sunday output.
        if fd.rite == "I":
            for entry in (
                ("The Lord be with you.",   "And with thy spirit."),
                ("Lift up your hearts.",    "We lift them up unto the Lord."),
                ("Let us give thanks unto our Lord God.", "It is meet and right so to do."),
            ):
                add_celebrant_line(doc, "Celebrant", entry[0])
                add_people_line(doc, "People", entry[1])
            add_spacer(doc)
            add_rubric(doc, "Then, facing the Holy Table, the Celebrant proceeds")
            opening = _flow(eucharistic_prayers.get("preface_opening_rite_i", ""))
            add_body(doc, _join_preface(opening, commemoration_preface))
            add_spacer(doc)
            add_body(doc, _flow(eucharistic_prayers.get("sanctus_transition_rite_i", "")))
            add_spacer(doc)
            add_rubric(doc, "Celebrant and People")
            add_sanctus_text(doc, common.get("sanctus_rite_i", common.get("sanctus", [])))
        else:
            # Prayer D — its own three-paragraph preface, no insert.
            for entry in eucharistic_prayers.get("sursum_corda", []):
                add_celebrant_line(doc, "Celebrant", entry["celebrant"])
                add_people_line(doc, "People", entry["people"])
            add_spacer(doc)
            add_rubric(doc, "Then, facing the Holy Table, the Celebrant proceeds")
            for k in ("preface_1", "preface_2", "preface_3"):
                if prayer.get(k):
                    add_body(doc, _flow(prayer[k]))
                    add_spacer(doc)
            add_rubric(doc, "Celebrant and People")
            add_sanctus_text(doc, common.get("sanctus", []))

        add_spacer(doc)
        # Post-Sanctus body — custom render with spacers between
        # every paragraph (Prayer D / Rite I shapes).
        _add_eucharistic_prayer_body(doc, prayer, fd)

    # Lord's Prayer — said in unison; render bold via the People
    # character style (matches the Sunday convention).
    add_spacer(doc)
    intro = common.get("lords_prayer_intro", {}).get("option_1",
            "And now, as our Savior Christ has taught us, we are bold to say,")
    add_body(doc, intro)
    lp_text = " ".join(line.strip() for line in common.get("lords_prayer", []))
    _add_unison_prayer(doc, lp_text)

    # Breaking of the Bread
    add_spacer(doc)
    add_heading2(doc, "The Breaking of the Bread")
    fraction = common.get("fraction", {}).get("with_alleluia") or {}
    add_celebrant_line(doc, "Celebrant", fraction.get("celebrant", ""))
    add_people_line(doc, "People", fraction.get("people", ""))

    # Invitation
    add_spacer(doc)
    add_heading2(doc, "The Invitation")
    add_rubric(doc, "Facing the people, the Celebrant says the following Invitation")
    invitation = common.get("invitation_to_communion", "") + " " + common.get("invitation_addition", "")
    add_body(doc, _flow(invitation))

    # Communion music
    comms = fd.music.get("communion_music") or []
    if comms:
        add_spacer(doc)
        add_heading2(doc, "Communion Music")
        for i, title in enumerate(comms):
            if i > 0:
                add_spacer(doc)
            song = song_lookup_fn(title, "11 am")
            add_song_smart(doc, song)

    # Postcommunion Prayer — also said in unison.
    add_spacer(doc)
    pc = rite_texts["postcommunion_prayer"]
    add_heading2(doc, pc.get("label", "Postcommunion Prayer"))
    for entry in pc.get("intro_dialogue", []):
        add_celebrant_line(doc, "Celebrant", entry.get("celebrant", ""))
    pc_text = _flow(pc["text"]) + " " + pc.get("amen", "Amen.")
    _add_unison_prayer(doc, pc_text)


def _add_eucharistic_prayer_body(doc: Document, prayer: dict,
                                   fd: FuneralData) -> None:
    """Render the post-Sanctus body of a eucharistic prayer.

    Each prayer (A, B, D, I, II) has a slightly different shape; we
    render whatever keys are present in canonical order so each
    prayer's structure dictates the layout.
    """
    # Order of keys to emit, in canonical order. Missing keys are
    # silently skipped — every prayer has a different subset.
    order = [
        "post_sanctus", "post_sanctus_1", "post_sanctus_2", "post_sanctus_3",
        "institution_1",
        "institution_rubric",
        "institution_bread",
        "institution_cup",
        "anamnesis",
        "memorial_intro",
        # memorial_acclamation is rendered separately because it's a
        # list of lines spoken by Celebrant + People together.
        "epiclesis",
        "epiclesis_1", "epiclesis_2",
        "sacrifice",
        "self_offering",
        "unworthiness",
        "intercession_main",
        # intercession_optional / intercession_close are nested dicts
        # with their own rendering rules; we render them inline below.
    ]

    # Helper closure: emit one paragraph + a spacer afterwards. The
    # spacer is what gives every prayer paragraph the same visual
    # breathing room as the Sunday output.
    def emit(text: str) -> None:
        if text:
            add_body(doc, _flow(text))
            add_spacer(doc)

    for key in order:
        if key not in prayer:
            continue
        val = prayer[key]
        if isinstance(val, str):
            emit(val)
            # Memorial intro flows into the acclamation
            if key == "memorial_intro" and "memorial_acclamation" in prayer:
                add_rubric(doc, prayer.get("memorial_acclamation_rubric",
                                            "Celebrant and People"))
                p = doc.add_paragraph(style="Body")
                for i, line in enumerate(prayer["memorial_acclamation"]):
                    if i > 0:
                        p.add_run().add_break()
                    run = p.add_run(line)
                    run.bold = True
                add_spacer(doc)
        elif isinstance(val, dict):
            # e.g. intercession sub-trees with text variants
            if "text" in val:
                emit(val["text"])

    # Intercession close (optional, has text_with_saints / _without_saints)
    if "intercession_close" in prayer:
        ic = prayer["intercession_close"]
        text = ic.get("text_without_saints") or ic.get("text_with_saints") or ""
        emit(text)

    # Doxology + bold AMEN, matching the Sunday `add_doxology_amen`.
    if "doxology" in prayer:
        add_doxology_amen(doc, prayer)


# =====================================================================
# Section: The Commendation
# =====================================================================

def _add_commendation(doc: Document, fd: FuneralData,
                       rite_texts: dict) -> None:
    """Render the Commendation: heading, intro rubric, anthem, prayer.

    Anthem layout matches the spoken / responsive convention Andrew
    spelled out:
      • First refrain — three lines. Line 1 is said by the
        Celebrant (regular type); lines 2-3 are the People's
        response (Adobe Garamond Pro **Bold** via the People char
        style, not Word's faux-bold).
      • Middle verse — collapsed to a single flowing paragraph,
        said by the Celebrant.
      • Second refrain — all three lines bold (the People join in).

    Applies to both rites: the only thing that differs between
    rite_I.commendation and rite_II.commendation is the BCP wording
    (thy/your, thee/you, etc.).
    """
    cm = rite_texts["commendation"]
    add_heading(doc, cm.get("label", "The Commendation"))

    # The two BCP rubrics ("The Celebrant and other ministers take
    # their places at the body." and "This anthem, or some other
    # suitable anthem, or a hymn, may be sung or said.") are
    # combined into a single Body - Introductory Rubric paragraph.
    setup = (cm.get("setup_rubric") or "").strip()
    anthem_rubric = (cm.get("anthem", {}).get("anthem_rubric") or "").strip()
    intro_parts = [s for s in (setup, anthem_rubric) if s]
    if intro_parts:
        add_introductory_rubric(doc, " ".join(intro_parts))

    anthem = cm.get("anthem", {})
    refrain_lines = anthem.get("refrain", "").strip().split("\n")
    verse = anthem.get("verse", "")
    people_style = doc.styles["People"]

    if refrain_lines:
        # First refrain — line 1 Celebrant, lines 2-3 People (bold).
        add_body(doc, refrain_lines[0])
        for line in refrain_lines[1:]:
            p = doc.add_paragraph(style="Body")
            run = p.add_run(line)
            run.style = people_style
        add_spacer(doc)

    # Middle verse — collapsed to one flowing paragraph, Celebrant.
    if verse:
        add_body(doc, _flow(verse))
        add_spacer(doc)

    # Refrain repeat — every line bold (People join in throughout).
    for line in refrain_lines:
        p = doc.add_paragraph(style="Body")
        run = p.add_run(line)
        run.style = people_style

    # Commendation prayer.
    prayer = cm.get("prayer", {})
    add_spacer(doc)
    if prayer.get("prayer_rubric"):
        add_rubric(doc, prayer["prayer_rubric"])
    text = _sub(_flow(prayer.get("text", "")), fd)
    add_body_with_amen(doc, text + " " + prayer.get("amen", "Amen."))


# =====================================================================
# Section: Procession hymn (only when committal is at the church)
# =====================================================================

def _add_procession_hymn_if_present(doc: Document, fd: FuneralData,
                                       song_lookup_fn) -> None:
    """When the committal happens at the church (columbarium etc.), a
    procession hymn covers the move from sanctuary to the burial spot.
    """
    title = fd.music.get("procession_hymn")
    if not title:
        return
    add_heading2(doc, "Hymn")
    rubric = (
        "During the hymn, the congregation will move to the columbarium "
        "in the Narthex (entrance area) of the church. The clergy and "
        "ministers will follow the cross, and the family will follow "
        "behind them. After the family, please join the procession out "
        "for the Committal."
    )
    add_rubric(doc, rubric)
    song = song_lookup_fn(title, "11 am")
    add_song_smart(doc, song)


# =====================================================================
# Section: The Committal
# =====================================================================

def _add_committal(doc: Document, fd: FuneralData, rite_texts: dict,
                    common: dict) -> None:
    """Render the Committal section (anthem + earth-cast prayer + Lord's
    Prayer + Rest eternal + dismissal blessing)."""
    cm = rite_texts["committal"]
    add_spacer(doc)
    add_heading(doc, cm.get("label", "The Committal"))
    add_spacer(doc)

    anthem = cm.get("anthem", {})
    if anthem.get("anthem_rubric"):
        add_rubric(doc, anthem["anthem_rubric"])
    for stanza in anthem.get("stanzas", []):
        for line in stanza.strip().split("\n"):
            add_body(doc, line)
        add_spacer(doc)

    # Earth-cast committal prayer
    prayer = cm.get("prayer", {})
    if prayer.get("prayer_rubric"):
        add_rubric(doc, prayer["prayer_rubric"])
    text = _sub(_flow(prayer.get("text", "")), fd)
    text = text.replace("{location}", _committal_location_word(fd))
    add_body_with_amen(doc, text + " " + prayer.get("amen", "Amen."))

    # Lord's Prayer — said in unison; bold via the People char style.
    add_spacer(doc)
    for entry in cm.get("lords_prayer_dialogue", []):
        if "people" in entry:
            add_celebrant_line(doc, "Celebrant", entry["celebrant"])
            add_people_line(doc, "People", entry["people"])
        else:
            add_celebrant_line(doc, "Celebrant", entry["celebrant"])
    lp_text = " ".join(line.strip() for line in common.get("lords_prayer", []))
    _add_unison_prayer(doc, lp_text)

    # Rest eternal antiphon
    re = cm.get("rest_eternal", {})
    if re:
        add_spacer(doc)
        if re.get("rubric"):
            add_rubric(doc, re["rubric"])
        for entry in re.get("antiphon", []):
            cel = _sub(entry.get("celebrant", ""), fd)
            ppl = _sub(entry.get("people", ""), fd)
            add_celebrant_line(doc, "Celebrant", cel)
            add_people_line(doc, "People", ppl)
        if re.get("conclusion"):
            add_spacer(doc)
            add_body_with_amen(doc, _sub(_flow(re["conclusion"]), fd)
                                 + " " + re.get("conclusion_amen", "Amen."))

    # Dismissal blessing — pulled from the top-level blessing/dismissal
    add_spacer(doc)
    if cm.get("dismissal_rubric"):
        add_rubric(doc, cm["dismissal_rubric"])
    blessing = rite_texts["blessing"]["god_of_peace"]
    add_body_with_amen(doc, _flow(blessing["text"]) + " " + blessing.get("amen", "Amen."))


def _committal_location_word(fd: FuneralData) -> str:
    """Map service.committal_location → the BCP footnote word for the
    committal prayer's `to the {location}` slot."""
    loc = fd.service.get("committal_location") or "ground"
    return {
        "columbarium": "ground",
        "cemetery":    "ground",
        "scattering":  "deep",
        "sea":         "deep",
        "cremation":   "elements",
    }.get(loc, "resting place")


# =====================================================================
# Section: blessing close (when no committal)
# =====================================================================

def _add_blessing_close(doc: Document, fd: FuneralData, rite_texts: dict,
                          song_lookup_fn) -> None:
    """For services without a Committal section, close with the
    Blessing → Closing Hymn → Dismissal sequence."""
    blessing = rite_texts["blessing"]
    add_heading2(doc, blessing.get("label", "The Blessing"))
    if blessing.get("rubric"):
        add_rubric(doc, blessing["rubric"])
    god_of_peace = blessing["god_of_peace"]
    add_body_with_amen(doc, _flow(god_of_peace["text"]) + " "
                          + god_of_peace.get("amen", "Amen."))

    closing_title = fd.music.get("closing_hymn")
    if closing_title:
        add_spacer(doc)
        add_introductory_rubric(doc, "Remain standing.")
        add_heading2(doc, "Closing Hymn")
        song = song_lookup_fn(closing_title, "11 am")
        add_song_smart(doc, song)

    # Dismissal
    add_spacer(doc)
    dismissal = rite_texts["dismissal"]["standard"]
    add_heading2(doc, "Dismissal")
    if rite_texts["dismissal"].get("rubric"):
        add_rubric(doc, rite_texts["dismissal"]["rubric"])
    for entry in dismissal.get("exchange", []):
        add_celebrant_line(doc, "Deacon", entry["celebrant"])
        add_people_line(doc, "People", entry["people"])


# =====================================================================
# Postlude + small notices + participants page
# =====================================================================

def _add_postlude(doc: Document) -> None:
    add_spacer(doc)
    add_heading2(doc, "Postlude")


def _add_interment_notice(doc: Document, text: str) -> None:
    """Italic, centered notice for No-Committal services telling
    attendees when/where the actual interment will happen."""
    add_spacer(doc)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph(style="Body")
        p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.italic = True


def _add_reception_notice(doc: Document, text: str) -> None:
    """Italic, centered notice on its own page near the back
    inviting attendees to the reception."""
    if not text:
        return
    add_spacer(doc)
    add_spacer(doc)
    for line in text.strip().split("\n"):
        p = doc.add_paragraph(style="Body")
        p.alignment = 1  # CENTER
        run = p.add_run(line)
        run.italic = True


def _add_participants_page(doc: Document, participants: dict) -> None:
    """Participants page near the back of the bulletin."""
    if not participants:
        return
    add_spacer(doc)
    add_spacer(doc)
    add_heading(doc, "Participants")

    # Roles in display order. Each tuple: (yaml-key, display-label,
    # is_list).
    rows = [
        ("celebrant",            "Celebrant",            False),
        ("assisting",            "Assisting Priests",    True),
        ("deacon",               "Deacon",               False),
        ("readers",              "Readers",              True),
        ("crucifer",             "Crucifer",             False),
        ("chalice_bearers",      "Chalice Bearers",      True),
        ("musician",             "Musician",             False),
        ("vocalist",             "Vocalist",             False),
        ("pall_bearers",         "Pall Bearers",         True),
    ]
    for key, label, is_list in rows:
        val = participants.get(key)
        if not val:
            continue
        if is_list:
            names = [n for n in val if n]
            if not names:
                continue
            for n in names:
                p = doc.add_paragraph(style="Body")
                p.alignment = 1  # CENTER
                run = p.add_run(_format_clergy_title(n) if key in
                                 ("celebrant", "assisting", "deacon") else n)
                run.bold = True
        else:
            p = doc.add_paragraph(style="Body")
            p.alignment = 1
            run = p.add_run(_format_clergy_title(val) if key in
                             ("celebrant", "assisting", "deacon") else val)
            run.bold = True
        p2 = doc.add_paragraph(style="Body")
        p2.alignment = 1
        run = p2.add_run(label)
        run.italic = True
        add_spacer(doc)

    # BCP-authority boilerplate at the bottom of the participants page.
    # Mirrors the wording on every existing St. Andrew's bulletin.
    add_spacer(doc)
    add_spacer(doc)
    p = doc.add_paragraph(style="Body")
    p.alignment = 1
    run = p.add_run(
        "The text of the service of worship contained in this program "
        "is taken from the Book of Common Prayer 1979. The Book of "
        "Common Prayer alone is of authority in the worship of the "
        "Episcopal Church. This program is provided for convenience of "
        "use on this occasion. CCLI License No. 2626819"
    )
    run.italic = True
    run.font.size = run.font.size  # leave Body default; comment-only
    # Note: the parish logo block belongs here too — to be added once a
    # back_cover_funeral.docx template exists.
