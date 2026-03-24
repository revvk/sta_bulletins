"""
Palm Sunday section module — The Sunday of the Passion.

Composes the Palm Sunday liturgy which has a unique opening structure:

Flow:
  1. Liturgy of the Palms (gathering, Palm Gospel, blessing, procession)
  2. The Word of God (Collect, Epistle, Psalm, Passion Gospel in parts, Sermon)
  3. Nicene Creed
  4. Prayers of the People + Confession + Peace
  5. Holy Communion (standard)

The Passion Gospel is read by multiple voices, each with its own
paragraph style ("Passion Gospel - Narrator", "Passion Gospel - Jesus", etc.)
for later reading-sheet generation.
"""

from docx import Document
from docx.shared import Pt

from bulletin.config import FONT_BODY_BOLD
from bulletin.data.loader import load_common_prayers, load_great_litany
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line, add_scripture_text,
)
from bulletin.document.sections.word_of_god import (
    add_reading, add_psalm, add_song_smart,
    add_body_with_amen, add_nicene_creed, add_pop,
    add_confession, add_celebrant_with_cross,
)
from bulletin.logic.rules import SeasonalRules


def add_liturgy_of_the_palms(doc: Document, rules: SeasonalRules,
                              wog_data: dict, ps_data: dict):
    """Add the Liturgy of the Palms before the Word of God.

    Args:
        doc: The Document.
        rules: SeasonalRules.
        wog_data: Standard Word of God data dict.
        ps_data: Palm Sunday specific data:
            - palm_texts: dict from palm_sunday.yaml
            - passion_gospel_lines: list of {part, text} dicts
            - passion_gospel_ref: str (e.g., "Matthew 27:11-54")
            - palm_gospel_text: ScriptureReading or str (the triumphal entry)
            - palm_gospel_ref: str (e.g., "Matthew 21:1-11")
    """
    palm_texts = ps_data["palm_texts"]
    liturgy = palm_texts["liturgy_of_the_palms"]

    # === Liturgy of the Palms ===
    add_heading(doc, "The Liturgy of the Palms")
    add_spacer(doc)

    # Gathering rubric
    gathering = liturgy.get("gathering_rubric", "")
    if gathering:
        add_introductory_rubric(doc, gathering)

    # Opening acclamation
    acclamation = liturgy["acclamation"]
    add_celebrant_line(doc, "Celebrant", acclamation["celebrant"])
    add_people_line(doc, "People", acclamation["people"])

    # Collect
    add_spacer(doc)
    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, liturgy["collect"])

    # Palm Gospel reading
    add_spacer(doc)
    palm_gospel_ref = ps_data["palm_gospel_ref"]
    add_heading2(doc, f"The Scriptures: {palm_gospel_ref}")

    palm_gospel_text = ps_data.get("palm_gospel_text")
    if palm_gospel_text:
        from bulletin.document.sections.word_of_god import _add_reading_text
        _add_reading_text(doc, palm_gospel_text)

    # Blessing of the Palms
    add_spacer(doc)
    add_rubric(doc, "The Celebrant then says the following blessing")
    blessing_dialogue = liturgy["blessing_dialogue"]
    for line in blessing_dialogue:
        if "celebrant" in line:
            add_celebrant_line(doc, "Celebrant", line["celebrant"])
        elif "people" in line:
            add_people_line(doc, "People", line["people"])

    add_spacer(doc)
    add_body_with_amen(doc, liturgy["blessing_prayer"])
    add_spacer(doc)

    # Blessing response
    response = liturgy["blessing_response"]
    add_celebrant_with_cross(doc, "Celebrant", response["celebrant"])
    add_people_line(doc, "People", response["people"])

    # Procession
    add_spacer(doc)
    add_heading2(doc, "The Procession")
    procession = liturgy["procession"]
    add_rubric(doc, "The Deacon or a Priest says")
    add_celebrant_line(doc, "", procession["deacon"])
    add_people_line(doc, "People", procession["people"])

    add_spacer(doc)
    add_introductory_rubric(doc, procession["rubric"])

    # Processional hymn(s)
    processional_songs = wog_data.get("processional_songs", [])
    if processional_songs:
        for i, song in enumerate(processional_songs):
            add_spacer(doc)
            add_song_smart(doc, song)
    else:
        add_spacer(doc)
        add_song_smart(doc, wog_data.get("processional"))


def add_palm_sunday_word_of_god(doc: Document, rules: SeasonalRules,
                                 wog_data: dict, ps_data: dict):
    """Add the Word of God section for Palm Sunday.

    This replaces the standard Word of God. It includes:
    - Collect of the Day
    - First Reading (Epistle)
    - Psalm
    - Passion Gospel (read in parts)
    - Sermon
    - Nicene Creed
    - Prayers of the People + Confession + Peace
    """
    prayers = load_common_prayers()
    service_time = wog_data.get("service_time", "9 am")

    # === The Word of God ===
    add_heading(doc, "The Word of God")

    # Collect of the Day
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Collect of the Day")
    add_celebrant_line(doc, "Celebrant", "The Lord be with you.")
    add_people_line(doc, "People", "And also with you.")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, wog_data["collect_text"])

    # Be seated for readings
    add_spacer(doc)
    if service_time == "11 am":
        add_introductory_rubric(
            doc,
            "Elementary-aged children are invited to attend Children\u2019s Church. "
            "They will rejoin the service during The Peace."
        )
        add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # First Reading (Epistle)
    add_reading(doc, wog_data["reading_1_ref"], wog_data["reading_1_text"])

    # Psalm
    add_spacer(doc)
    add_psalm(doc, wog_data["psalm_ref"], wog_data.get("psalm_rubric", ""),
              wog_data["psalm_text"])

    # Sequence Hymn
    add_spacer(doc)
    if service_time != "8 am":
        add_introductory_rubric(doc, "Please stand.")
        add_heading2(doc, "Sequence Hymn")
        add_song_smart(doc, wog_data.get("sequence_hymn"))

    # === Passion Gospel ===
    _add_passion_gospel(doc, ps_data, service_time)

    # Musical Reflection (replaces Sermon on Palm Sunday)
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "Musical Reflection")
    add_rubric(doc,
               "In response to our Lord\u2019s passion and death, we will "
               "observe a time of quiet prayer and reflection.")

    # Nicene Creed
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Nicene Creed")
    add_nicene_creed(doc, prayers)

    # The Great Litany (replaces Prayers of the People on Palm Sunday)
    add_spacer(doc)
    add_heading2(doc, "The Great Litany")
    great_litany = load_great_litany()
    add_pop(doc, great_litany.get("elements", []))

    # Confession & Absolution (if not done in Penitential Order)
    if not rules.no_confession_after_pop:
        add_spacer(doc)
        add_confession(doc, prayers)

    # The Peace
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Peace")
    add_celebrant_line(doc, "Celebrant", "The peace of the Lord be always with you.")
    add_people_line(doc, "People", "And also with you.")


def _add_passion_gospel(doc: Document, ps_data: dict, service_time: str):
    """Add the Passion Gospel reading in parts.

    Each speaking part is rendered with its own named style
    ("Passion Gospel - Narrator", etc.) using a hanging indent
    so the part name sits on the left and text wraps cleanly.
    """
    palm_texts = ps_data["palm_texts"]
    passion_config = palm_texts["passion_gospel"]
    passion_lines = ps_data["passion_gospel_lines"]
    passion_ref = ps_data["passion_gospel_ref"]
    gospel_book = passion_ref.split()[0]

    # Gospel heading
    add_spacer(doc)
    if service_time == "8 am":
        add_introductory_rubric(doc, "Please stand.")
    else:
        add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, f"The Gospel: {passion_ref}")

    # Rubric about seating/standing
    seating = passion_config.get("seating_rubric", "")
    if seating:
        add_rubric(doc, seating)

    # Announcement rubric
    omission = passion_config.get("omission_rubric", "")
    if omission:
        add_rubric(doc, omission)

    add_spacer(doc)

    _add_passion_gospel_lines(doc, passion_lines)


def _add_passion_gospel_lines(doc: Document, passion_lines: list[dict]):
    """Render passion gospel lines with part-specific styles.

    Each line is rendered as:
        PART_NAME:  [tab]  Spoken text here...

    Using a hanging-indent style so the part name sits at the left
    margin and the spoken text wraps cleanly at the tab stop.

    This function is shared between Palm Sunday and Good Friday.

    Args:
        doc: The Document.
        passion_lines: List of dicts with 'part' and 'text' keys.
    """
    for line in passion_lines:
        part = line["part"]
        text = line["text"]

        if part == "Rubric":
            # Inline rubric (e.g., "All stand.")
            add_spacer(doc)
            add_introductory_rubric(doc, text)
            add_spacer(doc)
            continue

        # Use the part-specific style
        style_name = f"Passion Gospel - {part}"

        # Fall back to Narrator style if the part doesn't have its own style
        try:
            doc.styles[style_name]
        except KeyError:
            style_name = "Passion Gospel - Narrator"

        p = doc.add_paragraph(style=style_name)
        # Part name (lowercase + small caps renders as capitals, 9pt) + tab + spoken text
        label_run = p.add_run(f"{part.lower()}:\t")
        label_run.font.small_caps = True
        label_run.font.size = Pt(9)
        p.add_run(text)
