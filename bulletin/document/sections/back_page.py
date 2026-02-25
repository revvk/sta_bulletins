"""
Back page of the bulletin.

Contains:
  - Welcome blurb
  - Altar Guild acknowledgment
  - BCP attribution
  - Staff & Vestry listing
  - Optional flower memorial
"""

from docx import Document

from bulletin.data.loader import load_staff
from bulletin.document.formatting import add_spacer, add_heading


def add_back_page(doc: Document, data: dict = None):
    """Add the back page content.

    Args:
        doc: The Document.
        data: Optional dict with:
            - flower_memorial: str | None (e.g., "the flowers on the altar...")
    """
    staff_data = load_staff()
    data = data or {}

    # Welcome section
    add_heading(doc, "Welcome")

    p = doc.add_paragraph(style="Cover Note")
    p.add_run(
        "Have questions about our worship or St. Andrew's? Feel free to "
        "contact Father Andrew or Father Logan at the contact information below."
    )

    add_spacer(doc)

    p = doc.add_paragraph(style="Cover Note")
    p.add_run(
        "Our worship is made possible each week by the faithful service of "
        "the members of our Altar Guild, directed by Chris Pruitt."
    )

    add_spacer(doc)

    p = doc.add_paragraph(style="Cover Note")
    run = p.add_run(
        "The text of the service of worship contained in this program is taken "
        "from the Book of Common Prayer 1979. The Book of Common Prayer is the "
        "official liturgical book of The Episcopal Church."
    )
    run.italic = True

    # Staff & Vestry heading
    add_heading(doc, "St. Andrew's Staff & Vestry")

    # Clergy
    for person in staff_data.get("clergy", []):
        doc.add_paragraph(person["name"], style="Staff - Name")
        doc.add_paragraph(person["title"], style="Staff - Title")
        if person.get("email"):
            doc.add_paragraph(person["email"], style="Staff - Email")

    # Staff
    for person in staff_data.get("staff", []):
        doc.add_paragraph(person["name"], style="Staff - Name")
        doc.add_paragraph(person["title"], style="Staff - Title")
        if person.get("email"):
            doc.add_paragraph(person["email"], style="Staff - Email")

    # Vestry
    vestry = staff_data.get("vestry", {})
    if vestry:
        doc.add_paragraph(vestry.get("senior_warden", ""), style="Staff - Name")
        doc.add_paragraph("Senior Warden", style="Staff - Title")

        doc.add_paragraph(vestry.get("junior_warden", ""), style="Staff - Name")
        doc.add_paragraph("Junior Warden", style="Staff - Title")

        members = vestry.get("members", [])
        if members:
            doc.add_paragraph("Vestry", style="Staff - Name")
            # Format as comma-separated list
            doc.add_paragraph(", ".join(members), style="Staff - Email")

    # Optional flower memorial
    flower = data.get("flower_memorial")
    if flower:
        add_spacer(doc)
        for line in flower if isinstance(flower, list) else [flower]:
            p = doc.add_paragraph(style="Spacer - Small")
            run = p.add_run(line)
            run.bold = True
