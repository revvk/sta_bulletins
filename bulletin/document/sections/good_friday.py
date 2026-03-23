"""
Good Friday section module.

Composes the Good Friday liturgy from shared building blocks.
This is NOT a Eucharist — there is no communion.

Flow:
  1. Silent entrance
  2. Opening Acclamation + Collect of the Day
  3. First Reading (Isaiah) + Psalm + Sequence Hymn
  4. The Passion Gospel (John) — dramatic multi-voice reading
  5. Sermon
  6. The Solemn Collects (5 bidding prayers)
  7. Veneration of the Cross (hymns + antiphon)
  8. The Lord's Prayer (standalone, no Eucharistic context)
  9. Final Prayer + silent departure
"""

from docx import Document

from bulletin.config import CROSS_SYMBOL, FONT_BODY_BOLD
from bulletin.data.loader import load_common_prayers
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line,
)
from bulletin.document.sections.word_of_god import (
    add_reading, add_psalm, add_gospel, add_song_smart,
    add_body_with_amen,
)
from bulletin.logic.rules import SeasonalRules


def add_good_friday(doc: Document, rules: SeasonalRules,
                     wog_data: dict, gf_data: dict):
    """Add the complete Good Friday liturgy.

    Args:
        doc: The Document.
        rules: SeasonalRules (with Good Friday acclamation overrides).
        wog_data: Standard Word of God data dict (readings, psalm, preacher).
        gf_data: Good Friday specific data:
            - entrance_rubric: str
            - acclamation: dict with celebrant, people
            - solemn_collects: dict with introduction, exhortation,
                              rubric, collects (list)
            - veneration_hymns: list of song data dicts
            - veneration: dict with antiphon, antiphon_psalm
            - closing_prayer: str
            - closing_rubric: str
    """
    # ==========================================
    # Part 1: Entrance + Acclamation + Collect
    # ==========================================
    _add_entrance_and_collect(doc, rules, wog_data, gf_data)

    # ==========================================
    # Part 2: Readings + Sermon
    # ==========================================
    _add_readings_and_sermon(doc, wog_data, gf_data)

    # ==========================================
    # Part 3: The Solemn Collects
    # ==========================================
    _add_solemn_collects(doc, gf_data)

    # ==========================================
    # Part 4: Veneration of the Cross
    # ==========================================
    _add_veneration(doc, gf_data)

    # ==========================================
    # Part 5: Lord's Prayer + Final Prayer
    # ==========================================
    _add_lords_prayer_and_closing(doc, gf_data)


def _add_entrance_and_collect(doc: Document, rules: SeasonalRules,
                                wog_data: dict, gf_data: dict):
    """Silent entrance, acclamation, and Collect of the Day."""
    # No "The Word of God" heading on Good Friday
    # No prelude, no processional

    # Silent entrance rubric
    entrance_rubric = gf_data.get("entrance_rubric", "")
    if entrance_rubric:
        add_introductory_rubric(doc, entrance_rubric)
    add_rubric(doc, "All then kneel for silent prayer, after which the "
               "Celebrant alone stands and begins the liturgy.")

    # Collect of the Day (no Sursum Corda dialogue on Good Friday)
    add_spacer(doc)
    add_heading2(doc, "The Collect of the Day")

    # Opening Acclamation (embedded in the Collect section)
    acclamation = gf_data.get("acclamation", {})
    add_celebrant_line(doc, "Celebrant",
                       acclamation.get("celebrant", rules.acclamation_celebrant))
    add_people_line(doc, "People",
                    acclamation.get("people", rules.acclamation_people))
    add_spacer(doc)

    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, wog_data["collect_text"])


def _add_readings_and_sermon(doc: Document, data: dict, gf_data: dict = None):
    """Readings, Psalm, Sequence Hymn, Passion Gospel, Sermon."""
    # Be seated for readings
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # First Reading (Isaiah on Good Friday)
    add_reading(doc, data["reading_1_ref"], data["reading_1_text"])

    # Psalm
    add_spacer(doc)
    add_psalm(doc, data["psalm_ref"], data.get("psalm_rubric", ""),
               data["psalm_text"])

    # Sequence Hymn
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Sequence Hymn")
    add_song_smart(doc, data.get("sequence_hymn"))

    # Be seated for Passion Gospel
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # The Passion Gospel — dramatic multi-voice reading
    add_heading2(doc, f"The Gospel: {data['gospel_ref']}")
    add_rubric(doc, "The congregation may be seated for the first part of "
               "the Passion. At the verse which mentions the arrival at "
               "Golgotha all stand.")
    add_rubric(doc, "The customary responses before and after the Gospel "
               "are omitted.")
    add_spacer(doc)

    # Render passion gospel in parts if available
    from bulletin.document.sections.palm_sunday import _add_passion_gospel_lines
    passion_lines = gf_data.get("passion_gospel_lines", []) if gf_data else []
    if passion_lines:
        _add_passion_gospel_lines(doc, passion_lines)
    else:
        # Fallback: render as plain text
        gospel = data.get("gospel_text")
        if gospel and hasattr(gospel, "paragraphs"):
            from bulletin.document.formatting import add_scripture_text
            for i, para in enumerate(gospel.paragraphs):
                add_scripture_text(doc, para, indent=(i > 0))
        elif gospel:
            add_body(doc, "[Passion Gospel text — to be formatted with parts]")

    # Sermon
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "The Sermon")
    add_body(doc, data.get("preacher", ""))


def _add_solemn_collects(doc: Document, gf_data: dict):
    """Add the Solemn Collects — 5 bidding prayers with silence + collect."""
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Solemn Collects")

    solemn = gf_data.get("solemn_collects", {})

    # Introductory address
    introduction = solemn.get("introduction", "")
    if introduction:
        add_rubric(doc, "The Deacon, or other person appointed, "
                   "says to the people")
        add_body(doc, introduction)
        add_spacer(doc)

    exhortation = solemn.get("exhortation", "")
    if exhortation:
        add_body(doc, exhortation)
        add_spacer(doc)

    # Main rubric
    rubric = solemn.get("rubric", "")
    if rubric:
        add_rubric(doc, rubric)
        add_spacer(doc)

    # Each bidding + silence + collect
    collects = solemn.get("collects", [])
    for i, collect in enumerate(collects):
        # Bidding
        bidding = collect.get("bidding", "")
        if bidding:
            add_body(doc, bidding)

        # Silence
        add_spacer(doc)
        add_rubric(doc, "Silence")
        add_spacer(doc)

        # Celebrant's collect
        collect_text = collect.get("collect", "")
        if collect_text:
            add_body_with_amen(doc, collect_text)

        # Spacer between collects (but not after the last)
        if i + 1 < len(collects):
            add_spacer(doc)


def _add_veneration(doc: Document, gf_data: dict):
    """Add the Veneration of the Cross with hymns and antiphon."""
    add_heading(doc, "Veneration of the Cross")

    # Rubric about the cross being brought in
    add_rubric(doc, "During the hymn, a wooden cross is brought into the "
               "church and placed in the sight of the people.")
    add_spacer(doc)

    # First hymn (during the cross procession)
    veneration_hymns = gf_data.get("veneration_hymns", [])
    if veneration_hymns:
        add_song_smart(doc, veneration_hymns[0])
        add_spacer(doc)

    # Antiphon (after the cross is placed)
    veneration = gf_data.get("veneration", {})
    antiphon = veneration.get("antiphon", "")
    antiphon_psalm = veneration.get("antiphon_psalm", "")

    add_rubric(doc, "After the cross is nailed together, the anthem below "
               "is recited, the congregation reading the part in bold.")
    add_spacer(doc)

    if antiphon:
        # Antiphon (versicle — normal text)
        add_body(doc, antiphon)

    if antiphon_psalm:
        # Psalm verse (normal text)
        add_body(doc, antiphon_psalm)

    if antiphon:
        # Antiphon repeated (bold — people's part)
        p = doc.add_paragraph(style="Body")
        run = p.add_run(antiphon)
        run.bold = True
        run.font.name = FONT_BODY_BOLD

    # Additional hymn(s) during veneration
    if len(veneration_hymns) > 1:
        add_spacer(doc)
        add_rubric(doc, "Led by the cantor, the people now sing together "
                   "the following hymn. During the singing of the hymn, "
                   "the congregation is invited to come forward to the "
                   "Cross for private devotion and veneration.")
        for hymn in veneration_hymns[1:]:
            add_spacer(doc)
            add_song_smart(doc, hymn)


def _add_lords_prayer_and_closing(doc: Document, gf_data: dict):
    """Standalone Lord's Prayer and Final Prayer — no Eucharist."""
    prayers = load_common_prayers()

    # Lord's Prayer (standalone — no eucharistic context)
    add_spacer(doc)
    add_heading2(doc, "The Lord's Prayer")
    text = " ".join(line.strip() for line in prayers["lords_prayer"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]

    # Final Prayer
    add_spacer(doc)
    add_heading2(doc, "Final Prayer")
    add_rubric(doc, "The Celebrant concludes the service with the following "
               "prayer. No blessing or dismissal is added.")
    closing_prayer = gf_data.get("closing_prayer", "")
    if closing_prayer:
        add_body_with_amen(doc, closing_prayer)

    # Silent departure
    closing_rubric = gf_data.get("closing_rubric", "")
    if closing_rubric:
        add_spacer(doc)
        add_rubric(doc, closing_rubric)
