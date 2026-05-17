"""
The Holy Communion section of the bulletin.

Covers everything from the Offertory through the Dismissal:
  - Offertory
  - Offering Music
  - Doxology
  - The Great Thanksgiving (Sursum Corda + Proper Preface + Sanctus + Eucharistic Prayer)
  - Lord's Prayer
  - Breaking of the Bread
  - Invitation
  - Holy Communion (rubric)
  - Communion Music
  - Closing Prayer
  - Prayer for Lay Eucharistic Visitor (optional)
  - Blessing / Prayer over the People
  - Closing Hymn
  - Dismissal
  - Postlude
"""

from docx import Document

from bulletin.config import (
    CROSS_SYMBOL, FONT_BODY, FONT_BODY_BOLD, FONT_HEADER_FOOTER,
    GIVING_URL, CONNECT_URL,
)
from bulletin.data.loader import (
    load_common_prayers, load_eucharistic_prayers, load_blessings,
    get_proper_preface_text,
)
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_body_with_bold_ending,
    add_celebrant_line, add_people_line, add_cross_symbol,
    add_song, add_song_two_column, add_hymn_header,
    add_no_split_block,
)
from bulletin.logic.rules import SeasonalRules


def add_holy_communion(doc: Document, rules: SeasonalRules, data: dict):
    """Add the entire Holy Communion section.

    Args:
        doc: The Document.
        rules: SeasonalRules for this service.
        data: Dict with keys:
            - offertory_song: song data dict (or None)
            - sanctus_song: song data dict (or None) — the Sanctus setting
            - communion_songs: list of song data dicts
            - closing_hymn: song data dict (or None)
            - postlude_songs: list of song data dicts (medley for postlude)
            - eucharistic_prayer: "A", "B", or "C"
            - proper_preface_text: str (already resolved)
            - fraction_song: song data dict | None (Lent: Agnus Dei)
            - blessing_text: str | None
            - dismissal_deacon: str
            - dismissal_people: str
            - post_communion_prayer: "a" or "b"
            - include_lev: bool (Lay Eucharistic Visitor prayer)
    """
    prayers = load_common_prayers()
    ep_data = load_eucharistic_prayers()

    # Section heading
    add_heading(doc, "The Holy Communion")

    # --- Offertory ---
    add_spacer(doc)
    be_seated_p = add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "Offertory")
    service_time = data.get("service_time", "9 am")
    if service_time in ("8 am", "11 am"):
        add_short_offertory_rubric(doc)
        # Anchor the same QR code that 9am uses, but to the "Be seated."
        # rubric just above the Offertory heading.
        _add_qr_code(doc, be_seated_p)
    else:
        add_offertory_rubric(doc)
    add_spacer(doc)

    # Offering Music (skip at 8am — no music)
    if service_time != "8 am":
        add_heading2(doc, "Offering Music")
        if service_time == "11 am":
            _add_11am_offertory(doc, data)
        else:
            add_communion_song_smart(doc, data.get("offertory_song"))

    # Doxology (skip at 8am — no music)
    if service_time != "8 am":
        add_spacer(doc)
        add_introductory_rubric(doc, "Please stand.")
        add_heading2(doc, "Doxology")
        _DOXOLOGY = [
            "Praise God, from Whom all blessings flow;",
            "Praise Him, all creatures here below;",
            "Praise Him above, ye heavenly host:",
            "Praise Father, Son, and Holy Ghost.",
        ]
        def _add_doxology(cell):
            for line in _DOXOLOGY:
                cell.add_paragraph(line, style="Body - Lyrics")
        add_no_split_block(doc, _add_doxology)

    # --- The Great Thanksgiving ---
    add_spacer(doc)
    if service_time == "8 am":
        # Congregation is still seated (no Doxology), so "Please stand"
        add_introductory_rubric(doc, "Please stand.")
    else:
        add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "The Great Thanksgiving")

    # Sursum Corda
    for exchange in ep_data["sursum_corda"]:
        add_celebrant_line(doc, "Celebrant", exchange["celebrant"])
        add_people_line(doc, "People", exchange["people"])
    add_spacer(doc)

    # Eucharistic Prayer
    prayer_key = data.get("eucharistic_prayer", "A").upper()
    if prayer_key == "C":
        add_prayer_c(doc, ep_data, data, prayers)
    else:
        add_prayer_a_or_b(doc, ep_data, data, prayers, prayer_key)

    # --- Lord's Prayer ---
    add_spacer(doc)
    add_rubric(doc, "Please stand and, as you are comfortable, join hands "
               "with those around you.")
    add_body(doc, prayers["lords_prayer_intro"]["option_1"])
    text = " ".join(line.strip() for line in prayers["lords_prayer"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]

    # --- Breaking of the Bread ---
    add_spacer(doc)
    add_heading2(doc, "Breaking of the Bread")
    if rules.use_fraction_anthem:
        # Lent: fraction anthem (Agnus Dei)
        if service_time == "11 am":
            # 11am: Cantor/People call-and-response with music notation images
            _add_agnus_dei_images(doc)
        elif service_time == "8 am":
            # 8am: spoken Agnus Dei with split bold
            add_agnus_dei_spoken(doc)
        else:
            fraction_song = data.get("fraction_song")
            if fraction_song:
                add_communion_song_smart(doc, fraction_song,
                                force_single_column=True)
            else:
                # Default Agnus Dei text
                for line in [
                    "Lamb of God, you take away the sins of the world: have mercy on us.",
                    "Lamb of God, you take away the sins of the world: have mercy on us.",
                    "Lamb of God, you take away the sins of the world: grant us peace.",
                ]:
                    p = doc.add_paragraph(style="Body - Dialogue")
                    run = p.add_run(line)
                    run.bold = True
                    run.font.name = FONT_BODY_BOLD
    else:
        add_celebrant_line(doc, "Celebrant", rules.fraction_celebrant)
        add_people_line(doc, "People", rules.fraction_people)

    # --- Invitation ---
    add_spacer(doc)
    add_heading2(doc, "Invitation")
    add_rubric(doc, "Facing the people, the Celebrant says the following Invitation")
    add_body(doc, prayers["invitation_to_communion"] + " " +
             prayers["invitation_addition"])

    # --- Holy Communion ---
    add_spacer(doc)
    add_heading2(doc, "Holy Communion")
    add_rubric(doc, "All baptized Christians are invited to the altar rail to "
               "receive Holy Communion. If you are not going to receive, you "
               "are still welcome at the rail for a blessing; just cross your "
               "arms over your chest.")

    # --- Communion Music (skip at 8am — no music) ---
    if service_time != "8 am":
        add_spacer(doc)
        add_heading2(doc, "Communion Music")
        comm_songs = data.get("communion_songs", [])
        for i, song in enumerate(comm_songs):
            if i > 0:
                add_spacer(doc)
            add_communion_song_smart(doc, song)

    # --- Closing Prayer ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Closing Prayer")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")

    prayer_choice = data.get("post_communion_prayer", "a")
    if prayer_choice == "b":
        text = " ".join(prayers["post_communion_prayer_b"])
    else:
        text = " ".join(prayers["post_communion_prayer_a"])
    p = doc.add_paragraph(style="Body - People Recitation")
    run = p.add_run(text)
    run.style = doc.styles["People"]

    # --- Prayer for Lay Eucharistic Visitor (optional) ---
    if data.get("include_lev", True):
        add_spacer(doc)
        add_heading2(doc, "Prayer for Lay Eucharistic Visitor")
        p = doc.add_paragraph(style="Body - Rubric")
        run = p.add_run("The Celebrant commissions the Lay Eucharistic Visitor, saying")
        run.italic = True
        add_celebrant_line(
            doc, "",
            "In the name of this congregation, I send you forth bearing these "
            "holy gifts that those to whom you go may share with us in the "
            "communion of Christ's body and blood.",
        )
        add_people_line(doc, "People",
                        "We who are many are one body because we all share "
                        "one bread and one cup.")

    # --- Blessing / Prayer over the People ---
    add_spacer(doc)
    add_heading2(doc, rules.blessing_label)
    if rules.use_prayer_over_people:
        add_rubric(doc, "The Deacon or a Priest says")
        add_celebrant_line(doc, "", "Bow down before the Lord.")
        add_spacer(doc)
        add_rubric(doc, "The people bow their heads and the Celebrant says "
                   "the following prayer")
        blessing_text = data.get("blessing_text", "")
        if blessing_text:
            add_body_with_amen(doc, blessing_text)
    else:
        blessing_text = data.get("blessing_text", "")
        if not blessing_text:
            blessing_text = (
                "\u2026 and the blessing of God Almighty, "
                f"{CROSS_SYMBOL} the Father, the Son, and the Holy Spirit, "
                "be among you, and remain with you always."
            )
        else:
            # Insert cross before trinitarian formula
            blessing_text = blessing_text.replace(
                "the Father, the Son, and the Holy Spirit",
                f"{CROSS_SYMBOL} the Father, the Son, and the Holy Spirit",
            )
        add_blessing_line(doc, blessing_text)

    # --- Closing Hymn (skip at 8am — no music) ---
    if service_time != "8 am":
        add_spacer(doc)
        add_heading2(doc, "Closing Hymn")
        add_communion_song_smart(doc, data.get("closing_hymn"))

    # --- Dismissal ---
    add_spacer(doc)
    add_heading2(doc, "Dismissal")
    add_rubric(doc, "The Deacon or a Priest dismisses the people with these words")
    add_celebrant_line(doc, "", data.get("dismissal_deacon", ""))
    add_people_line(doc, "People", data.get("dismissal_people", ""))

    # --- Postlude (skip at 8am — no music) ---
    if service_time != "8 am":
        add_spacer(doc)
        add_heading2(doc, "Postlude")
        postlude_songs = data.get("postlude_songs", [])
        for song in postlude_songs:
            add_communion_song_smart(doc, song)


def add_prayer_a_or_b(doc: Document, ep_data: dict, data: dict,
                       prayers: dict, prayer_key: str):
    """Add Eucharistic Prayer A or B."""
    key = f"prayer_{prayer_key.lower()}"
    prayer = ep_data[key]

    # Preface opening + proper preface + Sanctus transition
    # The proper preface is joined to the opening as a single paragraph.
    # Most prefaces start with connecting words ("Through", "For", "Because")
    # and join with a comma: "...Creator of heaven and earth, through..."
    # Some prefaces (e.g. Lent's "Paschal preparation") start a new sentence:
    # "...Creator of heaven and earth. You bid..."
    preface_opening = ep_data["preface_opening"]
    preface_text = data.get("proper_preface_text", "")
    if preface_text:
        # Connecting words that join naturally with a comma + lowercase
        _CONNECTING_WORDS = {"through", "for", "because", "but", "who"}
        first_word = preface_text.split()[0].lower()
        if first_word in _CONNECTING_WORDS:
            # Strip trailing period, join with comma, lowercase first char
            opening = preface_opening.rstrip().rstrip(".")
            joined = preface_text[0].lower() + preface_text[1:]
            add_body(doc, f"{opening}, {joined}")
        else:
            # New sentence: keep the period, keep original capitalization
            add_body(doc, f"{preface_opening.rstrip()} {preface_text}")
    else:
        add_body(doc, preface_opening)
    add_spacer(doc)

    add_body(doc, ep_data["sanctus_transition"])
    add_spacer(doc)

    # Sanctus hymn
    sanctus_song = data.get("sanctus_song")
    if sanctus_song:
        add_communion_song_smart(doc, sanctus_song, force_single_column=True)
    elif data.get("service_time") == "8 am":
        add_sanctus_spoken(doc, prayers["sanctus"])
    else:
        # Text version of Sanctus
        add_sanctus_text(doc, prayers["sanctus"])

    # Kneeling rubric
    add_spacer(doc)
    add_rubric(doc, "Please kneel or remain standing. The Celebrant continues")

    # Post-Sanctus
    add_body(doc, prayer["post_sanctus"])
    add_spacer(doc)

    # Institution narrative (Prayer A has institution_1, B does not)
    if "institution_1" in prayer:
        add_body(doc, prayer["institution_1"])
        add_spacer(doc)

    # Words of institution - bread
    _add_institution_words(doc, prayer["institution_bread"])
    add_spacer(doc)

    # Words of institution - cup
    _add_institution_words(doc, prayer["institution_cup"])
    add_spacer(doc)

    # Memorial acclamation
    add_body(doc, prayer["memorial_intro"])
    p = doc.add_paragraph(style="Body")
    for i, line in enumerate(prayer["memorial_acclamation"]):
        if i > 0:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line)
        run.bold = True
        run.font.name = FONT_BODY_BOLD
    add_spacer(doc)

    # Epiclesis
    add_rubric(doc, "The Celebrant continues")
    add_body(doc, prayer["epiclesis_1"] if "epiclesis_1" in prayer else prayer.get("epiclesis", ""))
    add_spacer(doc)
    add_body(doc, prayer["epiclesis_2"])
    add_spacer(doc)

    # Doxology
    add_doxology_amen(doc, prayer)


def add_prayer_c(doc: Document, ep_data: dict, data: dict, prayers: dict):
    """Add Eucharistic Prayer C (responsive format).

    Prayer C has a unique structure different from A/B — it's fully
    responsive with Celebrant/People exchanges throughout.
    """
    # Prayer C text — we output the full responsive text
    # The preface is woven into the prayer itself, not separate
    pc = [
        ("celebrant", "God of all power, Ruler of the Universe, you are worthy of glory and praise."),
        ("people", "Glory to you for ever and ever."),
        ("celebrant", "At your command all things came to be: the vast expanse of interstellar space, galaxies, suns, the planets in their courses, and this fragile earth, our island home."),
        ("people", "By your will they were created and have their being."),
        ("celebrant", "From the primal elements you brought forth the human race, and blessed us with memory, reason, and skill. You made us the rulers of creation. But we turned against you, and betrayed your trust; and we turned against one another."),
        ("people", "Have mercy, Lord, for we are sinners in your sight."),
        ("celebrant", "Again and again, you called us to return. Through prophets and sages you revealed your righteous Law. And in the fullness of time you sent your only Son, born of a woman, to fulfill your Law, to open for us the way of freedom and peace."),
        ("people", "By his blood, he reconciled us. By his wounds, we are healed."),
        ("celebrant", "And therefore we praise you, joining with the heavenly chorus, with prophets, apostles, and martyrs, and with all those in every generation who have looked to you in hope, to proclaim with them your glory, in their unending hymn:"),
    ]

    for role, text in pc:
        if role == "celebrant":
            add_body(doc, text)
        else:
            p = doc.add_paragraph(style="Body")
            run = p.add_run(text)
            run.bold = True
            run.font.name = FONT_BODY_BOLD
        add_spacer(doc)

    # Sanctus
    sanctus_song = data.get("sanctus_song")
    if sanctus_song:
        add_communion_song_smart(doc, sanctus_song, force_single_column=True)
    elif data.get("service_time") == "8 am":
        add_sanctus_spoken(doc, prayers["sanctus"])
    else:
        add_sanctus_text(doc, prayers["sanctus"])

    # Post-Sanctus continuation
    add_spacer(doc)
    add_rubric(doc, "Please kneel or remain standing. The Celebrant continues")

    pc2 = [
        ("celebrant", "And so, Father, we who have been redeemed by him, and made a new people by water and the Spirit, now bring before you these gifts. Sanctify them by your Holy Spirit to be the Body and Blood of Jesus Christ our Lord."),
    ]
    for role, text in pc2:
        add_body(doc, text)
    add_spacer(doc)

    # Institution - bread
    _add_institution_words(doc,
        "On the night he was betrayed he took bread, said the blessing, "
        "broke the bread, and gave it to his friends, and said, "
        '"Take, eat: This is my Body, which is given for you. '
        f'Do this for the remembrance of me." {CROSS_SYMBOL}')
    add_spacer(doc)

    # Institution - cup
    _add_institution_words(doc,
        "After supper, he took the cup of wine, gave thanks, and said, "
        '"Drink this, all of you: This is my Blood of the new Covenant, '
        "which is shed for you and for many for the forgiveness of sins. "
        f'Whenever you drink it, do this for the remembrance of me." {CROSS_SYMBOL}')
    add_spacer(doc)

    # Memorial acclamation
    pc3 = [
        ("celebrant", "Remembering now his work of redemption, and offering to you this sacrifice of thanksgiving,"),
        ("people", "We celebrate his death and resurrection, as we await the day of his coming."),
    ]
    for role, text in pc3:
        if role == "celebrant":
            add_body(doc, text)
        else:
            p = doc.add_paragraph(style="Body")
            run = p.add_run(text)
            run.bold = True
            run.font.name = FONT_BODY_BOLD
    add_spacer(doc)

    # Epiclesis
    pc4 = [
        ("celebrant", "Lord God of our Fathers; God of Abraham, Isaac, and Jacob; God and Father of our Lord Jesus Christ: Open our eyes to see your hand at work in the world about us. Deliver us from the presumption of coming to this Table for solace only, and not for strength; for pardon only, and not for renewal. Let the grace of this Holy Communion make us one body, one spirit in Christ ✠, that we may worthily serve the world in his name."),
        ("people", "Risen Lord, be known to us in the breaking of the Bread."),
    ]
    for role, text in pc4:
        if role == "celebrant":
            add_body(doc, text)
        else:
            p = doc.add_paragraph(style="Body")
            run = p.add_run(text)
            run.bold = True
            run.font.name = FONT_BODY_BOLD
    add_spacer(doc)

    # Doxology
    p = doc.add_paragraph(style="Body")
    p.add_run(
        "Accept these prayers and praises, Father, through Jesus Christ our "
        "great High Priest, to whom, with you and the Holy Spirit, your "
        "Church gives honor, glory, and worship, from generation to generation. "
    )
    run = p.add_run("AMEN.")
    run.bold = True
    run.font.name = FONT_BODY_BOLD


def add_sanctus_text(doc: Document, lines: list[str]):
    """Add text Sanctus in a no-split block, rendering ✠ as a bold cross."""
    def _add_sanctus_lines(cell):
        for line in lines:
            if CROSS_SYMBOL in line:
                p = cell.add_paragraph(style="Body - Lyrics")
                parts = line.split(CROSS_SYMBOL)
                for i, part in enumerate(parts):
                    if i > 0:
                        add_cross_symbol(p)
                    if part:
                        p.add_run(part)
            else:
                cell.add_paragraph(line, style="Body - Lyrics")
    add_no_split_block(doc, _add_sanctus_lines)


def _add_institution_words(doc: Document, text: str):
    """Add institution narrative as plain body text."""
    add_body(doc, text)


def add_doxology_amen(doc: Document, prayer: dict):
    """Add the prayer doxology with bold AMEN."""
    p = doc.add_paragraph(style="Body")
    p.add_run(prayer["doxology"] + " ")
    run = p.add_run(prayer["doxology_response"])
    run.bold = True
    run.font.name = FONT_BODY_BOLD


def add_communion_song_smart(doc: Document, song_data: dict | None,
                             force_single_column: bool = False):
    """Add a song inside the Holy Communion liturgy.

    Differs from ``word_of_god.add_song_smart`` in ONE place — when
    ``force_single_column=True``, this variant lays the song out as a
    single column with all sections in **one** cell (not one row per
    section). That matches what Sunday-morning Communion needs for
    short liturgical pieces (Sanctus, fraction anthem, Agnus Dei) where
    the cell-per-row treatment would add gratuitous vertical space.

    Hidden Springs has the opposite need (large-print, one verse per
    row), so HS explicitly imports the default ``add_song_smart`` from
    ``word_of_god`` — that path uses ``multi_row=True`` under the hood.
    Keeping the two functions named differently makes the divergence
    visible at every call site, instead of a silent behaviour change
    based on which module's symbol you happened to import.

    Args:
        force_single_column: If True, always use single-column layout
            (one table cell containing every section). For ANY layout
            that wants one row per section, call
            ``word_of_god.add_song_smart`` with ``force_single_column=True``
            instead.
    """
    if not song_data:
        add_body(doc, "[Song lyrics not found]")
        return

    sections = song_data.get("sections", [])

    # Hymnal-only: just the header line, no lyrics
    if not sections:
        add_hymn_header(
            doc,
            song_data["title"],
            song_data.get("tune_name"),
            song_data.get("hymnal_number"),
            song_data.get("hymnal_name"),
        )
        return

    max_line_len = max(
        (len(line) for s in sections for line in s["lines"]),
        default=0
    )

    # Two-column if 2+ sections and lines fit half-width
    if not force_single_column and len(sections) >= 2 and max_line_len <= 52:
        add_song_two_column(doc, song_data)
    else:
        add_song(doc, song_data)


# ``add_body_with_amen`` lives in word_of_god — re-export here so that
# code already inside this module (and any future caller that imports
# it from here for convenience) keeps working without the byte-identical
# local copy that used to live at this spot.
from bulletin.document.sections.word_of_god import add_body_with_amen  # noqa: E402


def add_offertory_rubric(doc: Document):
    """Add the offertory rubric with mixed font styling and QR code.

    The rubric text uses the standard rubric font (italic Adobe Garamond Pro)
    but URLs are rendered in Gill Sans Nova Light.  A QR code image (from the
    front-cover template) is anchored at the right edge of the text area.
    """
    from docx.shared import Pt, Inches, Emu
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    p = doc.add_paragraph(style="Body - Rubric")

    # --- Rubric text with mixed styling ---
    # Part 1: plain rubric text
    run1 = p.add_run(
        "Please place your offerings and Connection Cards in the "
        "offering plates as they are passed. You can also easily give "
        "online at "
    )

    # Part 2: URL in Gill Sans Nova Light
    run_url1 = p.add_run(GIVING_URL)
    run_url1.font.name = FONT_HEADER_FOOTER
    run_url1.italic = False

    # Part 3: connecting text
    run3 = p.add_run(
        " or by scanning the QR code. If you are online with us, "
        "please also fill out a digital Connection Card at "
    )

    # Part 4: second URL in Gill Sans Nova Light
    run_url2 = p.add_run(CONNECT_URL)
    run_url2.font.name = FONT_HEADER_FOOTER
    run_url2.italic = False

    # Part 5: closing period
    p.add_run(".")

    # --- QR code image (anchored, floating) ---
    _add_qr_code(doc, p)


def _add_qr_code(doc: Document, paragraph):
    """Add the QR code as a floating image anchored to *paragraph*.

    The QR code is extracted from the front-cover template (which is the
    document's base).  It is positioned 5.75" from the left edge of the
    page with tight text wrapping, sized at 0.75" × 0.75".
    """
    from pathlib import Path
    from docx.shared import Emu
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # Find the QR code image (EMF) already in the document from the front cover
    part = doc.part
    qr_rId = None
    for rId, rel in list(part.rels.items()):
        if (rel.reltype ==
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
                and hasattr(rel, "target_part")
                and str(rel.target_part.partname).endswith(".emf")):
            qr_rId = rId
            break

    if not qr_rId:
        # If the EMF isn't found (shouldn't happen), skip silently
        return

    # Dimensions
    cx = cy = int(0.75 * 914400)      # 0.75" in EMUs = 685,800
    pos_h = int(5.75 * 914400)        # 5.75" from left edge of page

    # Namespace declarations needed for the anchor element
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A = "http://schemas.openxmlformats.org/drawingml/2006/main"
    PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    anchor_xml = (
        f'<wp:anchor xmlns:wp="{WP}" xmlns:a="{A}" xmlns:pic="{PIC}" xmlns:r="{R}"'
        f'  distT="0" distB="0" distL="114300" distR="114300"'
        f'  simplePos="0" relativeHeight="251658240" behindDoc="0"'
        f'  locked="0" layoutInCell="1" allowOverlap="1">'
        f'  <wp:simplePos x="0" y="0"/>'
        f'  <wp:positionH relativeFrom="page">'
        f'    <wp:posOffset>{pos_h}</wp:posOffset>'
        f'  </wp:positionH>'
        f'  <wp:positionV relativeFrom="paragraph">'
        f'    <wp:posOffset>45720</wp:posOffset>'
        f'  </wp:positionV>'
        f'  <wp:extent cx="{cx}" cy="{cy}"/>'
        f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'  <wp:wrapTight wrapText="bothSides">'
        f'    <wp:wrapPolygon edited="0">'
        f'      <wp:start x="0" y="0"/>'
        f'      <wp:lineTo x="0" y="21600"/>'
        f'      <wp:lineTo x="21600" y="21600"/>'
        f'      <wp:lineTo x="21600" y="0"/>'
        f'      <wp:lineTo x="0" y="0"/>'
        f'    </wp:wrapPolygon>'
        f'  </wp:wrapTight>'
        f'  <wp:docPr id="100" name="QR Code"/>'
        f'  <wp:cNvGraphicFramePr>'
        f'    <a:graphicFrameLocks noChangeAspect="1"/>'
        f'  </wp:cNvGraphicFramePr>'
        f'  <a:graphic>'
        f'    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'      <pic:pic>'
        f'        <pic:nvPicPr>'
        f'          <pic:cNvPr id="0" name="QR Code"/>'
        f'          <pic:cNvPicPr/>'
        f'        </pic:nvPicPr>'
        f'        <pic:blipFill>'
        f'          <a:blip r:embed="{qr_rId}"/>'
        f'          <a:stretch><a:fillRect/></a:stretch>'
        f'        </pic:blipFill>'
        f'        <pic:spPr>'
        f'          <a:xfrm>'
        f'            <a:off x="0" y="0"/>'
        f'            <a:ext cx="{cx}" cy="{cy}"/>'
        f'          </a:xfrm>'
        f'          <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'        </pic:spPr>'
        f'      </pic:pic>'
        f'    </a:graphicData>'
        f'  </a:graphic>'
        f'</wp:anchor>'
    )

    anchor = parse_xml(anchor_xml)

    # Wrap in a drawing element and insert into a run in the paragraph
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    drawing = parse_xml(f'<w:drawing xmlns:w="{W_NS}"/>')
    drawing.append(anchor)

    # Insert the drawing into the first run of the paragraph
    first_run = paragraph.runs[0]._element
    first_run.insert(0, drawing)


def _add_11am_offertory(doc: Document, data: dict):
    """Add the 11am offertory section.

    The anthem/offering music at 11am may be:
      - A hymnal song (e.g. '#482 Lord of all Hopefulness') → header only
      - A choir anthem (e.g. 'Offertory, John Ness Beck') → title + stub
      - A song with YAML lyrics → full song rendering

    If the anthem has a hymnal number, it's treated as a hymn.
    Otherwise, the anthem title is shown with a stub for pasting lyrics.
    """
    from docx.shared import Pt
    from bulletin.sources.music_11am import parse_11am_identifier

    anthem_title = data.get("offertory_anthem_title", "")

    if anthem_title:
        parsed = parse_11am_identifier(anthem_title)
        if parsed["hymnal_number"]:
            # It's a hymnal song — render as header only
            add_hymn_header(
                doc,
                parsed["title"] or "",
                parsed.get("setting"),
                parsed["hymnal_number"],
                parsed["hymnal_name"],
            )
        else:
            # It's a choir anthem — render title + stub
            p = doc.add_paragraph(style="Body")
            p.add_run(anthem_title)

            p = doc.add_paragraph(style="Body")
            run = p.add_run("TODO: Paste lyrics here using Paste Unformatted text")
            run.italic = True
            run.font.name = FONT_BODY
            run.font.size = Pt(11)
    else:
        p = doc.add_paragraph(style="Body")
        p.add_run("[Anthem]")

    # Also render the regular offertory hymn if one was assigned
    offertory_song = data.get("offertory_song")
    if offertory_song:
        add_spacer(doc)
        add_communion_song_smart(doc, offertory_song)


def _add_agnus_dei_images(doc: Document):
    """Add the Agnus Dei fraction anthem with inline music notation images.

    For 11am Lent, the People's response is rendered as inline PNG images
    of the music notation rather than as text. The images are:
      - Agnus Dei Responses 1-2.png  (used for verses 1 and 2)
      - Agnus Dei Response final.png (used for verse 3)

    Layout:
      Cantor  Lamb of God, you take away the sins of the world.
      People  [inline image — 4" wide]
      (spacer)
      (repeat for verses 2 and 3)
    """
    from pathlib import Path
    from docx.shared import Inches
    from docx.oxml.ns import qn

    source_dir = Path(__file__).parent.parent.parent.parent / "source_documents"
    img_1_2 = source_dir / "Agnus Dei Responses 1-2.png"
    img_final = source_dir / "Agnus Dei Response final.png"

    cantor_lines = [
        "Lamb of God, you take away the sins of the world.",
        "Lamb of God, you take away the sins of the world.",
        "Lamb of God, you take away the sins of the world.",
    ]
    image_paths = [img_1_2, img_1_2, img_final]

    for i, (cantor_text, img_path) in enumerate(zip(cantor_lines, image_paths)):
        # Cantor line
        add_celebrant_line(doc, "Cantor", cantor_text)

        # People line with inline image
        p = doc.add_paragraph(style="Body - Dialogue")
        run_label = p.add_run("People")
        run_label.bold = True
        run_label.font.name = FONT_BODY_BOLD
        p.add_run("\t")

        if img_path.exists():
            run_img = p.add_run()
            run_img.add_picture(str(img_path), width=Inches(4))
        else:
            run_fallback = p.add_run("[Music notation image not found]")
            run_fallback.italic = True

        # Spacer between verses (but not after the last)
        if i < 2:
            add_spacer(doc)


def add_blessing_line(doc: Document, text: str):
    """Add a blessing as body text (no Celebrant label), with ✠ cross symbol.

    The final 'Amen.' is rendered bold (spoken by all the people in unison).
    """
    # Ensure the text ends with "Amen."
    if not text.rstrip().endswith("Amen."):
        text = text.rstrip() + " Amen."

    # Split off the trailing "Amen." to bold it
    body = text.rstrip()[:-5]  # everything before "Amen."
    p = doc.add_paragraph(style="Body")
    if CROSS_SYMBOL in body:
        parts = body.split(CROSS_SYMBOL)
        for i, part in enumerate(parts):
            if i > 0:
                add_cross_symbol(p)
            if part:
                p.add_run(part)
    else:
        p.add_run(body)
    run = p.add_run("Amen.")
    run.bold = True
    run.font.name = FONT_BODY_BOLD


def add_short_offertory_rubric(doc: Document):
    """Add the shorter offertory rubric for 8am and 11am services.

    No QR code, no Connection Card URL — just the giving URL in
    Gill Sans Nova Light inline with the italic rubric text.
    """
    p = doc.add_paragraph(style="Body - Rubric")

    p.add_run(
        "Please place your offerings and Connection Cards in the "
        "offering plates as they are passed. You can also easily give "
        "online at "
    )

    run_url = p.add_run(GIVING_URL)
    run_url.font.name = FONT_HEADER_FOOTER
    run_url.italic = False

    p.add_run(".")


def add_sanctus_spoken(doc: Document, lines: list[str]):
    """Add the Sanctus as bold poetry (8am) in a no-split block.

    Same line-by-line layout as the 9/11am text Sanctus, but all text
    is bold (spoken in unison by the people). ✠ is rendered via the
    bold cross symbol helper.
    """
    def _add_spoken_lines(cell):
        for line in lines:
            p = cell.add_paragraph(style="Body - Lyrics")
            if CROSS_SYMBOL in line:
                parts = line.split(CROSS_SYMBOL)
                for i, part in enumerate(parts):
                    if i > 0:
                        add_cross_symbol(p)
                    if part:
                        run = p.add_run(part)
                        run.bold = True
                        run.font.name = FONT_BODY_BOLD
            else:
                run = p.add_run(line)
                run.bold = True
                run.font.name = FONT_BODY_BOLD
    add_no_split_block(doc, _add_spoken_lines)


def add_agnus_dei_spoken(doc: Document):
    """Add the Agnus Dei (Fraction Anthem) as spoken text (8am Lent).

    Three lines with split bold — the response portion is bold:
      Lamb of God, you take away the sins of the world: **have mercy on us.**
      Lamb of God, you take away the sins of the world: **have mercy on us.**
      Lamb of God, you take away the sins of the world: **grant us peace.**
    """
    _AGNUS_DEI = [
        ("Lamb of God, you take away the sins of the world: ", "have mercy on us."),
        ("Lamb of God, you take away the sins of the world: ", "have mercy on us."),
        ("Lamb of God, you take away the sins of the world: ", "grant us peace."),
    ]
    for plain, bold in _AGNUS_DEI:
        add_body_with_bold_ending(doc, plain, bold)
