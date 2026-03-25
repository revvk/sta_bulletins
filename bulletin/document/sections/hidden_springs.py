"""
Hidden Springs (senior living) bulletin sections.

Two service types:
  - LOW (Liturgy of the Word): no Eucharist, ends with General Thanksgiving
  - HE-II (Holy Eucharist Rite II): full communion, same as main campus

The structure mirrors add_word_of_god() but without service-time branching
(no 8am/11am distinctions — only one service at Hidden Springs).
All songs get full lyrics printed (no hymnals available).
"""

from docx import Document
from docx.shared import Pt

from bulletin.config import CROSS_SYMBOL, FONT_BODY_BOLD, PREACHER_NAMES
from bulletin.data.loader import load_common_prayers
from bulletin.document.formatting import (
    add_spacer, add_heading, add_heading2, add_rubric,
    add_introductory_rubric, add_body, add_celebrant_line,
    add_people_line, add_dialogue, add_cross_symbol,
    add_hymn_header, add_song, add_song_two_column,
    add_scripture_text,
)
from bulletin.document.sections.word_of_god import (
    add_reading, add_psalm, add_gospel, add_song_smart,
    add_body_with_amen, add_nicene_creed, add_pop,
    add_confession, add_celebrant_with_cross,
    _add_gloria_spoken, _add_kyrie_spoken,
)
from bulletin.logic.rules import SeasonalRules


def add_hidden_springs_low(doc: Document, rules: SeasonalRules, data: dict):
    """Build a Liturgy of the Word (non-communion) Hidden Springs bulletin.

    Flow:
      Prelude → Processional → Opening → Collect for Purity → Gloria/Kyrie →
      Collect of the Day → Reading → Psalm → Sequence → Gospel → Sermon →
      Creed → POP → Confession → Peace → Lord's Prayer →
      General Thanksgiving → Blessing → Closing Hymn → Dismissal → Postlude
    """
    prayers = load_common_prayers()

    # --- Prelude ---
    if data.get("prelude_title"):
        add_heading2(doc, "Prelude")
        add_rubric(doc, data["prelude_title"])
        add_spacer(doc)

    # --- Word of God ---
    if rules.use_penitential_order:
        _add_hs_penitential_opening(doc, rules, data, prayers)
    else:
        _add_hs_standard_opening(doc, rules, data, prayers)

    # --- Collect of the Day ---
    add_spacer(doc)
    add_heading2(doc, "Collect of the Day")
    add_celebrant_line(doc, "Celebrant", "The Lord be with you.")
    add_people_line(doc, "People", "And also with you.")
    add_celebrant_line(doc, "Celebrant", "Let us pray.")
    add_spacer(doc)
    add_body_with_amen(doc, data["collect_text"])

    # --- Be seated for readings ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")

    # --- First Reading ---
    add_reading(doc, data["reading_1_ref"], data["reading_1_text"])

    # --- Psalm ---
    add_spacer(doc)
    add_psalm(doc, data["psalm_ref"], data.get("psalm_rubric", ""),
              data["psalm_text"])

    # --- Sequence Hymn ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Sequence Hymn")
    add_song_smart(doc, data.get("sequence_hymn"), force_single_column=True)

    # --- Gospel ---
    add_gospel(doc, data["gospel_ref"], data["gospel_book"],
               data["gospel_text"], service_time="hidden_springs")

    # --- Sermon ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Be seated.")
    add_heading2(doc, "Sermon")
    add_body(doc, data.get("preacher", ""))

    # --- Nicene Creed ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Nicene Creed")
    add_nicene_creed(doc, prayers)

    # --- Prayers of the People ---
    add_spacer(doc)
    add_heading2(doc, "Prayers of the People")
    add_pop(doc, data.get("pop_elements", []))
    add_spacer(doc)
    add_rubric(doc, data.get("pop_concluding_rubric",
                             "The Celebrant concludes with a suitable Collect."))

    # --- Confession & Absolution ---
    if not rules.no_confession_after_pop:
        add_spacer(doc)
        add_confession(doc, prayers)

    # --- The Peace ---
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "The Peace")
    add_celebrant_line(doc, "Celebrant",
                       "The peace of the Lord be always with you.")
    add_people_line(doc, "People", "And also with you.")

    # --- Lord's Prayer ---
    add_spacer(doc)
    add_heading2(doc, "The Lord\u2019s Prayer")
    lp_intro = prayers.get("lords_prayer_intro", {})
    invitation = lp_intro.get("option_1", "")
    if invitation:
        add_body(doc, invitation)
        add_spacer(doc)
    lp_lines = prayers.get("lords_prayer", [])
    if lp_lines:
        lp_text = " ".join(line.strip() for line in lp_lines)
        p = doc.add_paragraph(style="Body - People Recitation")
        run = p.add_run(lp_text)
        run.style = doc.styles["People"]

    # --- General Thanksgiving ---
    add_spacer(doc)
    add_heading2(doc, "The General Thanksgiving")
    gen_thanks = prayers.get("general_thanksgiving", {})
    intro = gen_thanks.get("introduction", "")
    if intro:
        add_celebrant_line(doc, "Celebrant", intro)
        add_spacer(doc)
    text = gen_thanks.get("text", "")
    if text:
        p = doc.add_paragraph(style="Body - People Recitation")
        run = p.add_run(text)
        run.style = doc.styles["People"]

    # --- Blessing ---
    add_spacer(doc)
    _add_hs_blessing(doc, rules, data, prayers)

    # --- Closing Hymn ---
    add_spacer(doc)
    add_heading2(doc, "Closing Hymn")
    add_song_smart(doc, data.get("closing_hymn"), force_single_column=True)

    # --- Dismissal ---
    add_spacer(doc)
    add_heading2(doc, "The Dismissal")
    add_celebrant_line(doc, "Deacon", data.get("dismissal_deacon", ""))
    add_people_line(doc, "People", data.get("dismissal_people", ""))

    # --- Postlude ---
    if data.get("postlude_title"):
        add_spacer(doc)
        add_heading2(doc, "Postlude")
        add_rubric(doc, data["postlude_title"])


def _add_hs_standard_opening(doc: Document, rules: SeasonalRules,
                              data: dict, prayers: dict):
    """Standard (non-Lent) opening for Hidden Springs."""
    add_heading(doc, "The Word of God")
    add_spacer(doc)

    # Processional
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Processional")
    add_song_smart(doc, data.get("processional"), force_single_column=True)

    # Opening Acclamation
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Opening Acclamation")
    cel_text = rules.acclamation_celebrant.replace("{cross}", CROSS_SYMBOL)
    add_celebrant_with_cross(doc, "Celebrant", cel_text)
    add_people_line(doc, "People", rules.acclamation_people)
    add_spacer(doc)

    # Collect for Purity
    if rules.include_collect_for_purity:
        add_body_with_amen(doc, prayers["collect_for_purity"])

    # Song of Praise / Kyrie
    add_spacer(doc)
    add_heading2(doc, rules.song_of_praise_label)
    if rules.is_advent:
        # No Advent wreath at Hidden Springs — use Gloria as lyrics
        _add_gloria_as_lyrics(doc, prayers)
    else:
        # Try to use the song from the sheet; fall back to Gloria as lyrics
        sop = data.get("song_of_praise")
        if sop and sop.get("sections"):
            add_song_smart(doc, sop, force_single_column=True)
        else:
            _add_gloria_as_lyrics(doc, prayers)


def _add_hs_penitential_opening(doc: Document, rules: SeasonalRules,
                                 data: dict, prayers: dict):
    """Lenten opening for Hidden Springs."""
    add_heading(doc, "A Penitential Order")
    add_spacer(doc)

    # Processional
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Processional")
    add_song_smart(doc, data.get("processional"), force_single_column=True)

    # Opening Acclamation
    add_spacer(doc)
    add_introductory_rubric(doc, "Remain standing.")
    add_heading2(doc, "Opening Acclamation")
    cel_text = rules.acclamation_celebrant.replace("{cross}", CROSS_SYMBOL)
    add_celebrant_with_cross(doc, "Celebrant", cel_text)
    add_people_line(doc, "People", rules.acclamation_people)
    add_spacer(doc)

    # Penitential sentence
    sentence = data.get("penitential_sentence", "")
    if sentence:
        sentence_ref = data.get("penitential_sentence_ref", "")
        p = doc.add_paragraph(style="Body")
        p.add_run(sentence)
        if sentence_ref:
            run = p.add_run(f" ({sentence_ref})")
            run.italic = True

    # Confession
    add_spacer(doc)
    add_introductory_rubric(doc, "Please kneel or remain standing.")
    add_confession(doc, prayers)

    # Then "The Word of God" section heading
    add_heading(doc, "The Word of God")

    # Kyrie
    add_spacer(doc)
    add_introductory_rubric(doc, "Please stand.")
    add_heading2(doc, "Kyrie")
    sop = data.get("song_of_praise")
    if sop and sop.get("sections"):
        add_song_smart(doc, sop)
    else:
        _add_kyrie_spoken(doc)


def _add_gloria_as_lyrics(doc: Document, prayers: dict):
    """Render the Gloria as lyrics for Hidden Springs (large print).

    Each section of the Gloria gets its own row in a multi-row table,
    rendered as bold People text with line breaks within each section.
    """
    from bulletin.document.formatting import (
        add_no_split_block,
    )
    from bulletin.config import FONT_BODY_BOLD

    gloria = prayers["gloria"]
    if isinstance(gloria, dict):
        sections = gloria.get("sections", [])
    else:
        sections = [gloria]

    for section_lines in sections:
        def _add_section(cell, lines=section_lines):
            p = cell.add_paragraph(style="Body - People Recitation")
            p.paragraph_format.space_before = Pt(0)
            for i, line in enumerate(lines):
                if i > 0:
                    p.runs[-1].add_break()
                run = p.add_run(line.strip())
                run.style = doc.styles["People"]
        add_no_split_block(doc, _add_section)


def _add_hs_blessing(doc: Document, rules: SeasonalRules, data: dict,
                      prayers: dict):
    """Add blessing or Prayer over the People."""
    if rules.use_prayer_over_people:
        add_heading2(doc, "Prayer over the People")
        blessing_text = data.get("blessing_text", "")
        if blessing_text:
            add_rubric(doc, "The Celebrant says a Prayer over the People.")
            add_body(doc, blessing_text)
    else:
        add_heading2(doc, "The Blessing")
        blessing_text = data.get("blessing_text", "")
        if blessing_text:
            add_body(doc, blessing_text)
