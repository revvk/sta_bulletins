"""
Main document builder that orchestrates bulletin assembly.

Coordinates data fetching, rules computation, and section assembly
to produce a complete .docx bulletin.
"""

import re
from pathlib import Path
from datetime import date, datetime

from bulletin.config import CHURCH_NAME, GIVING_URL
from bulletin.document.styles import configure_document, configure_lp_document
from bulletin.document.templates import (
    load_front_cover, append_back_cover, append_template_page, setup_footers,
)
from bulletin.document.sections.word_of_god import add_word_of_god
from bulletin.document.sections.holy_communion import add_holy_communion
from bulletin.logic.rules import get_seasonal_rules, get_dismissal_text, detect_special_service
from bulletin.data.loader import (
    load_common_prayers, load_pop_forms, load_blessings,
    get_proper_preface_text, get_preface_option_labels,
    load_placeholders, get_canonical_hymn_title,
    load_maundy_thursday, load_good_friday,
    load_palm_sunday, load_passion_gospel,
)


def _format_hs_title(title: str, dt=None) -> str:
    """Format a Hidden Springs liturgical title.

    For regular Sunday propers, prepend 'Wednesday after the'.
    For feast days and special observances, use the title as-is.
    """
    if not title or not title.strip():
        return title
    # If the title contains "Sunday" or "Proper", it's a regular week
    is_sunday_title = any(
        kw in title.lower()
        for kw in ("sunday",)
    )
    if is_sunday_title and dt and dt.weekday() == 2:
        # Avoid "Wednesday after the The ..." — drop leading "The "
        base = title
        if base.startswith("The "):
            base = base[4:]
        return f"Wednesday after the {base}"
    return title


# Cover template overrides for special services
_COVER_TEMPLATES = {
    "palm_sunday": "front_cover_palm_sunday.docx",
    "maundy_thursday": "front_cover_maundy_thursday.docx",
    "good_friday": "front_cover_good_friday.docx",
}

# Inside back cover template overrides for special services
_INSIDE_BACK_COVER_TEMPLATES = {
    "maundy_thursday": "inside_back_cover_maundy_thursday.docx",
    "good_friday": "inside_back_cover_good_friday.docx",
}


class BulletinBuilder:
    """Builds a bulletin .docx from all data sources.

    Usage:
        builder = BulletinBuilder(target_date, sheet_data, music_data,
                                  scripture_data, song_lookup_fn,
                                  service_time="9 am")
        builder.resolve_interactive()  # prompts user for choices
        doc = builder.build()
        doc.save("output/bulletin.docx")
    """

    def __init__(self, target_date: date, sheet_data, music_data,
                 scripture_readings: dict, song_lookup_fn,
                 parish_ministries: str, service_time: str = "9 am",
                 hidden_springs_data=None, report=None):
        """
        Args:
            target_date: The Sunday date.
            sheet_data: BulletinData from google_sheet module.
            music_data: ServiceMusic9am (has .slots) or list[MusicSlot], or None.
            scripture_readings: Dict of {ref: ScriptureReading} from scripture module.
            song_lookup_fn: Callable(identifier, service) -> song_data dict or None.
            parish_ministries: Formatted ministry string for POP.
            service_time: "9 am" or "11 am" (default "9 am").
            hidden_springs_data: Optional (HiddenSpringsRow, upcoming_rows) tuple.
            report: Optional bulletin.report.RunReport that this builder
                will record warnings/blockers/manual TODOs to. None
                preserves the legacy "print to stdout only" behavior.
        """
        self.target_date = target_date
        self.sheet = sheet_data
        self.music = music_data
        self.scripture = scripture_readings
        self.song_lookup = song_lookup_fn
        self.parish_ministries = parish_ministries
        self.service_time = service_time
        self.hidden_springs_data = hidden_springs_data
        self.report = report

        # Sunrise services use the 9am pipeline but display their own time
        self.is_sunrise = service_time == "sunrise"
        if self.is_sunrise:
            self.service_time = "9 am"  # pipeline behavior
            self._display_time = "6:30 am"  # shown on cover/footer
        else:
            self._display_time = service_time

        # Derived data
        self.schedule = sheet_data.schedule
        self.clergy = sheet_data.clergy

        # Hidden Springs: check if this is a HS service
        self.is_hidden_springs = hidden_springs_data is not None

        # Detect special service type (Maundy Thursday, Good Friday, etc.)
        self.special_service = detect_special_service(self.schedule.title)

        # Compute seasonal rules
        self.rules = get_seasonal_rules(
            title=self.schedule.title,
            color=self.schedule.color,
            notes=self.schedule.notes or "",
            pop_form=self.schedule.pop_form or "",
        )

        # Data that may need user input
        self.proper_preface_text = ""
        self.advent_wreath_verse = None
        self.penitential_sentence = None
        self.penitential_sentence_ref = ""
        self.blessing_text = ""
        self.eucharistic_prayer = "A"
        self.post_communion_prayer = self._resolve_closing_prayer()
        self.pop_form_key = None  # Resolved in resolve_all()
        self._missing_songs = []
        self._aac_manifest = []  # (slot_name, aac_filename) for HS services

    def get_aac_manifest(self) -> list[tuple[str, str]]:
        """Return the list of (slot_name, aac_filename) pairs for HS services."""
        return self._aac_manifest

    # ------------------------------------------------------------------
    # Warning helpers — print AND (optionally) record to the report.
    # Each builder method that catches an exception should use one of
    # these instead of bare print() so the web UI's TODO list sees the
    # same content the CLI does.
    # ------------------------------------------------------------------

    def _warn(self, message: str, *, category: str = "lookup",
              fix_hint: str | None = None) -> None:
        """Emit a recoverable warning (placeholder used)."""
        print(f"  Warning: {message}")
        if self.report is not None:
            self.report.warning(category=category, message=message,
                                fix_hint=fix_hint)

    def _block(self, message: str, *, category: str,
               fix_hint: str | None = None) -> None:
        """Emit a blocker (something the user must fix before printing)."""
        print(f"  Warning: {message}")
        if self.report is not None:
            self.report.blocker(category=category, message=message,
                                fix_hint=fix_hint)

    def resolve_all(self, prompt_fn=None, shared_resolutions=None):
        """Resolve all data that might need user input.

        Args:
            prompt_fn: Callable(question, options) -> answer.
                       If None, makes best-guess choices silently.
            shared_resolutions: Dict of previously resolved values to reuse.
                       When provided, skips prompting for these values.
                       Used when generating multiple bulletins so the user
                       is only prompted once for shared liturgical choices.
        """
        shared = shared_resolutions or {}

        self._resolve_eucharistic_prayer()

        if "proper_preface_text" in shared:
            self.proper_preface_text = shared["proper_preface_text"]
        else:
            self._resolve_proper_preface(prompt_fn)

        if "penitential_sentence" in shared:
            self.penitential_sentence = shared["penitential_sentence"]
            self.penitential_sentence_ref = shared.get(
                "penitential_sentence_ref", "")
        else:
            self._resolve_penitential_sentence(prompt_fn)

        if "advent_wreath_verse" in shared:
            self.advent_wreath_verse = shared["advent_wreath_verse"]
        else:
            self._resolve_advent_wreath(prompt_fn)

        self._resolve_blessing()

        if "pop_form_key" in shared:
            self.pop_form_key = shared["pop_form_key"]
        else:
            self._resolve_pop_version(prompt_fn)

        self._resolve_songs(prompt_fn)

    def get_shared_resolutions(self) -> dict:
        """Return resolved liturgical choices that apply across all services.

        These values (proper preface, penitential sentence, POP form, etc.)
        are the same regardless of service time and can be passed to
        subsequent builders via resolve_all(shared_resolutions=...).
        """
        result = {}
        if self.proper_preface_text:
            result["proper_preface_text"] = self.proper_preface_text
        if self.penitential_sentence:
            result["penitential_sentence"] = self.penitential_sentence
            result["penitential_sentence_ref"] = self.penitential_sentence_ref
        if self.advent_wreath_verse:
            result["advent_wreath_verse"] = self.advent_wreath_verse
        if self.pop_form_key:
            result["pop_form_key"] = self.pop_form_key
        return result

    def build(self) -> "Document":
        """Build and return the complete document."""
        # Format the date nicely
        date_str = self.target_date.strftime("%B %-d, %Y")
        service_time = self.service_time
        display_time = self._display_time

        if self.is_hidden_springs:
            return self._build_hidden_springs(date_str)

        # Use special cover template if applicable
        cover_template = _COVER_TEMPLATES.get(self.special_service)

        # Start from the front-cover template (replaces placeholders)
        doc = load_front_cover(
            date_str=date_str,
            service_time=display_time,
            liturgical_title=self.schedule.title,
            subtitle=" ",  # single space to preserve spacing when unused
            cover_template=cover_template,
        )

        # Apply page setup and register all bulletin styles
        configure_document(doc)

        # Route to special service module or standard flow
        if self.special_service == "maundy_thursday":
            self._build_maundy_thursday(doc)
        elif self.special_service == "good_friday":
            self._build_good_friday(doc)
        elif self.special_service == "palm_sunday":
            self._build_palm_sunday(doc)
        else:
            self._build_standard(doc)

        # Inside back cover for special services (before back cover)
        inside_back_template = _INSIDE_BACK_COVER_TEMPLATES.get(self.special_service)
        if inside_back_template:
            append_template_page(doc, inside_back_template)

        # Back cover (template-based, new page section)
        append_back_cover(doc)

        # Footers — added after back cover so both sections exist
        setup_footers(doc, date_str, display_time, self.schedule.title)

        return doc

    def _build_standard(self, doc):
        """Standard Sunday bulletin flow."""
        from bulletin.document.formatting import (
            add_spacer, add_introductory_rubric, add_heading2,
        )
        service_time = self.service_time

        # Prelude items (bridge between cover and liturgy)
        if service_time == "8 am":
            # 8am: shorter rubric, no Prelude heading (no music)
            add_introductory_rubric(
                doc,
                "Prayers before worship can be found in the "
                "Book of Common Prayer, p. 833-35."
            )
            add_spacer(doc)
        else:
            add_introductory_rubric(
                doc,
                "Once the Prelude begins, please refrain from further visiting "
                "and conversation as we prepare our hearts and thoughts for worship. "
                "Prayers before worship can be found in the Book of Common Prayer, "
                "p. 833-35."
            )
            add_spacer(doc)
            add_heading2(doc, "Prelude")
            add_spacer(doc)

        # Word of God
        wog_data = self._prepare_word_of_god_data()
        add_word_of_god(doc, self.rules, wog_data)

        # Holy Communion
        hc_data = self._prepare_holy_communion_data()
        add_holy_communion(doc, self.rules, hc_data)

    def _build_maundy_thursday(self, doc):
        """Maundy Thursday liturgy flow."""
        from bulletin.document.formatting import (
            add_spacer, add_introductory_rubric, add_heading2,
        )
        from bulletin.document.sections.maundy_thursday import add_maundy_thursday

        # Prelude (Maundy Thursday has a prelude like a regular service)
        add_introductory_rubric(
            doc,
            "Once the Prelude begins, please refrain from further visiting "
            "and conversation as we prepare our hearts and thoughts for worship. "
            "Prayers before worship can be found in the Book of Common Prayer, "
            "p. 833-35."
        )
        add_spacer(doc)
        add_heading2(doc, "Prelude")
        add_spacer(doc)

        wog_data = self._prepare_word_of_god_data()
        hc_data = self._prepare_holy_communion_data()
        mt_data = self._prepare_maundy_thursday_data()

        add_maundy_thursday(doc, self.rules, wog_data, hc_data, mt_data)

    def _build_good_friday(self, doc):
        """Good Friday liturgy flow — no prelude, no Eucharist."""
        from bulletin.document.sections.good_friday import add_good_friday

        # No prelude on Good Friday — ministers enter in silence
        wog_data = self._prepare_word_of_god_data()
        gf_data = self._prepare_good_friday_data()

        add_good_friday(doc, self.rules, wog_data, gf_data)

    def _build_palm_sunday(self, doc):
        """Palm Sunday liturgy flow — Liturgy of the Palms + Word of God + Communion."""
        from bulletin.document.formatting import (
            add_spacer, add_introductory_rubric, add_heading2,
        )
        from bulletin.document.sections.palm_sunday import (
            add_liturgy_of_the_palms, add_palm_sunday_word_of_god,
        )

        service_time = self.service_time

        # Palm Sunday 9/11am begins outside — no initial rubric or Prelude.
        # 8am still gets the short prayer rubric (no prelude at 8am anyway).
        if service_time == "8 am":
            add_introductory_rubric(
                doc,
                "Prayers before worship can be found in the "
                "Book of Common Prayer, p. 833-35."
            )
            add_spacer(doc)

        wog_data = self._prepare_word_of_god_data()
        ps_data = self._prepare_palm_sunday_data()

        # Palm Sunday processional is sung outside without hymnals,
        # so always print full lyrics (even for 11am).
        if service_time == "11 am" and wog_data.get("processional_songs"):
            refreshed = []
            for proc in wog_data["processional_songs"]:
                if not proc.get("sections"):
                    full_song = self.song_lookup(proc.get("title", ""), self.service_time)
                    if full_song and full_song.get("sections"):
                        refreshed.append(self._apply_canonical_title(full_song))
                    else:
                        refreshed.append(proc)
                else:
                    refreshed.append(proc)
            wog_data["processional_songs"] = refreshed
            wog_data["processional"] = refreshed[0] if refreshed else None

        # Liturgy of the Palms
        add_liturgy_of_the_palms(doc, self.rules, wog_data, ps_data)

        # Word of God (with Passion Gospel in parts)
        add_palm_sunday_word_of_god(doc, self.rules, wog_data, ps_data)

        # Holy Communion (standard)
        hc_data = self._prepare_holy_communion_data()
        add_holy_communion(doc, self.rules, hc_data)

    def _build_hidden_springs(self, date_str: str) -> "Document":
        """Build a Hidden Springs (senior living) bulletin.

        Uses the senior_living front/back cover templates, large-print styles,
        and either LOW (Liturgy of the Word) or HE-II flow depending on
        the service type in the Hidden Springs Planner sheet.
        """
        from bulletin.document.sections.hidden_springs import (
            add_hidden_springs_low,
        )
        from bulletin.document.sections.holy_communion import add_holy_communion

        hs_row, upcoming = self.hidden_springs_data

        # Determine the liturgical title for the front cover
        # For feast days, use the feast name directly.
        # For regular weeks, format as "Wednesday after the [Sunday Name]"
        title = _format_hs_title(hs_row.title, hs_row.date)

        # Front cover — uses senior_living template
        doc = load_front_cover(
            date_str=date_str,
            service_time="11:00 am",
            liturgical_title=title,
            subtitle=date_str,
            cover_template="senior_living_front_cover.docx",
        )

        # Apply large-print page setup and styles
        configure_lp_document(doc)

        # Prepare data
        data = self._prepare_hidden_springs_wog_data()

        is_communion = hs_row.service_type.upper().startswith("HE")

        if is_communion:
            # HE-II: Word of God (reuse standard) + Holy Communion
            # Use the LOW section builder for Word of God (it handles
            # everything through The Peace), then add communion
            add_hidden_springs_low(doc, self.rules, data)

            # Remove the General Thanksgiving, Lord's Prayer, blessing,
            # closing hymn, dismissal, and postlude that LOW added —
            # Actually, for HE-II we should build it differently.
            # Let's use the standard flow instead.
            pass  # TODO: implement HE-II variant
        else:
            # LOW: Liturgy of the Word (no Eucharist)
            add_hidden_springs_low(doc, self.rules, data)

        # Back cover with upcoming services
        append_back_cover(
            doc,
            template_name="senior_living_back_cover.docx",
        )
        self._fill_upcoming_services(doc, upcoming)

        # Footers — LP page is 8.5" with 0.5" margins = 7.5" text width
        from bulletin.config import LP_PAGE_WIDTH_INCHES, LP_MARGIN_INCHES
        lp_text_width = LP_PAGE_WIDTH_INCHES - 2 * LP_MARGIN_INCHES
        setup_footers(doc, date_str, "11:00 am", title, page_width=lp_text_width)

        return doc

    def _prepare_hidden_springs_wog_data(self) -> dict:
        """Prepare Word of God data using Hidden Springs planner row."""
        from bulletin.config import PREACHER_NAMES
        from bulletin.sources.songs import (
            hs_lookup_song, resolve_aac_file, slice_song_verses,
        )

        hs_row, _ = self.hidden_springs_data

        # Look up songs from HS catalog (with AAC tracking), falling back
        # to songs.yaml for lyrics only
        def hs_lookup(title, slot_name=""):
            if not title:
                return None
            song = hs_lookup_song(title)
            if song:
                # Resolve AAC file and add to manifest
                vc = song.get("_hs_verse_count")
                inst = song.get("_hs_instrument")
                aac = resolve_aac_file(song, vc, inst)
                if aac:
                    self._aac_manifest.append((slot_name, aac))
                elif not song.get("_hs_fallback"):
                    self._aac_manifest.append(
                        (slot_name, "[NO AAC FILE FOUND]"))

                # Slice verses if a count was specified
                if vc is not None and song.get("sections"):
                    song = slice_song_verses(song, vc)

                return self._apply_canonical_title(song)
            return None

        processional = hs_lookup(hs_row.processional, "Processional")

        # Song of praise: if "Gloria" is specified, use the spoken Gloria
        song_of_praise = None
        sop_title = (hs_row.song_of_praise or "").strip()
        if sop_title.lower() not in ("gloria", ""):
            song_of_praise = hs_lookup(sop_title, "Song of Praise")
        # If it's "Gloria" or no song found, the section builder will
        # fall back to the spoken Gloria

        sequence = hs_lookup(hs_row.sequence, "Sequence")
        closing = hs_lookup(hs_row.recessional, "Closing")

        # Scripture readings
        reading_1_ref = hs_row.reading or ""
        reading_1 = self.scripture.get("reading")

        psalm_raw = hs_row.psalm or ""
        psalm_ref = re.split(r"[\n\r]+", psalm_raw)[0].strip()
        psalm_ref = re.sub(
            r"\s+(responsively|unison|in unison|antiphonally).*$",
            "", psalm_ref, flags=re.IGNORECASE,
        ).strip()
        psalm_text = []
        if psalm_ref:
            try:
                from bulletin.sources.psalms import get_psalm
                psalm_selection = get_psalm(psalm_ref)
                psalm_text = psalm_selection.to_lines()
            except Exception as e:
                self._warn(f"Could not look up psalm: {e}",
                           category="psalm",
                           fix_hint=f"Check that '{psalm_ref}' is a valid "
                                    "psalm reference; bulletin contains "
                                    "no psalm text.")

        gospel_ref = hs_row.gospel or ""
        gospel = self.scripture.get("gospel")
        gospel_book = gospel_ref.split()[0] if gospel_ref else ""

        # Psalm rubric from the sheet
        psalm_field = psalm_raw.lower()
        if "half verse" in psalm_field:
            psalm_rubric = "Read responsively by half verse."
        elif "responsiv" in psalm_field:
            psalm_rubric = "Read responsively by whole verse."
        elif "antiphon" in psalm_field:
            psalm_rubric = "Read antiphonally."
        else:
            psalm_rubric = "Read in unison."

        # Preacher
        preacher_short = hs_row.preacher or ""
        preacher = PREACHER_NAMES.get(preacher_short.strip(), preacher_short)

        # POP elements
        pop_elements = self._prepare_pop_elements()

        # Dismissal
        from bulletin.logic.rules import get_dismissal_text
        dismissal_num = hs_row.dismissal or "3"
        deacon_text, people_text = get_dismissal_text(
            dismissal_num, self.rules.dismissal_has_alleluia)

        # Resolve AAC files for prelude/postlude (no lyrics needed)
        if hs_row.prelude:
            pre_song = hs_lookup_song(hs_row.prelude)
            if pre_song:
                aac = resolve_aac_file(pre_song)
                if aac:
                    self._aac_manifest.append(("Prelude", aac))
        if hs_row.postlude:
            post_song = hs_lookup_song(hs_row.postlude)
            if post_song:
                aac = resolve_aac_file(post_song)
                if aac:
                    self._aac_manifest.append(("Postlude", aac))

        return {
            "service_time": "hidden_springs",
            "processional": processional,
            "processional_songs": [processional] if processional else [],
            "song_of_praise": song_of_praise,
            "sequence_hymn": sequence,
            "collect_text": self._get_collect_text(),
            "reading_1_ref": reading_1_ref,
            "reading_1_text": reading_1 or reading_1_ref,
            "psalm_ref": psalm_ref,
            "psalm_rubric": psalm_rubric,
            "psalm_text": psalm_text or [],
            "gospel_ref": gospel_ref,
            "gospel_book": gospel_book,
            "gospel_text": gospel or gospel_ref,
            "preacher": preacher,
            "pop_elements": pop_elements,
            "pop_concluding_rubric": self._get_pop_concluding_rubric(),
            "advent_wreath_verse": None,
            "advent_hymnal_ref": None,
            "penitential_sentence": self.penitential_sentence,
            "penitential_sentence_ref": self.penitential_sentence_ref,
            "closing_hymn": closing,
            "prelude_title": hs_row.prelude or "",
            "postlude_title": hs_row.postlude or "",
            "blessing_text": self.blessing_text,
            "dismissal_deacon": deacon_text,
            "dismissal_people": people_text,
        }

    def _fill_upcoming_services(self, doc, upcoming: list):
        """Fill upcoming services on the back cover with formatted paragraphs.

        Finds {{Upcoming_Service_N}} placeholder paragraphs in the document
        and replaces them with properly formatted content:
          Line 1: Month DD, YYYY: HH:MM am - Service Type  (bold)
          Line 2: Name of the day
          Line 3: Clergy Name  Role (role in italics)
        """
        from docx.shared import Pt
        from bulletin.config import PREACHER_NAMES, FONT_BODY_BOLD

        # Build data for each slot
        service_info = []
        for row in upcoming[:3]:
            if row.date is None:
                service_info.append(None)
                continue

            date_str = row.date.strftime("%B %-d, %Y")
            svc_type = ("Holy Communion"
                        if row.service_type.upper().startswith("HE")
                        else "Liturgy of the Word")
            title = _format_hs_title(row.title, row.date)

            preacher_short = row.preacher or ""
            preacher = PREACHER_NAMES.get(
                preacher_short.strip(), preacher_short)
            role = ("Preacher and Celebrant"
                    if row.service_type.upper().startswith("HE")
                    else "Preacher and Officiant")

            service_info.append({
                "header": f"{date_str}: 11:00 am \u2013 {svc_type}",
                "title": title,
                "clergy_name": preacher,
                "clergy_role": role,
            })

        # Pad to 3 slots
        while len(service_info) < 3:
            service_info.append(None)

        # Find and replace placeholder paragraphs
        body = doc.element.body
        from docx.oxml.ns import qn
        for p_elem in list(body.iter(qn("w:p"))):
            full_text = "".join(
                (t.text or "")
                for t in p_elem.iter(qn("w:t"))
            )
            for i in range(3):
                placeholder = f"{{{{Upcoming_Service_{i+1}}}}}"
                if placeholder in full_text:
                    info = service_info[i]
                    if info is None:
                        # Clear the placeholder
                        for t in p_elem.iter(qn("w:t")):
                            t.text = ""
                    else:
                        self._replace_with_formatted_service(
                            p_elem, doc, info)
                    break

    @staticmethod
    def _replace_with_formatted_service(p_elem, doc, info: dict):
        """Replace a placeholder paragraph with formatted service info.

        Clears all existing runs, then adds:
          Line 1 (bold): date/time - service type
          Line 2: liturgical title
          Line 3: clergy name + role (italic)
        """
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        from docx.shared import Pt
        from bulletin.config import FONT_BODY, FONT_BODY_BOLD

        # Clear existing runs
        for r in list(p_elem.findall(qn("w:r"))):
            p_elem.remove(r)

        def _add_run(text, bold=False, italic=False, break_before=False):
            """Add a formatted run to the paragraph."""
            r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve">{{}}</w:t></w:r>')
            rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
            if bold:
                rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
                font_name = FONT_BODY_BOLD
            else:
                font_name = FONT_BODY
            if italic:
                rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
            # Set font
            rFonts = parse_xml(
                f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" '
                f'w:hAnsi="{font_name}"/>'
            )
            rPr.append(rFonts)
            r.insert(0, rPr)
            # Set text
            t_elem = r.find(qn("w:t"))
            t_elem.text = text
            if break_before:
                br = parse_xml(f'<w:br {nsdecls("w")}/>')
                r.insert(list(r).index(t_elem), br)
            p_elem.append(r)

        # Line 1: bold header
        _add_run(info["header"], bold=True)
        # Line 2: title
        _add_run(info["title"], break_before=True)
        # Line 3: clergy name + role (italic)
        if info["clergy_name"]:
            _add_run(info["clergy_name"] + "  ", break_before=True)
            _add_run(info["clergy_role"], italic=True)

    # ------------------------------------------------------------------
    # Resolution methods
    # ------------------------------------------------------------------

    def _resolve_eucharistic_prayer(self):
        """Determine which Eucharistic Prayer to use from the sheet."""
        ep = self.schedule.eucharistic_prayer or "A"
        self.eucharistic_prayer = ep.strip().upper()

    def _resolve_closing_prayer(self) -> str:
        """Map the rota's closing_prayer field to prayer key 'a' or 'b'.

        The sheet uses:
          'Almighty'    → prayer_b (Almighty and everliving God)
          'Eternal God' → prayer_a (Eternal God, heavenly Father)
        Default: 'a'
        """
        raw = (self.schedule.closing_prayer or "").strip().lower()
        if raw.startswith("almighty"):
            return "b"
        # "eternal god", "eternal", or empty → default
        return "a"

    def _resolve_proper_preface(self, prompt_fn):
        """Resolve the proper preface text."""
        key = self.rules.proper_preface_key
        options = self.rules.proper_preface_options

        if not options:
            # Single preface for this season
            self.proper_preface_text = get_proper_preface_text(key)
        elif prompt_fn and self.rules.prompt_preface:
            # Multiple options — ask the user
            labels = get_preface_option_labels(key)
            question = f"Proper Preface for '{self.schedule.title}':"
            option_strs = [f"  {i+1}. {label}" for i, (_, label) in enumerate(labels)]
            answer = prompt_fn(question, [label for _, label in labels])
            # Match answer to option key
            for opt_key, label in labels:
                if label == answer or answer == opt_key:
                    self.proper_preface_text = get_proper_preface_text(key, opt_key)
                    return
            # Default to first option
            self.proper_preface_text = get_proper_preface_text(key, options[0])
        else:
            # No prompt function — default to first option
            self.proper_preface_text = get_proper_preface_text(key, options[0])

    def _resolve_penitential_sentence(self, prompt_fn):
        """For Lent 2-5, pick a Penitential Order scripture sentence."""
        if not self.rules.use_penitential_order or self.rules.use_decalogue:
            return

        prayers = load_common_prayers()
        sentences = prayers.get("penitential_sentences", [])

        if prompt_fn and sentences:
            question = "Penitential Order scripture sentence:"
            options = [f'{s["reference"]}: {s["text"][:60]}...' for s in sentences]
            answer = prompt_fn(question, options)
            for s in sentences:
                if s["reference"] in answer:
                    self.penitential_sentence = s["text"]
                    self.penitential_sentence_ref = s["reference"]
                    return

        # Default to first sentence
        if sentences:
            self.penitential_sentence = sentences[0]["text"]
            self.penitential_sentence_ref = sentences[0]["reference"]

    def _resolve_advent_wreath(self, prompt_fn):
        """For Advent, resolve the O Antiphon verse for the wreath lighting."""
        if not self.rules.is_advent:
            return

        if prompt_fn:
            answer = prompt_fn(
                "Advent wreath: Enter the O Antiphon verse lines "
                "(one per line, or leave blank for verse 1 only):",
                []
            )
            if answer:
                self.advent_wreath_verse = answer.strip().split("\n")

    def _resolve_blessing(self):
        """Look up the seasonal blessing from BOS data.

        For Lent, selects the specific weekly Prayer over the People.
        For other seasons, uses form2 (single-sentence short form).
        """
        blessings = load_blessings()
        season = self.rules.season

        # Map season to blessing key
        season_map = {
            "advent": "advent",
            "christmas": "christmas",
            "epiphany": "epiphany",
            "lent": "lent",
            "easter": "easter",
            "pentecost_day": "pentecost",
            "ordinary": None,
        }

        blessing_key = season_map.get(season)
        if not blessing_key or blessing_key not in blessings:
            return

        blessing_data = blessings[blessing_key]

        if blessing_key == "lent":
            # Lent uses specific weekly Prayer over the People
            week_key = self._get_lent_week_key(self.schedule.title)
            if week_key and week_key in blessing_data:
                week = blessing_data[week_key]
                self.blessing_text = (
                    week.get("prayer", "") if isinstance(week, dict)
                    else str(week)
                )
        else:
            # Non-Lent seasons use form2 (short form)
            form2 = blessing_data.get("form2", {})
            sentences = form2.get("sentences", [])
            if sentences:
                self.blessing_text = sentences[0].get("celebrant", "")

    @staticmethod
    def _get_lent_week_key(title: str) -> str:
        """Determine the Lent week key from the liturgical title."""
        t = title.lower()
        if "ash wednesday" in t:
            return "ash_wednesday"
        if "palm sunday" in t or "passion" in t or "maundy" in t:
            return "palm_sunday_through_maundy_thursday"
        ordinals = {
            "first": "1", "second": "2", "third": "3",
            "fourth": "4", "fifth": "5",
        }
        for word, num in ordinals.items():
            if word in t and "lent" in t:
                return f"lent_{num}"
        import re
        m = re.search(r"lent\s*(\d)", t)
        if m:
            return f"lent_{m.group(1)}"
        return ""

    def _resolve_songs(self, prompt_fn):
        """Look up all song lyrics from music data."""
        self._missing_songs = []
        if self.service_time == "8 am":
            return  # No music at 8am
        service_key = self.service_time  # e.g. "9 am" or "11 am"

        slots = self._get_music_slots()
        for slot in slots:
            if not slot.song_title:
                continue
            # For 11am, skip the Anthem slot (handled as raw title, not song)
            if self.service_time == "11 am" and slot.service_part.lower() == "anthem":
                continue
            song = self.song_lookup(slot.song_title, service_key)
            if not song:
                # For 11am, hymnal-only songs are expected to have no lyrics
                if self.service_time == "11 am":
                    from bulletin.sources.music_11am import parse_11am_identifier
                    parsed = parse_11am_identifier(slot.song_title)
                    if parsed["hymnal_number"]:
                        continue  # hymnal song — no lyrics expected
                self._missing_songs.append(
                    f"{slot.service_part}: {slot.song_title}")
                if self.report is not None:
                    # URL-encode just enough for the deep-link query string
                    from urllib.parse import urlencode
                    qs = urlencode({"title": slot.song_title})
                    self.report.warning(
                        category="song",
                        message=(
                            f"Missing lyrics for \"{slot.song_title}\" "
                            f"({self.service_time}, {slot.service_part})"
                        ),
                        fix_hint=(
                            "Bulletin will print the title only — paste "
                            "the lyrics into the .docx by hand, or add "
                            "the song to songs.yaml so future bulletins "
                            "include them automatically."
                        ),
                        link=f"/songs/new?{qs}",
                    )

        if self._missing_songs and prompt_fn:
            prompt_fn(
                f"Warning: Could not find lyrics for:\n" +
                "\n".join(f"  - {s}" for s in self._missing_songs),
                ["Continue anyway"]
            )

    def _resolve_pop_version(self, prompt_fn):
        """Resolve which Prayers of the People form version to use.

        The Google Sheet may specify a version in parentheses, e.g.,
        "III (immigration)". If no version is specified and there are
        alternative versions available in the YAML, the user is prompted
        to choose.

        Versioned forms are stored as separate top-level keys in the YAML:
          form_III          (default)
          form_III_immigration
          form_III_hidden_springs
        """
        pop_forms = load_pop_forms()
        key = self._get_pop_form_key()

        # If the exact key exists, use it directly
        if key in pop_forms:
            self.pop_form_key = key
        else:
            # Versioned key doesn't exist in YAML — fall back to base form
            # Base key is the part before any version suffix (e.g., form_III)
            base_key = key
            for prefix in ("form_VI", "form_IV", "form_V", "form_III",
                           "form_II", "form_I"):
                if key.startswith(prefix + "_") or key == prefix:
                    base_key = prefix
                    break
            self.pop_form_key = base_key if base_key in pop_forms else "form_I"

        # Check if there are versioned alternatives for the resolved base form
        # Only applies to standard forms (form_I through form_VI)
        base = self.pop_form_key
        # If pop_form_key is itself a version (e.g., form_III_immigration),
        # extract the base for scanning alternatives
        for prefix in ("form_VI", "form_IV", "form_V", "form_III",
                       "form_II", "form_I"):
            if base.startswith(prefix):
                base = prefix
                break

        version_keys = sorted(
            k for k in pop_forms
            if k.startswith(base + "_") and k != base
        )

        if not version_keys or not prompt_fn:
            return

        # Build options for the user
        base_title = pop_forms[base].get("title", base)
        options = [f"{base_title} (default)"]
        for vk in version_keys:
            vt = pop_forms[vk].get("title", vk)
            options.append(vt)

        answer = prompt_fn(
            f"Multiple versions of {base_title} are available:", options
        )

        # Match user's answer to a version
        if answer:
            if "default" in answer:
                self.pop_form_key = base
                return
            for vk in version_keys:
                vt = pop_forms[vk].get("title", vk)
                if vt in answer or vk in answer:
                    self.pop_form_key = vk
                    return
        # Keep current selection (base or sheet-specified version)

    # ------------------------------------------------------------------
    # Data preparation for section modules
    # ------------------------------------------------------------------

    def _prepare_word_of_god_data(self) -> dict:
        """Prepare the data dict for add_word_of_god."""
        # Look up songs from music data
        # Look up processional(s) — may be "Processional" or "Processional 1"/"Processional 2"
        processional_songs = []
        for slot_name in ["Processional 1", "Processional 2", "Processional 3"]:
            song = self._lookup_slot(slot_name)
            if song:
                processional_songs.append(song)
        if not processional_songs:
            single = self._lookup_slot("Processional")
            if single:
                processional_songs = [single]
        # For backward compatibility, also provide single processional
        processional = processional_songs[0] if processional_songs else None

        song_of_praise = self._lookup_slot("Song of Praise")
        sequence = self._lookup_slot("Sequence")

        # Scripture readings
        # Note: fetch_readings returns {label: ScriptureReading} using
        # "reading" and "gospel" as keys (set up in generate.py)
        reading_1_ref = self.schedule.reading or ""
        reading_1 = self.scripture.get("reading")

        psalm_raw = self.schedule.psalm or ""
        # Clean: sheet may contain "Psalm 121\nunison" — extract just the reference
        psalm_ref = re.split(r"[\n\r]+", psalm_raw)[0].strip()
        psalm_ref = re.sub(
            r"\s+(responsively|unison|in unison|antiphonally).*$",
            "", psalm_ref, flags=re.IGNORECASE,
        ).strip()
        psalm_text: list[str] = []
        if psalm_ref:
            try:
                from bulletin.sources.psalms import get_psalm
                psalm_selection = get_psalm(psalm_ref)
                psalm_text = psalm_selection.to_lines()
            except (ValueError, Exception) as e:
                self._warn(f"Could not look up psalm: {e}",
                           category="psalm",
                           fix_hint=f"Check that '{psalm_ref}' is a valid "
                                    "psalm reference; bulletin contains "
                                    "no psalm text.")
                psalm_text = []

        gospel_ref = self.schedule.gospel or ""
        gospel = self.scripture.get("gospel")

        # Extract gospel book name from reference
        gospel_book = gospel_ref.split()[0] if gospel_ref else ""

        # Preacher from clergy rota — expand short name to full title
        from bulletin.config import PREACHER_NAMES
        preacher_short = ""
        if self.clergy:
            if self.service_time == "8 am":
                preacher_short = self.clergy.preacher_8am or ""
            elif self.service_time == "11 am":
                preacher_short = self.clergy.preacher_11am or ""
            else:
                preacher_short = self.clergy.preacher_9am or ""
        preacher = PREACHER_NAMES.get(preacher_short.strip(), preacher_short)

        # POP elements
        pop_elements = self._prepare_pop_elements()

        return {
            "service_time": self.service_time,
            "processional": processional,
            "processional_songs": processional_songs,
            "song_of_praise": song_of_praise,
            "sequence_hymn": sequence,
            "collect_text": self._get_collect_text(),
            "reading_1_ref": reading_1_ref,
            "reading_1_text": reading_1 or reading_1_ref,
            "psalm_ref": psalm_ref,
            "psalm_rubric": self._get_psalm_rubric(),
            "psalm_text": psalm_text or [],
            "gospel_ref": gospel_ref,
            "gospel_book": gospel_book,
            "gospel_text": gospel or gospel_ref,
            "preacher": preacher,
            "pop_elements": pop_elements,
            "pop_concluding_rubric": self._get_pop_concluding_rubric(),
            "advent_wreath_verse": self.advent_wreath_verse,
            "advent_hymnal_ref": "#56 (Hymnal 1982)",
            "penitential_sentence": self.penitential_sentence,
            "penitential_sentence_ref": self.penitential_sentence_ref,
        }

    def _prepare_holy_communion_data(self) -> dict:
        """Prepare the data dict for add_holy_communion."""
        offertory = self._lookup_slot("Offertory")
        # At 8am and 11am, always use the text Sanctus (with cross)
        # rather than a hymnal reference.
        if self.service_time in ("8 am", "11 am"):
            sanctus = None
        else:
            sanctus = self._lookup_slot("Sanctus")
        closing = self._lookup_slot("Recessional")
        fraction = self._lookup_slot("Fraction") if self.rules.use_fraction_anthem else None

        # Communion songs
        comm_songs = []
        for slot_name in ["Communion 1", "Communion 2", "Communion 3", "Communion 4"]:
            song = self._lookup_slot(slot_name)
            if song:
                comm_songs.append(song)

        # Postlude: reuse various song snippets (this is a medley typically)
        postlude_songs = []

        # Dismissal text
        dismissal_num = self.schedule.dismissal or "3"
        deacon_text, people_text = get_dismissal_text(
            dismissal_num, self.rules.dismissal_has_alleluia)

        # Offertory anthem title (11am only — raw text from the Anthem field)
        offertory_anthem_title = None
        if self.service_time == "11 am":
            slots = self._get_music_slots()
            for slot in slots:
                if slot.service_part and slot.service_part.lower() == "anthem":
                    offertory_anthem_title = slot.song_title
                    break

        return {
            "service_time": self.service_time,
            "offertory_song": offertory,
            "offertory_anthem_title": offertory_anthem_title,
            "sanctus_song": sanctus,
            "communion_songs": comm_songs,
            "closing_hymn": closing,
            "postlude_songs": postlude_songs,
            "eucharistic_prayer": self.eucharistic_prayer,
            "proper_preface_text": self.proper_preface_text,
            "fraction_song": fraction,
            "blessing_text": self.blessing_text,
            "dismissal_deacon": deacon_text,
            "dismissal_people": people_text,
            "post_communion_prayer": self.post_communion_prayer,
            "include_lev": True,
        }

    def _get_music_slots(self) -> list:
        """Return the list of MusicSlot objects regardless of music data format.

        Handles both ServiceMusic9am (has .slots) and list[MusicSlot].
        """
        if not self.music:
            return []
        if hasattr(self.music, "slots"):
            return self.music.slots
        if isinstance(self.music, list):
            return self.music
        return []

    @staticmethod
    def _apply_canonical_title(song_data: dict) -> dict:
        """Replace the title with the canonical Hymnal 1982 first line.

        Only applies to regular hymns (all-digit hymnal numbers), not
        service music (S-prefix) or songs without hymnal numbers.

        Returns a new dict if the title was changed, otherwise the original.
        """
        if not song_data:
            return song_data
        hymnal_num = str(song_data.get("hymnal_number", "") or "")
        if hymnal_num and hymnal_num.isdigit():
            canonical = get_canonical_hymn_title(int(hymnal_num))
            if canonical and canonical != song_data.get("title"):
                original = song_data.get("title", "")
                song_data = dict(song_data)  # shallow copy to avoid mutating cache
                song_data["title"] = canonical
                if original:
                    print(f"  Hymn #{hymnal_num}: '{original}' → '{canonical}'")
        return song_data

    def _lookup_slot(self, slot_name: str) -> dict | None:
        """Look up a song from the music data for a given service slot.

        For 11am services, distinguishes between:
          - Service music (S-prefix hymnal numbers like S91, S129):
            prints full lyrics from the YAML.
          - Regular hymns (all-digit hymnal numbers like 686, 473):
            renders header-only (title + hymnal reference, no lyrics).
          - Non-hymnal songs (no hymnal number):
            prints full lyrics from the YAML.

        When no YAML entry is found for an 11am song with a hymnal number,
        constructs a minimal stub dict for header-only rendering.
        """
        slots = self._get_music_slots()
        # Special weekday services (7 pm) use the 11am music pool
        service_key = self.service_time
        uses_hymnals = self.service_time in ("11 am", "7 pm")

        for slot in slots:
            if slot.service_part and slot.service_part.lower() == slot_name.lower():
                if slot.song_title:
                    song = self.song_lookup(slot.song_title, service_key)
                    if song:
                        # For 11am: regular hymns → header only
                        if uses_hymnals:
                            hymnal_num = str(song.get("hymnal_number", "") or "")
                            hymnal_name = song.get("hymnal_name")
                            # If the song dict doesn't carry a hymnal number,
                            # check the identifier from the slot (e.g., "#493 ...")
                            if not (hymnal_num and hymnal_num.isdigit()):
                                from bulletin.sources.music_11am import parse_11am_identifier
                                parsed_id = parse_11am_identifier(slot.song_title)
                                if parsed_id["hymnal_number"] and parsed_id["hymnal_number"].isdigit():
                                    hymnal_num = parsed_id["hymnal_number"]
                                    hymnal_name = parsed_id.get("hymnal_name") or "Hymnal 1982"
                            # All-digit number = regular hymn → strip lyrics
                            if hymnal_num and hymnal_num.isdigit():
                                return self._apply_canonical_title({
                                    "title": song["title"],
                                    "hymnal_number": hymnal_num,
                                    "hymnal_name": hymnal_name,
                                    "tune_name": song.get("tune_name"),
                                    "sections": [],  # header only
                                })
                            # S-prefix or no number → keep full lyrics
                        return self._apply_canonical_title(song)
                    # For 11am: construct a hymnal-only stub if it has a hymnal number
                    if uses_hymnals:
                        from bulletin.sources.music_11am import parse_11am_identifier
                        parsed = parse_11am_identifier(slot.song_title)
                        if parsed["hymnal_number"]:
                            # Include setting in title with parentheses if available
                            title = parsed["title"] or ""
                            setting = parsed.get("setting")
                            if setting and title:
                                title = f"{title} ({setting})"
                            return self._apply_canonical_title({
                                "title": title,
                                "hymnal_number": parsed["hymnal_number"],
                                "hymnal_name": parsed["hymnal_name"],
                                "tune_name": None,
                                "sections": [],  # empty → header only
                            })
                    return None
        return None

    def _get_psalm_rubric(self) -> str:
        """Determine the psalm rubric from the sheet's psalm field.

        The 8am service always reads the psalm in unison, regardless
        of what the Google Sheet specifies.
        """
        if self.service_time == "8 am":
            return "Read in unison."
        psalm_field = (self.schedule.psalm or "").lower()
        if "half verse" in psalm_field:
            return "Read responsively by half verse."
        if "responsiv" in psalm_field:
            return "Read responsively by whole verse."
        if "antiphon" in psalm_field:
            return "Read antiphonally."
        if "men and women" in psalm_field or "alternating" in psalm_field:
            return "Read alternating between men and women."
        # "unison" or default
        return "Read in unison."

    def _get_collect_text(self) -> str:
        """Look up the Collect of the Day from BCP data."""
        try:
            from bulletin.sources.collects import get_collect
            collect = get_collect(self.schedule.title, self.target_date)
            if collect:
                return collect
        except Exception as e:
            self._block(f"Could not look up collect: {e}",
                        category="collect",
                        fix_hint=f"No matching collect found for "
                                 f"'{self.schedule.title}'. Bulletin shows "
                                 "'[Collect of the Day]' placeholder — "
                                 "either fix the title in the schedule or "
                                 "add the collect to collects.yaml.")
        return "[Collect of the Day]"

    def _get_pop_concluding_rubric(self) -> str:
        """Get the concluding rubric for the resolved POP form."""
        pop_forms = load_pop_forms()
        form_key = self.pop_form_key or self._get_pop_form_key()
        form = pop_forms.get(form_key, {})
        return form.get(
            "concluding_rubric",
            "The Celebrant concludes with a suitable Collect.",
        )

    def _prepare_pop_elements(self) -> list[dict]:
        """Build POP elements from the appropriate form."""
        pop_forms = load_pop_forms()
        placeholders = load_placeholders()
        liturgical_names = placeholders.get("liturgical_names", {})

        # Use the version resolved during resolve_all(), or fall back
        form_key = self.pop_form_key or self._get_pop_form_key()
        form = pop_forms.get(form_key)

        if not form:
            return [{"type": "leader",
                     "text": f"[Prayers of the People form '{form_key}' not found]"}]

        elements = form.get("elements", [])

        # Substitute placeholders. Static name/role values come from
        # placeholders.yaml; {parish_ministries} is computed weekly and
        # {departed} is left blank for hand-editing.
        subs = {
            "{presiding_bishop}": liturgical_names.get("presiding_bishop", "N."),
            "{bishop}": liturgical_names.get("bishop", "N."),
            "{bishop_coadjutor}": liturgical_names.get("bishop_coadjutor", ""),
            "{clergy_names}": ", ".join(liturgical_names.get("clergy_first_names", [])),
            "{president}": liturgical_names.get("president", "N."),
            "{governor}": liturgical_names.get("governor", "N."),
            "{mayor}": liturgical_names.get("mayor", "N."),
            "{parish_ministries}": self.parish_ministries,
            "{departed}": "",  # Left blank; filled in manually
            "{immigration_detainees}": placeholders.get(
                "immigration_detainees", ""),
        }

        result = []
        for elem in elements:
            new_elem = dict(elem)
            for field in ("text", "leader_text", "people_text"):
                if field in new_elem:
                    for placeholder, value in subs.items():
                        new_elem[field] = new_elem[field].replace(placeholder, value)
                    # Clean up spacing after empty placeholder substitution:
                    # collapse multiple spaces to one, remove space between
                    # cross symbol and following punctuation (e.g. "✠ ," → "✠,")
                    new_elem[field] = re.sub(r" {2,}", " ", new_elem[field])
                    new_elem[field] = re.sub(r"✠\s+([,;.])", r"✠\1", new_elem[field])
            result.append(new_elem)

        return result

    def _get_pop_form_key(self) -> str:
        """Map the sheet's POP form designation to a YAML key.

        Handles:
          - "VI (w/ confession)" → form_VI (special case via rules)
          - "III" → form_III
          - "III (immigration)" → form_III_immigration (versioned form)
          - Advent forms detected by season

        Versioned forms use a suffix: form_III_immigration,
        form_V_hidden_springs, etc. The version name is derived from
        any parenthetical in the sheet that isn't "w/ confession".
        """
        form = (self.schedule.pop_form or "").strip()

        # Check for BCP Form VI w/ confession (handled in YAML under form_VI)
        if self.rules.pop_use_bcp_form_vi:
            return "form_VI"

        # Check for Advent forms
        if self.rules.is_advent:
            title_lower = self.schedule.title.lower()
            if "first" in title_lower or "1" in title_lower.replace(" ", ""):
                return "advent_I"
            elif "second" in title_lower or "2" in title_lower.replace(" ", ""):
                return "advent_II"
            elif "third" in title_lower or "3" in title_lower.replace(" ", ""):
                return "advent_III"
            elif "fourth" in title_lower or "4" in title_lower.replace(" ", ""):
                return "advent_IV"

        # Named (non-numbered) forms: "Easter", etc.
        named_map = {
            "easter": "easter",
        }
        form_lower = form.lower()
        if form_lower in named_map:
            return named_map[form_lower]

        # Standard forms: "I", "II", "III", "IV", "V", "VI"
        roman_map = {
            "I": "form_I", "II": "form_II", "III": "form_III",
            "IV": "form_IV", "V": "form_V", "VI": "form_VI",
            "1": "form_I", "2": "form_II", "3": "form_III",
            "4": "form_IV", "5": "form_V", "6": "form_VI",
        }

        # Extract base form and optional version from parenthetical
        paren_match = re.search(r'\(([^)]+)\)', form)
        version_suffix = ""
        if paren_match:
            paren_content = paren_match.group(1).strip().lower()
            # "w/ confession" and "with confession" are handled by rules
            if paren_content not in ("w/ confession", "with confession"):
                version_suffix = "_" + paren_content.replace(" ", "_")

        form_clean = form.split("(")[0].strip()
        base_key = roman_map.get(form_clean, "form_I")

        return base_key + version_suffix

    def _prepare_maundy_thursday_data(self) -> dict:
        """Prepare the Maundy-Thursday-specific data dict."""
        mt_texts = load_maundy_thursday()

        # Psalm 22 for the Stripping of the Altar
        psalm_22_text = []
        try:
            from bulletin.sources.psalms import get_psalm
            ps = get_psalm("Psalm 22")
            psalm_22_text = ps.to_lines()
        except Exception as e:
            self._warn(f"Could not look up Psalm 22: {e}",
                       category="psalm",
                       fix_hint="Stripping of the Altar text will be missing "
                                "the psalm; check psalms.yaml.")

        # Post-communion hymn uses the Recessional slot
        post_comm = self._lookup_slot("Recessional")

        # Foot washing songs — look up from music slots
        # The sheet may use slots like "Communion 1", "Communion 2", etc.
        # for foot washing music, or we may need custom slot names.
        # For now, foot washing songs are not populated from the sheet.
        foot_washing_songs = []

        return {
            "foot_washing": mt_texts.get("foot_washing", {}),
            "foot_washing_songs": foot_washing_songs,
            "anthem": mt_texts.get("anthem", {}),
            "post_communion_hymn": post_comm,
            "stripping": mt_texts.get("stripping", {}),
            "psalm_22_text": psalm_22_text,
            "watch": mt_texts.get("watch", {}),
        }

    def _prepare_good_friday_data(self) -> dict:
        """Prepare the Good-Friday-specific data dict."""
        gf_texts = load_good_friday()

        # Veneration hymns from music slots
        # First hymn: processional slot (cross procession hymn)
        # Second hymn: recessional slot (veneration/devotion hymn)
        veneration_hymns = []
        proc = self._lookup_slot("Processional")
        if proc:
            veneration_hymns.append(proc)
        recess = self._lookup_slot("Recessional")
        if recess:
            veneration_hymns.append(recess)

        # Passion Gospel in parts (John — same every year on Good Friday)
        passion_lines = []
        try:
            passion_data = load_passion_gospel("john")
            passion_lines = passion_data.get("lines", [])
        except Exception as e:
            self._block(
                f"Could not load Good Friday passion gospel: {e}",
                category="passion_gospel",
                fix_hint="Check passion_gospel_john.yaml exists and is "
                         "well-formed; the bulletin's Veneration section "
                         "will be missing the Passion narrative.")

        return {
            "entrance_rubric": gf_texts.get("entrance_rubric", ""),
            "acclamation": gf_texts.get("acclamation", {}),
            "solemn_collects": gf_texts.get("solemn_collects", {}),
            "veneration_hymns": veneration_hymns,
            "veneration": gf_texts.get("veneration", {}),
            "closing_prayer": gf_texts.get("closing_prayer", ""),
            "closing_rubric": gf_texts.get("closing_rubric", ""),
            "passion_gospel_lines": passion_lines,
        }

    def _prepare_palm_sunday_data(self) -> dict:
        """Prepare the Palm-Sunday-specific data dict."""
        palm_texts = load_palm_sunday()

        # Determine liturgical year from the schedule title or date
        year = self._get_liturgical_year()

        # Palm Gospel (triumphal entry) — varies by year
        liturgy = palm_texts["liturgy_of_the_palms"]
        palm_gospel_ref = liturgy["palm_gospel"].get(year, "Matthew 21:1-11")
        palm_gospel_text = self.scripture.get("palm_gospel")

        # Passion Gospel (read in parts)
        passion_config = palm_texts["passion_gospel"]
        passion_ref = passion_config["references"].get(year, "Matthew 27:11-54")

        # Load the passion gospel lines from the year-specific YAML
        try:
            passion_data = load_passion_gospel(year)
            passion_lines = passion_data.get("lines", [])
        except Exception as e:
            self._block(
                f"Could not load passion gospel for year {year}: {e}",
                category="passion_gospel",
                fix_hint=f"Check passion_gospel_{year}.yaml exists; "
                         "the Palm Sunday Passion narrative will be missing.")
            passion_lines = []

        return {
            "palm_texts": palm_texts,
            "palm_gospel_ref": palm_gospel_ref,
            "palm_gospel_text": palm_gospel_text,
            "passion_gospel_ref": passion_ref,
            "passion_gospel_lines": passion_lines,
        }

    def _get_liturgical_year(self) -> str:
        """Determine the liturgical year (A, B, or C) from the schedule.

        The RCL (Revised Common Lectionary) year cycle:
          A: years divisible by 3 with remainder 1 (2023, 2026, 2029, ...)
          B: years divisible by 3 with remainder 2 (2024, 2027, 2030, ...)
          C: years divisible by 3 with remainder 0 (2025, 2028, 2031, ...)

        Note: the liturgical year starts in Advent (late Nov/early Dec),
        but for simplicity we use the calendar year since Palm Sunday
        is always well into the liturgical year.
        """
        year_num = self.target_date.year
        remainder = year_num % 3
        if remainder == 1:
            return "A"
        elif remainder == 2:
            return "B"
        else:
            return "C"

    def get_reading_sheet_data(self) -> dict:
        """Return the data subset needed for reading sheet generation.

        Must be called after resolve_all() so POP version is resolved.
        """
        wog = self._prepare_word_of_god_data()
        return {
            "reading_1_ref": wog["reading_1_ref"],
            "reading_1_text": wog["reading_1_text"],
            "psalm_ref": wog["psalm_ref"],
            "psalm_text": wog["psalm_text"],
            "psalm_rubric": wog["psalm_rubric"],
            "pop_elements": wog["pop_elements"],
            "pop_concluding_rubric": wog["pop_concluding_rubric"],
        }

    def get_psalm_rubric_for_service(self, service_time: str) -> str:
        """Return the psalm rubric for a given service time.

        8am always returns unison; other services read from the sheet.
        """
        if service_time == "8 am":
            return "Read in unison."
        return self._get_psalm_rubric()
