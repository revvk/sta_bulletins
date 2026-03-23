"""
Template handling for pre-designed .docx cover pages.

Opens a Word template, replaces placeholder strings (e.g. {{DATE}}),
and returns a Document ready for additional content to be appended.

Placeholder replacement handles the common case where Word splits a
placeholder across multiple XML runs (due to spell-check or grammar
marks).  The algorithm concatenates all run text in a paragraph, does
the replacement on the joined string, puts the result into the first
run, and clears the rest.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

from bulletin.config import (
    FONT_HEADER_FOOTER, PAGE_WIDTH_INCHES, MARGIN_INCHES,
)

# Resolve templates directory relative to the project root.
# bulletin/document/templates.py  ->  ../../templates/
_TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "templates"

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def load_front_cover(
    date_str: str,
    service_time: str,
    liturgical_title: str,
    subtitle: str = " ",
    cover_template: str = None,
) -> Document:
    """Open the front-cover template and fill in dynamic fields.

    Args:
        date_str:          e.g. "March 1, 2026"
        service_time:      e.g. "9 am"
        liturgical_title:  e.g. "Second Sunday in Lent"
        subtitle:          Optional subtitle line; defaults to a single
                           space to preserve the template's spacing.
        cover_template:    Optional template filename (e.g. "palm_sunday_cover.docx").
                           Defaults to "front_cover.docx" if not provided.

    Returns:
        A python-docx Document with placeholders replaced.
    """
    filename = cover_template if cover_template else "front_cover.docx"
    template_path = _TEMPLATES_DIR / filename
    if not template_path.exists():
        raise FileNotFoundError(f"Front cover template not found: {template_path}")

    doc = Document(str(template_path))

    replacements = {
        "{{DATE}}": date_str,
        "{{SERVICE_TIME}}": service_time,
        "{{TITLE OF DAY}}": liturgical_title,
        "{{SUBTITLE}}": subtitle,
    }

    _replace_all_placeholders(doc, replacements)
    return doc


def append_back_cover(doc: Document):
    """Append the back-cover template as a new page section at the end.

    Inserts a section break (next-page) after the existing content, then
    copies all content and page setup from the back-cover template into
    the new section.  Images and hyperlinks are properly transferred so
    the resulting document is valid.
    """
    template_path = _TEMPLATES_DIR / "back_cover.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Back cover template not found: {template_path}")

    back_doc = Document(str(template_path))
    body = doc.element.body
    back_body = back_doc.element.body

    # --- Step 1: copy images & hyperlinks, build rId remapping ----------
    rid_map = _copy_related_parts(back_doc, doc)

    # --- Step 2: close the current section (next-page break) ------------
    main_sect_pr = body.find(qn("w:sectPr"))
    if main_sect_pr is not None:
        body.remove(main_sect_pr)

        # Set the break type to "nextPage" so the back cover starts on a
        # fresh page.
        sect_type = main_sect_pr.find(qn("w:type"))
        if sect_type is not None:
            sect_type.set(qn("w:val"), "nextPage")
        else:
            main_sect_pr.insert(
                0, parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
            )

        # Find the last paragraph and attach the sectPr to its pPr.
        paragraphs = list(body.iterchildren(qn("w:p")))
        if paragraphs:
            last_p = paragraphs[-1]
            pPr = last_p.find(qn("w:pPr"))
            if pPr is None:
                pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
                last_p.insert(0, pPr)
            pPr.append(main_sect_pr)

    # --- Step 3: append back-cover elements with remapped rIds ----------
    for element in list(back_body):
        elem_copy = deepcopy(element)
        if rid_map:
            _remap_rids(elem_copy, rid_map)
        body.append(elem_copy)


def setup_footers(doc: Document, date_str: str, service_time: str,
                  liturgical_title: str):
    """Add odd/even page footers to the bulletin body section.

    The document has two sections:
      Section 0 — front cover (first page) + bulletin content
      Section 1 — back cover

    Footers only appear on the bulletin content pages (not front or back
    cover).  The front cover is handled by ``different_first_page_header_footer``
    being True, so its first-page footer is left empty.

    Layout:
      Odd pages:  page number (left)   |   "March 8, 2026 | 9 am" (right)
      Even pages: "Third Sunday in Lent" (left)   |   page number (right)

    Page numbering starts at 1 for the first bulletin content page.
    """
    # Enable document-wide even/odd headers & footers.
    settings_elem = doc.settings.element
    # Avoid duplicating if already present
    if settings_elem.find(qn("w:evenAndOddHeaders")) is None:
        settings_elem.append(
            parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>')
        )

    section = doc.sections[0]

    # Set page number start = 0 so the front cover is page 0 and the
    # first bulletin content page is page 1.
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="0"/>')
        sect_pr.append(pg_num)
    else:
        pg_num.set(qn("w:start"), "0")

    # ---- Odd footer (default footer when even/odd is enabled) --------
    footer_odd = section.footer
    footer_odd.is_linked_to_previous = False
    _build_footer_paragraph(
        footer_odd.paragraphs[0], doc,
        left_field="page",
        right_text=f"{date_str} | {service_time}",
    )

    # ---- Even footer ---------------------------------------------------
    footer_even = section.even_page_footer
    footer_even.is_linked_to_previous = False
    _build_footer_paragraph(
        footer_even.paragraphs[0], doc,
        left_text=liturgical_title,
        right_field="page",
    )

    # ---- First-page footer (front cover): empty -------------------------
    footer_first = section.first_page_footer
    footer_first.is_linked_to_previous = False
    # Leave the first-page footer empty — no text or page number.

    # ---- Back cover section: suppress footers ----------------------------
    if len(doc.sections) > 1:
        back = doc.sections[-1]
        back.footer.is_linked_to_previous = False
        back.even_page_footer.is_linked_to_previous = False
        if back.different_first_page_header_footer:
            back.first_page_footer.is_linked_to_previous = False


def _build_footer_paragraph(paragraph, doc, *,
                            left_text: str = "",
                            left_field: str = "",
                            right_text: str = "",
                            right_field: str = ""):
    """Configure a footer paragraph with left and right-aligned content.

    Args:
        paragraph: The footer paragraph to populate.
        doc: The Document (for style access).
        left_text/left_field: Content for the left side ("page" = PAGE field).
        right_text/right_field: Content for the right side.
    """
    paragraph.style = doc.styles["Header & Footer"]

    # Add a right tab stop at the full text width (page width − margins)
    tab_pos = Inches(PAGE_WIDTH_INCHES - 2 * MARGIN_INCHES)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        tab_pos, WD_TAB_ALIGNMENT.RIGHT
    )

    # Left content
    if left_field == "page":
        _add_page_number_field(paragraph)
    elif left_text:
        paragraph.add_run(left_text)

    # Tab to right side
    paragraph.add_run("\t")

    # Right content
    if right_field == "page":
        _add_page_number_field(paragraph)
    elif right_text:
        paragraph.add_run(right_text)


def _add_page_number_field(paragraph):
    """Insert a PAGE field code into a paragraph for automatic page numbering."""
    fld_begin = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="begin"/>'
        f'</w:r>'
    )
    fld_instr = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:instrText xml:space="preserve"> PAGE </w:instrText>'
        f'</w:r>'
    )
    fld_sep = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="separate"/>'
        f'</w:r>'
    )
    fld_text = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:t>1</w:t>'
        f'</w:r>'
    )
    fld_end = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:fldChar w:fldCharType="end"/>'
        f'</w:r>'
    )

    p_elem = paragraph._element
    p_elem.append(fld_begin)
    p_elem.append(fld_instr)
    p_elem.append(fld_sep)
    p_elem.append(fld_text)
    p_elem.append(fld_end)


# ------------------------------------------------------------------
# Internals – document merging helpers
# ------------------------------------------------------------------

_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# Relationship types we need to transfer from one document to another.
_IMAGE_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
)
_HYPERLINK_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _copy_related_parts(source_doc: Document, target_doc: Document) -> dict:
    """Copy images and hyperlinks from *source_doc* into *target_doc*.

    Returns a ``{old_rId: new_rId}`` mapping so that XML elements copied
    from the source can have their relationship references remapped.
    """
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    rid_map: dict[str, str] = {}
    src_part = source_doc.part
    tgt_part = target_doc.part

    # Collect existing media part names in the target to avoid collisions.
    existing_names = {str(p.partname) for p in tgt_part.package.iter_parts()}

    for rId, rel in list(src_part.rels.items()):
        if rel.reltype == _IMAGE_RELTYPE:
            img = rel.target_part

            # If the partname already exists in the target, generate a
            # unique one (e.g. /word/media/image1_2.png).
            new_name = str(img.partname)
            if new_name in existing_names:
                base, ext = new_name.rsplit(".", 1)
                counter = 2
                while f"{base}_{counter}.{ext}" in existing_names:
                    counter += 1
                new_name = f"{base}_{counter}.{ext}"

            new_part = OpcPart(
                PackURI(new_name), img.content_type, img.blob, tgt_part.package
            )
            new_rId = tgt_part.relate_to(new_part, _IMAGE_RELTYPE)
            rid_map[rId] = new_rId
            existing_names.add(new_name)

        elif rel.reltype == _HYPERLINK_RELTYPE and rel.is_external:
            # External hyperlinks (e.g. mailto: links)
            new_rId = tgt_part.relate_to(
                rel.target_ref, _HYPERLINK_RELTYPE, is_external=True
            )
            rid_map[rId] = new_rId

    return rid_map


def _remap_rids(element, rid_map: dict):
    """Walk *element* and replace old rIds with new ones.

    Handles ``r:embed``, ``r:link``, and ``r:id`` attributes used by
    images, hyperlinks, and other relationship references in OpenXML.
    """
    attrs = [f"{{{_R_NS}}}embed", f"{{{_R_NS}}}link", f"{{{_R_NS}}}id"]
    for node in element.iter():
        for attr in attrs:
            old = node.get(attr)
            if old and old in rid_map:
                node.set(attr, rid_map[old])


# ------------------------------------------------------------------
# Internals – placeholder replacement
# ------------------------------------------------------------------

def _replace_all_placeholders(doc: Document, replacements: dict[str, str]):
    """Walk every <w:p> in the document (body + text boxes) and replace."""
    body = doc.element.body
    for para in body.iter(qn("w:p")):
        _replace_in_paragraph(para, replacements)


def _replace_in_paragraph(para_elem, replacements: dict[str, str]):
    """Replace placeholders in one paragraph, handling cross-run splits.

    Strategy: concatenate all run texts, do the replacement on the
    joined string, deposit the result into the first <w:t> element,
    and blank out the remaining <w:t> elements.  This preserves the
    first run's character formatting (font, colour, size, etc.).
    """
    runs = para_elem.findall(f"{{{_W_NS}}}r")
    if not runs:
        return

    # Collect text from each run's <w:t>
    t_elements = []
    for run in runs:
        t = run.find(f"{{{_W_NS}}}t")
        t_elements.append(t)

    full_text = "".join((t.text or "") if t is not None else "" for t in t_elements)

    # Check whether any placeholder appears in this paragraph
    changed = False
    for placeholder, replacement in replacements.items():
        if placeholder in full_text:
            full_text = full_text.replace(placeholder, replacement)
            changed = True

    if not changed:
        return

    # Deposit the replaced text into the first <w:t>, blank the rest.
    first_set = False
    for t in t_elements:
        if t is None:
            continue
        if not first_set:
            t.text = full_text
            t.set(qn("xml:space"), "preserve")
            first_set = True
        else:
            t.text = ""
