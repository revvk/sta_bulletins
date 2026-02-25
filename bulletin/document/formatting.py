"""
Helper functions for adding formatted content to a bulletin document.

These functions encapsulate the common patterns used throughout the bulletin:
  - Rubrics (italic directions like "Please stand.")
  - Celebrant/People dialogue pairs
  - Hymn headers with tune names and hymnal references
  - Scripture text with superscript verse numbers
  - Two-column lyrics layout using borderless tables
  - Cross symbols
  - Spacers
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.table import Table
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

from bulletin.config import CROSS_SYMBOL, FONT_BODY, FONT_LYRICS


def add_spacer(doc: Document):
    """Add a small vertical spacer paragraph."""
    doc.add_paragraph("", style="Spacer - Small")


def add_heading(doc: Document, text: str):
    """Add a major section heading (e.g., 'The Word of God', 'The Holy Communion')."""
    doc.add_paragraph(text, style="Heading")


def add_heading2(doc: Document, text: str):
    """Add a section header (e.g., 'Processional', 'Sermon', 'The Peace')."""
    doc.add_paragraph(text, style="Heading 2")


def add_rubric(doc: Document, text: str):
    """Add a rubric (italic instruction within the service)."""
    doc.add_paragraph(text, style="Body - Rubric")


def add_introductory_rubric(doc: Document, text: str):
    """Add an introductory rubric ('Please stand.', 'Be seated.', etc.)."""
    doc.add_paragraph(text, style="Body - Introductory Rubric")


def add_body(doc: Document, text: str):
    """Add a body text paragraph."""
    doc.add_paragraph(text, style="Body")


def add_body_with_bold_ending(doc: Document, text: str, bold_text: str):
    """Add a body paragraph where the last portion is bold (e.g., ending with 'Amen.')."""
    p = doc.add_paragraph(style="Body")
    p.add_run(text)
    run = p.add_run(bold_text)
    run.bold = True
    return p


def add_celebrant_line(doc: Document, label: str, text: str):
    """Add a Celebrant dialogue line.

    Args:
        label: Usually 'Celebrant' or 'The Deacon or a Priest'
        text: The words spoken
    """
    p = doc.add_paragraph(style="Body - Celebrant")
    label_run = p.add_run(label)
    label_run.font.size = Pt(12)
    p.add_run("\t" + text)
    return p


def add_people_line(doc: Document, label: str, text: str):
    """Add a People response line (bold).

    Args:
        label: Usually 'People'
        text: The response (already bold from style)
    """
    p = doc.add_paragraph(style="Body - People")
    p.add_run(label + "\t" + text)
    return p


def add_dialogue(doc: Document, celebrant_text: str, people_text: str,
                 celebrant_label: str = "Celebrant",
                 people_label: str = "People"):
    """Add a Celebrant/People dialogue exchange."""
    add_celebrant_line(doc, celebrant_label, celebrant_text)
    add_people_line(doc, people_label, people_text)


def add_cross_symbol(paragraph, font_name=FONT_BODY):
    """Add a bold-italic cross symbol (✠) as a run in the given paragraph."""
    run = paragraph.add_run(CROSS_SYMBOL)
    run.bold = True
    run.italic = True
    run.font.name = font_name
    return run


def add_hymn_header(doc: Document, title: str, tune_name: str = None,
                    hymnal_number: str = None, hymnal_name: str = None):
    """Add a hymn title line with optional tune name and hymnal reference.

    Examples:
        add_hymn_header(doc, "Everlasting God")
        add_hymn_header(doc, "Come, Holy Spirit", "Saint Agnes", "510", "Hymnal 1982")
    """
    p = doc.add_paragraph(style="Body")
    p.add_run(title)
    if tune_name:
        run = p.add_run("\t" + tune_name)
        run.italic = True
    if hymnal_number:
        num_text = f"\t#{hymnal_number}"
        if hymnal_name:
            num_text += " "
        p.add_run(num_text)
        if hymnal_name:
            run = p.add_run(f"({hymnal_name})")
            run.italic = True
    return p


def add_lyric_verse(doc: Document, lines: list[str]):
    """Add a verse (non-italic) of a hymn/song, one paragraph per line."""
    for line in lines:
        doc.add_paragraph(line, style="Body - Lyrics")


def add_lyric_chorus(doc: Document, lines: list[str]):
    """Add a chorus/refrain (italic) of a hymn/song, one paragraph per line."""
    for line in lines:
        p = doc.add_paragraph(style="Body - Lyrics")
        run = p.add_run(line)
        run.italic = True


def add_song(doc: Document, song_data: dict):
    """Add a complete song (header + all verses/choruses).

    Args:
        song_data: dict with keys:
            - title: str
            - tune_name: str | None
            - hymnal_number: str | None
            - hymnal_name: str | None
            - sections: list of dicts with:
                - type: "verse" | "chorus"
                - lines: list[str]
    """
    add_hymn_header(
        doc,
        song_data["title"],
        song_data.get("tune_name"),
        song_data.get("hymnal_number"),
        song_data.get("hymnal_name"),
    )
    for i, section in enumerate(song_data["sections"]):
        if i > 0:
            add_spacer(doc)
        if section["type"] == "chorus":
            add_lyric_chorus(doc, section["lines"])
        else:
            add_lyric_verse(doc, section["lines"])


def add_song_two_column(doc: Document, song_data: dict):
    """Add a song using a two-column borderless table to save vertical space.

    Distributes verse sections evenly across left and right columns.
    Chorus sections are repeated or placed as-is.
    """
    # Add the header normally
    add_hymn_header(
        doc,
        song_data["title"],
        song_data.get("tune_name"),
        song_data.get("hymnal_number"),
        song_data.get("hymnal_name"),
    )

    sections = song_data["sections"]
    if len(sections) < 2:
        # Too short for two columns -- just do single column
        for section in sections:
            if section["type"] == "chorus":
                add_lyric_chorus(doc, section["lines"])
            else:
                add_lyric_verse(doc, section["lines"])
        return

    # Create a borderless two-column table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    _remove_table_borders(table)

    # Set column widths (equal split of available space)
    available_width = Inches(PAGE_WIDTH_INCHES - 2 * MARGIN_INCHES)
    col_width = available_width // 2
    for cell in table.rows[0].cells:
        cell.width = col_width

    # Split sections into left and right columns
    mid = (len(sections) + 1) // 2
    left_sections = sections[:mid]
    right_sections = sections[mid:]

    _fill_lyric_cell(table.cell(0, 0), left_sections)
    _fill_lyric_cell(table.cell(0, 1), right_sections)


def _fill_lyric_cell(cell, sections: list[dict]):
    """Fill a table cell with song sections (verses/choruses)."""
    # Remove the default empty paragraph
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    for i, section in enumerate(sections):
        if i > 0:
            # Add spacer between sections
            sp = cell.add_paragraph("", style="Spacer - Small")

        for line_text in section["lines"]:
            p = cell.add_paragraph(style="Body - Lyrics")
            run = p.add_run(line_text)
            if section["type"] == "chorus":
                run.italic = True


def _remove_table_borders(table: Table):
    """Remove all borders from a table (makes it invisible)."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Scripture text with verse numbers
# ---------------------------------------------------------------------------

# Non-breaking space used between superscript verse number and text
_NBSP = "\u00A0"


def add_scripture_text(doc: Document, text: str, style: str = "Reading/Gospel Text",
                       indent: bool = False):
    """Add scripture text with superscript verse numbers.

    The text should have verse numbers inline (as from bible.oremus.org).
    Numbers at the start of sentences are detected and formatted as superscript.

    Formatting rules:
      - Verse numbers are superscript, followed by a non-breaking space.
      - Continuation paragraphs (indent=True, or 2nd+ \\n\\n-separated paragraphs)
        are indented with a tab.
      - Asterisk characters (oremus.org artifacts) are stripped.

    Args:
        text: The scripture text with inline verse numbers.
        style: The paragraph style to use.
        indent: If True, indent the first paragraph with a tab.
                Used when the caller iterates paragraphs and knows
                this is not the first paragraph of the reading.
    """
    # Remove asterisks (liturgical break markers from oremus.org)
    text = text.replace("*", "")

    paragraphs = text.strip().split("\n\n")  # Double newline = new paragraph

    para_count = 0
    p = None
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        p = doc.add_paragraph(style=style)

        # Indent with a tab if:
        # - caller said this is a continuation paragraph (indent=True), OR
        # - this is the 2nd+ paragraph within this text block
        if indent or para_count > 0:
            p.add_run("\t")

        _add_verse_numbered_text(p, para_text)
        para_count += 1

    return p


def _add_verse_numbered_text(paragraph, text: str):
    """Add text to a paragraph, converting verse numbers to superscript.

    Verse numbers appear as bare integers followed by a space and text:
      "2 And suddenly from heaven..." or "17 'In the last days..."
    The space between the verse number and text is consumed by the regex;
    a non-breaking space is added as part of the superscript run instead.
    """
    segments = _split_verse_numbers(text)

    for verse_num, segment_text in segments:
        if verse_num:
            run = paragraph.add_run(verse_num + _NBSP)
            run.font.superscript = True
            run.font.size = Pt(9)
        if segment_text:
            paragraph.add_run(segment_text)


def _split_verse_numbers(text: str) -> list[tuple[str | None, str]]:
    """Split scripture text into (verse_number, text) segments.

    Returns a list of tuples where verse_number may be None for text
    that doesn't start with a verse number.

    Input format (from bible.oremus.org):
      "4 From Mount Hor they set out... the way. 5 The people spoke..."
    Verse numbers are 1-3 digits, preceded by start-of-string or whitespace,
    followed by a space then an uppercase letter, opening quote, or bracket.
    """
    # The pattern matches: (start or after whitespace) + digits + space + (uppercase/quote/bracket)
    # The space after the digits is consumed so it doesn't appear in the text segment.
    pattern = re.compile(
        r'(?:^|(?<=\s))(\d{1,3})\s(?=[A-Z\u2018\u201C\'\"(\[])'
    )
    segments = []
    last_end = 0

    for match in pattern.finditer(text):
        # Text from last_end to match.start() belongs to previous segment
        before_text = text[last_end:match.start()]

        if before_text:
            if segments:
                prev_num, prev_text = segments[-1]
                segments[-1] = (prev_num, prev_text + before_text)
            else:
                segments.append((None, before_text))

        # Start new segment with this verse number
        num = match.group(1)
        last_end = match.end()  # Position right before the uppercase letter
        segments.append((num, ""))

    # Remaining text after last verse number
    if last_end < len(text):
        remaining = text[last_end:]
        if segments:
            prev_num, prev_text = segments[-1]
            segments[-1] = (prev_num, prev_text + remaining)
        else:
            segments.append((None, remaining))

    if not segments:
        segments.append((None, text))

    return segments


# Import here to avoid circular; these are only needed at module level
from bulletin.config import PAGE_WIDTH_INCHES, MARGIN_INCHES
