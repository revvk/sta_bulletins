"""
Reading sheet builder for lay readers.

Generates a large-print document containing the First Reading (with
proper book introduction), the Psalm (with reader instructions for
the mode of reading), and the Prayers of the People.

The same formatting helper functions used by the bulletin are reused
here — the reading sheet styles (configured via
configure_reading_sheet_document) use the same style names but with
larger fonts and red rubric colors.
"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from bulletin.config import FONT_BODY, FONT_BODY_BOLD
from bulletin.data.loader import (
    load_reading_introductions,
    load_psalm_reader_instructions,
    extract_book_name,
)
from bulletin.document.styles import configure_reading_sheet_document
from bulletin.document.formatting import (
    add_spacer, add_rubric,
    add_body, add_scripture_text,
    add_celebrant_line, add_people_line,
    _add_text_runs,
)
from bulletin.document.sections.word_of_god import add_pop
from bulletin.sources.psalms import parse_psalm_reference


# Maps the bulletin's psalm rubric strings to YAML mode keys.
_RUBRIC_TO_MODE = {
    "Read in unison.": "unison",
    "Read responsively by half verse.": "responsive_half_verse",
    "Read responsively by whole verse.": "responsive_whole_verse",
    "Read antiphonally.": "antiphonal",
    "Read alternating between men and women.": "men_and_women",
}


def build_reading_sheet(data: dict, psalm_rubric: str) -> Document:
    """Build a reading sheet document.

    Args:
        data: Dict from BulletinBuilder.get_reading_sheet_data() with keys:
            reading_1_ref, reading_1_text, psalm_ref, psalm_text,
            psalm_rubric, pop_elements, pop_concluding_rubric.
        psalm_rubric: The psalm rubric for this reading sheet variant
            (e.g., "Read in unison." or "Read antiphonally.").

    Returns:
        A Document ready to save.
    """
    doc = Document()
    configure_reading_sheet_document(doc)

    # --- Section marker: readings ---
    _add_section_marker(doc, "Readings begin on next page.")

    # --- First Reading ---
    _add_reading_with_intro(
        doc,
        data["reading_1_ref"],
        data["reading_1_text"],
    )

    # --- Psalm ---
    add_spacer(doc)
    _add_psalm_for_reader(
        doc,
        data["psalm_ref"],
        data["psalm_text"],
        psalm_rubric,
    )

    # --- Section marker: prayers ---
    _add_section_marker(doc, "Prayers begin on next page.")

    # --- Prayers of the People ---
    _add_pop_section(doc, data)

    return doc


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _strip_verse_numbers(text: str) -> str:
    """Remove verse number markers (\\x01N\\x01) from scripture text."""
    return re.sub(r'\x01\d{1,3}\x01\s?', '', text)


def _add_scripture_text_no_verses(doc: Document, text: str,
                                   style: str = "Reading/Gospel Text",
                                   indent: bool = False):
    """Add scripture text with verse numbers stripped.

    Works like add_scripture_text but removes verse numbers entirely
    instead of rendering them as superscript.
    """
    text = _strip_verse_numbers(text)
    add_scripture_text(doc, text, style=style, indent=indent)


def _add_section_marker(doc: Document, text: str):
    """Add a section marker heading followed by a page break."""
    add_spacer(doc)
    p = doc.add_paragraph(style="Heading")
    p.add_run(text)

    # Insert a page break after the marker
    p2 = doc.add_paragraph()
    run = p2.add_run()
    br = parse_xml(
        '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:type="page"/>'
    )
    run._element.append(br)


def _add_reading_with_intro(doc: Document, reference: str, reading):
    """Add a scripture reading with introduction and closing response."""
    # Rubric before reading
    add_rubric(doc, "Before the reading, the Reader will say")
    add_spacer(doc)

    # Book introduction — look up the proper form
    intros = load_reading_introductions()
    book_name = extract_book_name(reference)
    intro_phrase = intros.get(book_name, book_name)
    add_body(doc, f"A Reading from {intro_phrase}:")
    add_spacer(doc)

    # Full reading text (verse numbers stripped for reading sheets).
    # Same three-tier fallback the bulletin renderer uses
    # (word_of_god._add_reading_text):
    #   1. `segments` for interleaved prose/poetry — important for
    #      readings like Acts 17 where a quotation ("For we too are
    #      his offspring") is embedded mid-paragraph; we render
    #      each segment in document order so it lands in the right
    #      place rather than being appended at the end.
    #   2. `paragraphs` + (optional) `poetry_lines` — legacy
    #      structure for pure-prose or all-poetry readings.
    #   3. plain string fallback.
    if hasattr(reading, "segments") and reading.segments:
        for seg in reading.segments:
            if seg["type"] == "prose":
                _add_scripture_text_no_verses(doc, seg["text"])
            elif seg["type"] == "poetry":
                for line in seg["lines"]:
                    if isinstance(line, dict):
                        indent_level = line.get("indent", 0)
                        text = line["text"]
                    else:
                        indent_level = 0
                        text = line
                    if indent_level == 0:
                        style = "Reading (Poetry)"
                    elif indent_level == 1:
                        style = "Reading (Poetry Indent 1)"
                    else:
                        style = "Reading (Poetry Indent 2)"
                    _add_scripture_text_no_verses(doc, text, style=style)
    elif hasattr(reading, "paragraphs"):
        for i, para in enumerate(reading.paragraphs):
            _add_scripture_text_no_verses(doc, para, indent=(i > 0))
        if reading.has_poetry:
            for line in reading.poetry_lines:
                _add_scripture_text_no_verses(doc, line, style="Reading (Poetry)")
    else:
        _add_scripture_text_no_verses(doc, str(reading))

    # Closing response
    add_spacer(doc)
    add_rubric(doc, "After the reading, the Reader will say")
    add_spacer(doc)
    add_celebrant_line(doc, "", "The Word of the Lord.")
    add_people_line(doc, "People", "Thanks be to God.")


def _add_psalm_for_reader(doc: Document, psalm_ref: str,
                          psalm_text: list[str], rubric: str):
    """Add the psalm with reader instructions and mode-appropriate bold pattern."""
    # Determine the mode
    mode_key = _RUBRIC_TO_MODE.get(rubric, "unison")
    instructions = load_psalm_reader_instructions()
    mode_data = instructions.get(mode_key, instructions.get("unison", {}))

    # Extract psalm number for instruction text
    try:
        psalm_num, _ = parse_psalm_reference(psalm_ref)
    except ValueError:
        psalm_num = ""

    # Add reader instruction (no heading — just the introductory instruction)
    instruction_text = mode_data.get("instruction", "")
    if "{psalm_number}" in instruction_text:
        instruction_text = instruction_text.replace(
            "{psalm_number}", str(psalm_num)
        )
    add_body(doc, instruction_text)

    # Antiphonal has additional hand-gesture instructions
    if mode_key == "antiphonal":
        add_spacer(doc)
        left_setup = mode_data.get("left_side_setup", "")
        left_text = mode_data.get("left_side_text", "")
        right_setup = mode_data.get("right_side_setup", "")
        right_text = mode_data.get("right_side_text", "")

        if left_setup:
            add_rubric(doc, left_setup)
        if left_text:
            add_body(doc, left_text)
        if right_setup:
            add_rubric(doc, right_setup)
        if right_text:
            add_body(doc, right_text)

    add_spacer(doc)

    # Determine bold pattern for the psalm:
    #   - Unison: all bold (everyone reads together)
    #   - Responsive/antiphonal/alternating by whole verse: alternating bold
    #   - Responsive by half-verse: alternating bold by half-verse
    all_bold = mode_key == "unison"
    alternating = mode_key in (
        "responsive_whole_verse", "antiphonal", "men_and_women"
    )
    half_verse = mode_key == "responsive_half_verse"

    if isinstance(psalm_text, list):
        for verse_idx, verse_text in enumerate(psalm_text):
            if half_verse:
                # Bold the second half of each verse
                _add_psalm_verse_half_bold(doc, verse_text)
            elif all_bold:
                _add_psalm_verse(doc, verse_text, bold=True)
            else:
                bold_verse = alternating and (verse_idx % 2 == 1)
                _add_psalm_verse(doc, verse_text, bold_verse)


def _add_psalm_verse(doc: Document, verse_text: str, bold: bool):
    """Add a single psalm verse, optionally all bold.

    Uses _add_text_runs to properly render LORD as small-caps Lord.
    Handles \\v (vertical tab) as a first-half continuation line that
    starts a new paragraph at the left margin (no hanging indent).
    """
    from docx.shared import Pt

    p = doc.add_paragraph(style="Psalm")
    if "\n" in verse_text:
        sub_lines = verse_text.split("\n")
        for i, sub in enumerate(sub_lines):
            if sub.startswith("\v"):
                # First-half continuation: new paragraph at left margin
                p.paragraph_format.space_after = Pt(0)
                p = doc.add_paragraph(style="Psalm")
                p.paragraph_format.space_before = Pt(0)
                _add_text_runs(p, sub[1:], bold=bold)
            elif i > 0:
                run = p.add_run()
                run.add_break()
                _add_text_runs(p, sub, bold=bold)
            else:
                _add_text_runs(p, sub, bold=bold)
    else:
        _add_text_runs(p, verse_text, bold=bold)


def _add_psalm_verse_half_bold(doc: Document, verse_text: str):
    """Add a psalm verse with the second half bold (responsive by half-verse).

    Uses _add_text_runs to properly render LORD as small-caps Lord.
    Handles \\v (vertical tab) as a first-half continuation line.
    """
    from docx.shared import Pt

    p = doc.add_paragraph(style="Psalm")
    if "\n" in verse_text:
        sub_lines = verse_text.split("\n")
        for i, sub in enumerate(sub_lines):
            if sub.startswith("\v"):
                p.paragraph_format.space_after = Pt(0)
                p = doc.add_paragraph(style="Psalm")
                p.paragraph_format.space_before = Pt(0)
                _add_text_runs(p, sub[1:], bold=False)
            elif i > 0:
                run = p.add_run()
                run.add_break()
                # Tab-indented lines are the second half → bold
                is_bold = sub.startswith("\t")
                _add_text_runs(p, sub, bold=is_bold)
            else:
                _add_text_runs(p, sub, bold=False)
    else:
        _add_text_runs(p, verse_text)


def _add_pop_section(doc: Document, data: dict):
    """Add the Prayers of the People section."""
    pop_elements = data.get("pop_elements", [])
    concluding_rubric = data.get(
        "pop_concluding_rubric",
        "The Celebrant concludes with a suitable Collect.",
    )

    add_pop(doc, pop_elements)
    # Empty/null `pop_concluding_rubric` suppresses the rubric —
    # used by forms with a built-in concluding collect.
    if concluding_rubric:
        add_spacer(doc)
        add_rubric(doc, concluding_rubric)
