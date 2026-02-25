#!/usr/bin/env python3
"""
Comprehensive DOCX structure analyzer for church song lyrics documents.
Dumps full paragraph-level and run-level formatting details to understand
how songs, verses, choruses, titles, and metadata are encoded.
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter, defaultdict
import re


def emu_to_pt(emu):
    """Convert EMU to points."""
    if emu is None:
        return None
    return round(emu / 12700, 2)


def emu_to_inches(emu):
    """Convert EMU to inches."""
    if emu is None:
        return None
    return round(emu / 914400, 3)


def get_alignment_name(alignment):
    """Get human-readable alignment name."""
    if alignment is None:
        return "None (inherited)"
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
    }
    return mapping.get(alignment, str(alignment))


def has_page_break_before(paragraph):
    """Check if paragraph has a page break before it."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        pageBreakBefore = pPr.find(qn('w:pageBreakBefore'))
        if pageBreakBefore is not None:
            val = pageBreakBefore.get(qn('w:val'))
            return val is None or val in ('1', 'true')
    return False


def has_run_page_break(paragraph):
    """Check if any run in the paragraph contains a page break character."""
    for run in paragraph.runs:
        for child in run._element:
            if child.tag == qn('w:br'):
                br_type = child.get(qn('w:type'))
                if br_type == 'page':
                    return True
    return False


def has_section_break(paragraph):
    """Check if paragraph ends with a section break."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            return True
    return False


def get_font_size(run):
    """Get font size from run properties."""
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        sz = rPr.find(qn('w:sz'))
        if sz is not None:
            val = sz.get(qn('w:val'))
            if val:
                return int(val) / 2  # half-points to points
    return None


def get_font_name(run):
    """Get font name from run properties."""
    if run.font and run.font.name:
        return run.font.name
    return None


def get_underline(run):
    """Get underline status."""
    if run.font and run.font.underline:
        return True
    return False


def get_font_color(run):
    """Get font color if set."""
    if run.font and run.font.color and run.font.color.rgb:
        return str(run.font.color.rgb)
    return None


def get_spacing(paragraph):
    """Get paragraph spacing details."""
    fmt = paragraph.paragraph_format
    result = {}
    if fmt.space_before is not None:
        result['before'] = emu_to_pt(fmt.space_before)
    if fmt.space_after is not None:
        result['after'] = emu_to_pt(fmt.space_after)
    if fmt.line_spacing is not None:
        result['line'] = fmt.line_spacing
    return result if result else None


def detect_special_chars(text):
    """Detect any non-standard ASCII characters."""
    specials = []
    for i, ch in enumerate(text):
        if ord(ch) > 127:
            specials.append(f"  U+{ord(ch):04X} '{ch}' at pos {i}")
    return specials


def analyze_document(filepath):
    """Full structural analysis of a DOCX file."""
    print("=" * 100)
    print(f"ANALYZING: {filepath}")
    print("=" * 100)

    doc = Document(filepath)

    # -- Document-level info --
    print("\n--- DOCUMENT PROPERTIES ---")
    props = doc.core_properties
    print(f"  Title:    {props.title}")
    print(f"  Author:   {props.author}")
    print(f"  Created:  {props.created}")
    print(f"  Modified: {props.modified}")

    # -- Styles used --
    style_counter = Counter()
    font_size_counter = Counter()
    bold_patterns = []
    italic_patterns = []
    all_paragraphs_data = []

    # -- Section info --
    print(f"\n--- SECTIONS ({len(doc.sections)}) ---")
    for i, section in enumerate(doc.sections):
        print(f"  Section {i}:")
        print(f"    Page width:  {emu_to_inches(section.page_width)} in")
        print(f"    Page height: {emu_to_inches(section.page_height)} in")
        print(f"    Left margin: {emu_to_inches(section.left_margin)} in")
        print(f"    Right margin:{emu_to_inches(section.right_margin)} in")
        print(f"    Top margin:  {emu_to_inches(section.top_margin)} in")
        print(f"    Bot margin:  {emu_to_inches(section.bottom_margin)} in")
        print(f"    Columns:     {section._sectPr.findall(qn('w:cols'))}")

    # -- Paragraph-by-paragraph dump --
    print(f"\n--- ALL PARAGRAPHS ({len(doc.paragraphs)} total) ---\n")

    page_break_indices = []
    blank_line_indices = []
    song_boundaries = []
    prev_was_blank = False

    for idx, para in enumerate(doc.paragraphs):
        text = para.text
        style_name = para.style.name if para.style else "None"
        style_counter[style_name] += 1

        # Formatting details
        alignment = get_alignment_name(para.alignment)
        indent_left = emu_to_inches(para.paragraph_format.left_indent) if para.paragraph_format.left_indent else None
        indent_right = emu_to_inches(para.paragraph_format.right_indent) if para.paragraph_format.right_indent else None
        indent_first = emu_to_inches(para.paragraph_format.first_line_indent) if para.paragraph_format.first_line_indent else None
        spacing = get_spacing(para)

        pb_before = has_page_break_before(para)
        pb_in_run = has_run_page_break(para)
        sect_break = has_section_break(para)

        if pb_before or pb_in_run:
            page_break_indices.append(idx)

        if text.strip() == "":
            blank_line_indices.append(idx)

        # Print paragraph header
        pb_marker = ""
        if pb_before:
            pb_marker += " [PAGE-BREAK-BEFORE]"
        if pb_in_run:
            pb_marker += " [PAGE-BREAK-IN-RUN]"
        if sect_break:
            pb_marker += " [SECTION-BREAK]"

        print(f"P[{idx:3d}] Style='{style_name}' | Align={alignment} | "
              f"Indent(L={indent_left}, R={indent_right}, 1st={indent_first})"
              f"{pb_marker}")

        if spacing:
            print(f"        Spacing: {spacing}")

        # Run-level details
        if para.runs:
            for r_idx, run in enumerate(para.runs):
                bold = run.bold
                italic = run.italic
                font_sz = get_font_size(run)
                font_nm = get_font_name(run)
                underline = get_underline(run)
                color = get_font_color(run)

                if font_sz:
                    font_size_counter[font_sz] += 1

                fmt_flags = []
                if bold:
                    fmt_flags.append("BOLD")
                if italic:
                    fmt_flags.append("ITALIC")
                if underline:
                    fmt_flags.append("UNDERLINE")
                if color:
                    fmt_flags.append(f"COLOR={color}")
                if font_sz:
                    fmt_flags.append(f"SIZE={font_sz}pt")
                if font_nm:
                    fmt_flags.append(f"FONT={font_nm}")

                fmt_str = " [" + ", ".join(fmt_flags) + "]" if fmt_flags else ""

                run_text = run.text
                # Show special whitespace explicitly
                display_text = run_text.replace('\t', '<TAB>').replace('\n', '<NL>')

                print(f"   R[{r_idx}]{fmt_str}: '{display_text}'")

                # Check for special characters
                specials = detect_special_chars(run_text)
                if specials:
                    for s in specials:
                        print(f"       SPECIAL CHAR: {s}")
        else:
            if text.strip():
                print(f"   (no runs, but text exists): '{text}'")
            else:
                print(f"   (empty paragraph - no runs, no text)")

        # Full text as a convenience line
        if text.strip():
            print(f"   FULL TEXT: '{text}'")

        print()  # blank line between paragraphs

        # Collect data for analysis
        all_paragraphs_data.append({
            'index': idx,
            'style': style_name,
            'text': text,
            'alignment': alignment,
            'indent_left': indent_left,
            'page_break_before': pb_before,
            'page_break_in_run': pb_in_run,
            'runs': [(r.text, r.bold, r.italic, get_font_size(r), get_font_name(r)) for r in para.runs],
        })

    # -- Summary Analysis --
    print("\n" + "=" * 100)
    print("STRUCTURAL ANALYSIS SUMMARY")
    print("=" * 100)

    print(f"\n--- Styles Used ---")
    for style, count in style_counter.most_common():
        print(f"  '{style}': {count} paragraphs")

    print(f"\n--- Font Sizes Used ---")
    for size, count in font_size_counter.most_common():
        print(f"  {size}pt: {count} runs")

    print(f"\n--- Page Breaks at paragraph indices ---")
    print(f"  {page_break_indices}")
    print(f"  Total page breaks: {len(page_break_indices)}")

    print(f"\n--- Blank Lines at paragraph indices ---")
    print(f"  Total blank paragraphs: {len(blank_line_indices)}")
    if blank_line_indices:
        print(f"  Indices: {blank_line_indices}")

    # Try to identify song boundaries
    print(f"\n--- Song Boundary Detection ---")
    songs = []
    current_song_start = 0
    current_song_title = None

    for i, pdata in enumerate(all_paragraphs_data):
        # Heuristic: song boundary at page breaks
        if pdata['page_break_before'] or pdata['page_break_in_run']:
            if current_song_title or i > current_song_start:
                songs.append({
                    'start': current_song_start,
                    'end': i - 1,
                    'title': current_song_title or '(untitled)',
                })
            current_song_start = i
            current_song_title = None

        # Heuristic: title detection - look for bold, larger font, or heading styles
        if pdata['text'].strip():
            runs = pdata['runs']
            all_bold = all(r[1] for r in runs if r[0].strip()) if runs else False
            has_large_font = any(r[3] and r[3] >= 14 for r in runs) if runs else False
            is_heading = 'Heading' in pdata['style'] or 'Title' in pdata['style']

            if (all_bold or has_large_font or is_heading) and current_song_title is None:
                # Check if this looks like a title (short, no verse numbering)
                txt = pdata['text'].strip()
                if len(txt) < 80 and not re.match(r'^\d+\.', txt):
                    current_song_title = txt

    # Don't forget the last song
    if current_song_start < len(all_paragraphs_data):
        songs.append({
            'start': current_song_start,
            'end': len(all_paragraphs_data) - 1,
            'title': current_song_title or '(untitled)',
        })

    print(f"  Detected {len(songs)} song(s):")
    for s in songs:
        print(f"    [{s['start']}-{s['end']}] {s['title']}")

    # Verse/chorus pattern detection
    print(f"\n--- Verse/Chorus Pattern Detection ---")
    verse_indicators = []
    chorus_indicators = []
    refrain_indicators = []
    for i, pdata in enumerate(all_paragraphs_data):
        txt = pdata['text'].strip().lower()
        if re.match(r'^(verse|v\.?\s*\d|stanza|\d+\.)', txt):
            verse_indicators.append((i, pdata['text'].strip()))
        if re.match(r'^(chorus|refrain|response)', txt, re.IGNORECASE):
            chorus_indicators.append((i, pdata['text'].strip()))
        # Check for italic runs (often used for chorus/refrain)
        has_italic = any(r[2] for r in pdata['runs'] if r[0].strip())
        if has_italic and pdata['text'].strip():
            refrain_indicators.append((i, pdata['text'].strip()[:60]))

    print(f"  Verse-like labels found: {len(verse_indicators)}")
    for v in verse_indicators:
        print(f"    P[{v[0]}]: '{v[1]}'")

    print(f"  Chorus/refrain labels found: {len(chorus_indicators)}")
    for c in chorus_indicators:
        print(f"    P[{c[0]}]: '{c[1]}'")

    print(f"  Italic paragraphs (possible chorus/refrain): {len(refrain_indicators)}")
    for r in refrain_indicators[:20]:
        print(f"    P[{r[0]}]: '{r[1]}'")
    if len(refrain_indicators) > 20:
        print(f"    ... and {len(refrain_indicators) - 20} more")

    # Metadata pattern detection
    print(f"\n--- Metadata Pattern Detection ---")
    metadata_patterns = []
    for i, pdata in enumerate(all_paragraphs_data):
        txt = pdata['text'].strip()
        # Look for tune references, hymnal numbers, author credits
        if re.search(r'(tune|hymn|#\d|text:|music:|words:|arr\.|adapt\.)', txt, re.IGNORECASE):
            metadata_patterns.append((i, txt))
        # Small font metadata
        small_font = any(r[3] and r[3] <= 9 for r in pdata['runs'] if r[0].strip())
        if small_font and txt:
            metadata_patterns.append((i, f"[SMALL FONT] {txt}"))

    print(f"  Metadata-like paragraphs: {len(metadata_patterns)}")
    for m in metadata_patterns:
        print(f"    P[{m[0]}]: '{m[1]}'")

    print("\n")
    return {
        'total_paragraphs': len(all_paragraphs_data),
        'songs': songs,
        'styles': dict(style_counter),
        'page_breaks': page_break_indices,
        'blank_lines': blank_line_indices,
        'verse_indicators': verse_indicators,
        'chorus_indicators': chorus_indicators,
        'italic_paragraphs': refrain_indicators,
        'metadata': metadata_patterns,
    }


def main():
    files = [
        "/Users/andrewvankirk/Scratch/Development/Bulletin/Bulletin Formatted Song Lyrics - 9 am.docx",
        "/Users/andrewvankirk/Scratch/Development/Bulletin/Bulletin Formatted Song Lyrics - 11 am.docx",
    ]

    results = {}
    for f in files:
        results[f] = analyze_document(f)

    # -- Cross-file comparison --
    print("=" * 100)
    print("CROSS-FILE COMPARISON & RECOMMENDATIONS")
    print("=" * 100)

    for f, r in results.items():
        short_name = f.split("/")[-1]
        print(f"\n  {short_name}:")
        print(f"    Total paragraphs: {r['total_paragraphs']}")
        print(f"    Songs detected:   {len(r['songs'])}")
        print(f"    Page breaks:      {len(r['page_breaks'])}")
        print(f"    Blank paragraphs: {len(r['blank_lines'])}")
        print(f"    Styles used:      {r['styles']}")

    print("\n--- RECOMMENDED DATA STRUCTURE ---")
    print("""
Based on analysis, a suggested data model for the bulletin generator:

    Song:
        title: str              # Song title (detected from bold/large/heading text)
        subtitle: str | None    # Tune name or alternate title
        metadata: dict          # Author, hymnal ref, tune, text source, etc.
        sections: list[Section] # Ordered list of verses/choruses

    Section:
        type: str               # 'verse', 'chorus', 'refrain', 'bridge', 'intro'
        number: int | None      # Verse number if applicable
        lines: list[str]        # Lines of text
        formatting: str         # 'normal', 'italic', 'bold', etc.
        indent_level: int       # 0 = no indent, 1 = first level, etc.

    BulletinSongList:
        service_time: str       # '9am' or '11am'
        date: str               # Service date
        songs: list[Song]       # Ordered songs for this service
""")


if __name__ == "__main__":
    main()
