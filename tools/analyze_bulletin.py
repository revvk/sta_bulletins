#!/usr/bin/env python3
"""
Comprehensive analysis of a Word (.docx) bulletin document.
Extracts all formatting details needed to recreate it with python-docx.

This document uses many sections (16 sections with section breaks serving as page breaks).
The paragraphs are spread across sections. We need to look at the raw XML body to
get ALL paragraphs, not just doc.paragraphs (which only gives paragraphs NOT inside sections).
"""

from docx import Document
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from lxml import etree

DOC_PATH = "/Users/andrewvankirk/Scratch/Development/Bulletin/2025-06-08 - Pentecost C - 9 am (HEII-A) - Bulletin.docx"

doc = Document(DOC_PATH)

def emu_to_inches(emu):
    if emu is None:
        return None
    return round(emu / 914400, 4)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 2)

def twips_to_pt(twips):
    if twips is None:
        return None
    return round(int(twips) / 20, 2)

def twips_to_inches(twips):
    if twips is None:
        return None
    return round(int(twips) / 1440, 4)

def alignment_name(align):
    if align is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "DISTRIBUTE",
    }
    return mapping.get(align, str(align))

def line_spacing_rule_name(rule):
    if rule is None:
        return None
    mapping = {
        WD_LINE_SPACING.SINGLE: "SINGLE",
        WD_LINE_SPACING.ONE_POINT_FIVE: "ONE_POINT_FIVE",
        WD_LINE_SPACING.DOUBLE: "DOUBLE",
        WD_LINE_SPACING.AT_LEAST: "AT_LEAST",
        WD_LINE_SPACING.EXACTLY: "EXACTLY",
        WD_LINE_SPACING.MULTIPLE: "MULTIPLE",
    }
    return mapping.get(rule, str(rule))

def get_font_color(font):
    try:
        if font.color and font.color.rgb:
            return str(font.color.rgb)
    except:
        pass
    return None

print("=" * 100)
print("BULLETIN DOCUMENT ANALYSIS")
print("=" * 100)
print(f"File: {DOC_PATH}")
print()

# ============================================================
# 1. DOCUMENT-LEVEL SETTINGS
# ============================================================
print("=" * 100)
print("1. DOCUMENT-LEVEL SETTINGS")
print("=" * 100)

print(f"\nNumber of sections: {len(doc.sections)}")
print(f"Number of doc.paragraphs: {len(doc.paragraphs)}")
print(f"Number of tables: {len(doc.tables)}")

# Count ALL paragraphs in the body XML
body = doc.element.body
all_para_elements = body.findall(qn('w:p'))
print(f"Number of ALL <w:p> in body: {len(all_para_elements)}")

for i, section in enumerate(doc.sections):
    print(f"\n--- Section {i} ---")
    print(f"  Page width:      {emu_to_inches(section.page_width)} in")
    print(f"  Page height:     {emu_to_inches(section.page_height)} in")
    print(f"  Left margin:     {emu_to_inches(section.left_margin)} in")
    print(f"  Right margin:    {emu_to_inches(section.right_margin)} in")
    print(f"  Top margin:      {emu_to_inches(section.top_margin)} in")
    print(f"  Bottom margin:   {emu_to_inches(section.bottom_margin)} in")
    print(f"  Header distance: {emu_to_inches(section.header_distance)} in")
    print(f"  Footer distance: {emu_to_inches(section.footer_distance)} in")
    print(f"  Orientation:     {section.orientation}")
    print(f"  Start type:      {section.start_type}")
    print(f"  Different first page: {section.different_first_page_header_footer}")

    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is not None:
        num_cols = cols.get(qn('w:num'))
        col_space = cols.get(qn('w:space'))
        equal_width = cols.get(qn('w:equalWidth'))
        print(f"  Columns: num={num_cols}, space={col_space} ({twips_to_inches(col_space) if col_space else None} in), equalWidth={equal_width}")
        col_elems = cols.findall(qn('w:col'))
        for j, col_elem in enumerate(col_elems):
            col_w = col_elem.get(qn('w:w'))
            col_s = col_elem.get(qn('w:space'))
            print(f"    Column {j}: width={col_w} ({twips_to_inches(col_w) if col_w else None} in), space={col_s} ({twips_to_inches(col_s) if col_s else None} in)")
    else:
        print("  Columns: 1 (default)")

print()

# ============================================================
# 2. DOCUMENT STYLES
# ============================================================
print("=" * 100)
print("2. ALL NAMED STYLES IN DOCUMENT")
print("=" * 100)

for style in doc.styles:
    # Skip numbering styles
    style_class_name = type(style).__name__
    if style_class_name == '_NumberingStyle':
        print(f"\n  Style: '{style.name}' (type=NUMBERING, skipping details)")
        continue

    if style.type is not None:
        style_type_name = {1: 'PARAGRAPH', 2: 'CHARACTER', 3: 'TABLE', 4: 'LIST'}.get(style.type, str(style.type))
    else:
        style_type_name = 'None'

    base_style_name = None
    try:
        base_style_name = style.base_style.name if style.base_style else None
    except:
        pass

    print(f"\n  Style: '{style.name}' (style_id='{style.style_id}', type={style_type_name}, builtin={style.builtin})")
    print(f"    Base style: {base_style_name}")

    if hasattr(style, 'font') and style.font:
        f = style.font
        print(f"    Font: name={f.name}, size={emu_to_pt(f.size) if f.size else None}pt, "
              f"bold={f.bold}, italic={f.italic}, underline={f.underline}")
        color = get_font_color(f)
        if color:
            print(f"    Font color: {color}")

    if hasattr(style, 'paragraph_format') and style.paragraph_format:
        pf = style.paragraph_format
        print(f"    Paragraph: align={alignment_name(pf.alignment)}")
        print(f"      space_before={emu_to_pt(pf.space_before) if pf.space_before else None}pt, "
              f"space_after={emu_to_pt(pf.space_after) if pf.space_after else None}pt")
        print(f"      line_spacing={pf.line_spacing} (rule={line_spacing_rule_name(pf.line_spacing_rule)})")
        print(f"      left_indent={emu_to_inches(pf.left_indent) if pf.left_indent else None}in, "
              f"right_indent={emu_to_inches(pf.right_indent) if pf.right_indent else None}in, "
              f"first_line_indent={emu_to_inches(pf.first_line_indent) if pf.first_line_indent else None}in")
        print(f"      keep_together={pf.keep_together}, keep_with_next={pf.keep_with_next}, "
              f"page_break_before={pf.page_break_before}")

        try:
            if pf.tab_stops:
                for tab in pf.tab_stops:
                    print(f"      Tab stop: {emu_to_inches(tab.position)}in align={tab.alignment} leader={tab.leader}")
        except:
            pass

print()

# ============================================================
# 3. WALK ENTIRE BODY XML - GET ALL PARAGRAPHS
# ============================================================
print("=" * 100)
print("3. ALL PARAGRAPHS FROM BODY XML (with section boundaries)")
print("=" * 100)

# We need to walk the body's children in order.
# Each child is either a <w:p> (paragraph) or <w:tbl> (table) or <w:sdt> (structured doc tag).
# Section breaks are found inside <w:pPr><w:sectPr> of the last paragraph of each section.

from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT

para_idx = 0
section_idx = 0

for child in body:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

    if tag == 'p':
        para = Paragraph(child, doc)
        pf = para.paragraph_format

        # Check for section break in this paragraph
        pPr = child.find(qn('w:pPr'))
        has_section_break = False
        sect_break_type = None
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                has_section_break = True
                sect_type_elem = sectPr.find(qn('w:type'))
                sect_break_type = sect_type_elem.get(qn('w:val')) if sect_type_elem is not None else 'nextPage'

        # Check for page breaks in runs
        run_page_breaks = []
        for run in para.runs:
            for br in run._element.findall(qn('w:br')):
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    run_page_breaks.append(True)

        # Also check for column breaks
        col_breaks = []
        for br_elem in child.findall('.//' + qn('w:br')):
            br_type = br_elem.get(qn('w:type'))
            if br_type == 'column':
                col_breaks.append(True)

        # Get direct paragraph properties from XML
        direct_jc = None
        direct_spacing = {}
        direct_indent = {}
        if pPr is not None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                direct_jc = jc.get(qn('w:val'))

            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                for attr in ['before', 'after', 'line', 'lineRule']:
                    val = spacing.get(qn(f'w:{attr}'))
                    if val is not None:
                        direct_spacing[attr] = val

            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                for attr in ['left', 'right', 'firstLine', 'hanging']:
                    val = ind.get(qn(f'w:{attr}'))
                    if val is not None:
                        direct_indent[attr] = val

        # Get tab stops
        tab_stops_info = []
        if pPr is not None:
            tabs = pPr.find(qn('w:tabs'))
            if tabs is not None:
                for tab in tabs.findall(qn('w:tab')):
                    tab_val = tab.get(qn('w:val'))
                    tab_pos = tab.get(qn('w:pos'))
                    tab_leader = tab.get(qn('w:leader'))
                    tab_stops_info.append(f"{tab_val}@{twips_to_inches(tab_pos)}in" + (f"(leader={tab_leader})" if tab_leader else ""))

        # Format markers
        markers = []
        if has_section_break:
            markers.append(f"SECTION_BREAK({sect_break_type})")
        if run_page_breaks:
            markers.append("RUN_PAGE_BREAK")
        if col_breaks:
            markers.append("COLUMN_BREAK")
        if pf.page_break_before:
            markers.append("PAGE_BREAK_BEFORE")
        marker_str = " [" + ", ".join(markers) + "]" if markers else ""

        text = para.text[:120]

        # Spacing info
        sp_str = ""
        if direct_spacing:
            parts = []
            if 'before' in direct_spacing:
                parts.append(f"before={twips_to_pt(direct_spacing['before'])}pt")
            if 'after' in direct_spacing:
                parts.append(f"after={twips_to_pt(direct_spacing['after'])}pt")
            if 'line' in direct_spacing:
                line_val = direct_spacing['line']
                rule = direct_spacing.get('lineRule', 'auto')
                if rule == 'exact':
                    parts.append(f"line={twips_to_pt(line_val)}pt(EXACT)")
                elif rule == 'atLeast':
                    parts.append(f"line={twips_to_pt(line_val)}pt(AT_LEAST)")
                else:
                    # auto = multiple of line spacing; value is in 240ths of a line
                    parts.append(f"line={round(int(line_val)/240, 2)}x(MULTIPLE)")
            sp_str = " spacing=[" + ", ".join(parts) + "]"

        # Indent info
        ind_str = ""
        if direct_indent:
            parts = []
            for k, v in direct_indent.items():
                parts.append(f"{k}={twips_to_inches(v)}in")
            ind_str = " indent=[" + ", ".join(parts) + "]"

        tab_str = ""
        if tab_stops_info:
            tab_str = " tabs=[" + ", ".join(tab_stops_info) + "]"

        print(f"\nPara[{para_idx}] style='{para.style.name}' jc={direct_jc}{sp_str}{ind_str}{tab_str}{marker_str}")
        print(f"  Text: '{text}'")

        # Print runs
        for r_idx, run in enumerate(para.runs):
            f = run.font
            color = get_font_color(f)

            # Get direct run properties from XML
            rPr = run._element.find(qn('w:rPr'))
            direct_font_name = None
            direct_font_size = None
            direct_bold = None
            direct_italic = None
            direct_underline = None
            direct_caps = None
            direct_smallcaps = None
            direct_superscript = None
            direct_color = None

            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    direct_font_name = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi')) or rFonts.get(qn('w:cs'))

                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    # sz value is in half-points
                    half_pts = sz.get(qn('w:val'))
                    if half_pts:
                        direct_font_size = round(int(half_pts) / 2, 1)

                b = rPr.find(qn('w:b'))
                if b is not None:
                    bval = b.get(qn('w:val'))
                    direct_bold = bval != '0' if bval is not None else True

                i = rPr.find(qn('w:i'))
                if i is not None:
                    ival = i.get(qn('w:val'))
                    direct_italic = ival != '0' if ival is not None else True

                u = rPr.find(qn('w:u'))
                if u is not None:
                    direct_underline = u.get(qn('w:val'))

                caps = rPr.find(qn('w:caps'))
                if caps is not None:
                    direct_caps = True

                smallCaps = rPr.find(qn('w:smallCaps'))
                if smallCaps is not None:
                    direct_smallcaps = True

                vertAlign = rPr.find(qn('w:vertAlign'))
                if vertAlign is not None:
                    direct_superscript = vertAlign.get(qn('w:val'))

                color_elem = rPr.find(qn('w:color'))
                if color_elem is not None:
                    direct_color = color_elem.get(qn('w:val'))

                # Check for rStyle (character style reference)
                rStyle = rPr.find(qn('w:rStyle'))
                if rStyle is not None:
                    char_style = rStyle.get(qn('w:val'))
                    if char_style:
                        direct_font_name = f"{direct_font_name} [rStyle={char_style}]" if direct_font_name else f"[rStyle={char_style}]"

            # Check for special elements in run (breaks, images, etc.)
            special = []
            for elem in run._element:
                tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag_local not in ('rPr', 't'):
                    if tag_local == 'br':
                        br_type = elem.get(qn('w:type'))
                        special.append(f"br({br_type or 'line'})")
                    elif tag_local == 'tab':
                        special.append("tab")
                    elif tag_local == 'drawing':
                        special.append("drawing/image")
                    else:
                        special.append(tag_local)
            special_str = f" special=[{','.join(special)}]" if special else ""

            run_text = run.text[:80]

            print(f"  Run[{r_idx}]: font={direct_font_name} size={direct_font_size}pt "
                  f"B={direct_bold} I={direct_italic} U={direct_underline} "
                  f"sup={direct_superscript} color={direct_color} "
                  f"caps={direct_caps} smallcaps={direct_smallcaps}"
                  f"{special_str}")
            print(f"           '{run_text}'")

        # If paragraph has no runs but has text (unlikely), note it
        if len(para.runs) == 0 and para.text.strip():
            print(f"  (no runs, but text exists: '{para.text[:80]}')")

        # If paragraph has no runs and no text, check for special elements
        if len(para.runs) == 0:
            # Check for bookmark starts, hyperlinks, etc.
            for elem in child:
                tag_local = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                if tag_local not in ('pPr',):
                    if tag_local == 'hyperlink':
                        link_text = ''.join(t.text or '' for t in elem.findall('.//' + qn('w:t')))
                        rId = elem.get(qn('r:id'))
                        print(f"  Hyperlink: rId={rId}, text='{link_text[:80]}'")
                        # Get run formatting inside hyperlink
                        for h_run_elem in elem.findall(qn('w:r')):
                            h_rPr = h_run_elem.find(qn('w:rPr'))
                            if h_rPr is not None:
                                h_rFonts = h_rPr.find(qn('w:rFonts'))
                                h_font = h_rFonts.get(qn('w:ascii')) if h_rFonts is not None else None
                                h_sz = h_rPr.find(qn('w:sz'))
                                h_size = round(int(h_sz.get(qn('w:val'))) / 2, 1) if h_sz is not None else None
                                h_b = h_rPr.find(qn('w:b'))
                                h_bold = True if h_b is not None else None
                                h_i = h_rPr.find(qn('w:i'))
                                h_italic = True if h_i is not None else None
                                h_text = ''.join(t.text or '' for t in h_run_elem.findall(qn('w:t')))
                                print(f"    HyperlinkRun: font={h_font} size={h_size}pt B={h_bold} I={h_italic} '{h_text[:60]}'")
                    elif tag_local == 'r':
                        # Orphan run not in para.runs (should be rare)
                        r_text = ''.join(t.text or '' for t in elem.findall(qn('w:t')))
                        if r_text.strip():
                            print(f"  OrphanRun: '{r_text[:80]}'")

        if has_section_break:
            section_idx += 1
            print(f"  ======= END OF SECTION {section_idx - 1} / START OF SECTION {section_idx} =======")

        para_idx += 1

    elif tag == 'tbl':
        print(f"\n  [TABLE at position after Para[{para_idx-1}]]")
        # We'll detail tables later

    elif tag == 'sdt':
        print(f"\n  [STRUCTURED_DOC_TAG at position after Para[{para_idx-1}]]")

print(f"\nTotal paragraphs walked: {para_idx}")
print()

# ============================================================
# 4. UNIQUE RUN FORMATTING PATTERNS
# ============================================================
print("=" * 100)
print("4. UNIQUE RUN FORMATTING PATTERNS (from all paragraphs)")
print("=" * 100)

run_patterns = {}
for p_elem in all_para_elements:
    para = Paragraph(p_elem, doc)
    for run in para.runs:
        rPr = run._element.find(qn('w:rPr'))

        font_name = None
        font_size = None
        bold = None
        italic = None
        underline = None
        superscript = None
        color = None
        caps = None
        smallcaps = None

        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                font_name = rFonts.get(qn('w:ascii')) or rFonts.get(qn('w:hAnsi'))

            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                half_pts = sz.get(qn('w:val'))
                if half_pts:
                    font_size = round(int(half_pts) / 2, 1)

            b = rPr.find(qn('w:b'))
            if b is not None:
                bval = b.get(qn('w:val'))
                bold = bval != '0' if bval is not None else True

            i_elem = rPr.find(qn('w:i'))
            if i_elem is not None:
                ival = i_elem.get(qn('w:val'))
                italic = ival != '0' if ival is not None else True

            u = rPr.find(qn('w:u'))
            if u is not None:
                underline = u.get(qn('w:val'))

            vertAlign = rPr.find(qn('w:vertAlign'))
            if vertAlign is not None:
                superscript = vertAlign.get(qn('w:val'))

            color_elem = rPr.find(qn('w:color'))
            if color_elem is not None:
                color = color_elem.get(qn('w:val'))

            caps_elem = rPr.find(qn('w:caps'))
            if caps_elem is not None:
                caps = True

            sc_elem = rPr.find(qn('w:smallCaps'))
            if sc_elem is not None:
                smallcaps = True

        pattern_key = (font_name, font_size, bold, italic, underline, superscript, color, caps, smallcaps)

        if pattern_key not in run_patterns:
            run_patterns[pattern_key] = {
                'count': 0,
                'example': run.text[:60] if run.text else "(empty)",
                'para_style': para.style.name,
            }
        run_patterns[pattern_key]['count'] += 1

print(f"\nTotal unique run formatting patterns: {len(run_patterns)}")
for pattern, info in sorted(run_patterns.items(), key=lambda x: -x[1]['count']):
    fn, fs, b, it, u, sup, col, caps, sc = pattern
    print(f"\n  Pattern ({info['count']}x): font={fn} size={fs}pt B={b} I={it} U={u} sup={sup} color={col} caps={caps} smallcaps={sc}")
    print(f"    Example: '{info['example']}' (style='{info['para_style']}')")

print()

# ============================================================
# 5. HEADERS AND FOOTERS
# ============================================================
print("=" * 100)
print("5. HEADERS AND FOOTERS")
print("=" * 100)

for i, section in enumerate(doc.sections):
    print(f"\n--- Section {i} ---")

    for hf_type, hf_name in [
        (section.header, "Header"),
        (section.first_page_header, "First Page Header"),
        (section.even_page_header, "Even Page Header"),
        (section.footer, "Footer"),
        (section.first_page_footer, "First Page Footer"),
        (section.even_page_footer, "Even Page Footer"),
    ]:
        try:
            if hf_type is None:
                continue
            linked = hf_type.is_linked_to_previous
            paras = hf_type.paragraphs
            has_content = any(p.text.strip() for p in paras)

            if not has_content and linked:
                continue  # Skip linked empty headers

            print(f"\n  {hf_name} (linked_to_previous={linked}):")
            for p_idx, para in enumerate(paras):
                if para.text.strip() or len(para.runs) > 0:
                    pPr = para._element.find(qn('w:pPr'))
                    direct_jc = None
                    if pPr is not None:
                        jc = pPr.find(qn('w:jc'))
                        if jc is not None:
                            direct_jc = jc.get(qn('w:val'))

                    print(f"    Para[{p_idx}] style='{para.style.name}' jc={direct_jc}: '{para.text[:100]}'")
                    for r_idx, run in enumerate(para.runs):
                        f = run.font
                        rPr_elem = run._element.find(qn('w:rPr'))
                        rfn = None
                        rsz = None
                        if rPr_elem is not None:
                            rf = rPr_elem.find(qn('w:rFonts'))
                            if rf is not None:
                                rfn = rf.get(qn('w:ascii'))
                            sz = rPr_elem.find(qn('w:sz'))
                            if sz is not None:
                                rsz = round(int(sz.get(qn('w:val'))) / 2, 1)
                        print(f"      Run[{r_idx}]: font={rfn} size={rsz}pt B={f.bold} I={f.italic} '{run.text[:60]}'")
        except Exception as e:
            print(f"  {hf_name}: Error - {e}")

print()

# ============================================================
# 6. IMAGES AND DRAWINGS
# ============================================================
print("=" * 100)
print("6. IMAGES AND DRAWINGS")
print("=" * 100)

for idx, p_elem in enumerate(all_para_elements):
    drawings = p_elem.findall('.//' + qn('w:drawing'))
    if drawings:
        para = Paragraph(p_elem, doc)
        print(f"\n  Para[{idx}] has {len(drawings)} drawing(s)")
        print(f"    Text: '{para.text[:80]}'")
        for d_idx, drawing in enumerate(drawings):
            inline = drawing.find(qn('wp:inline'))
            anchor = drawing.find(qn('wp:anchor'))
            target = inline if inline is not None else anchor
            disp_type = "inline" if inline is not None else "anchor"
            if target is not None:
                extent = target.find(qn('wp:extent'))
                if extent is not None:
                    cx = extent.get('cx')
                    cy = extent.get('cy')
                    print(f"    Drawing[{d_idx}]: {disp_type}, w={emu_to_inches(int(cx)) if cx else '?'}in, h={emu_to_inches(int(cy)) if cy else '?'}in")

                # Try to get image description
                docPr = target.find(qn('wp:docPr'))
                if docPr is not None:
                    name = docPr.get('name')
                    descr = docPr.get('descr')
                    print(f"    Name: {name}, Description: {descr}")

print()

# ============================================================
# 7. TABLES
# ============================================================
print("=" * 100)
print("7. TABLES")
print("=" * 100)

if len(doc.tables) == 0:
    print("\n  No tables found.")
else:
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {t_idx} ---")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text[:80]
                if text.strip():
                    print(f"  Cell[{r_idx},{c_idx}]: '{text}'")

print()

# ============================================================
# 8. RAW XML OF FIRST 3 PARAGRAPHS
# ============================================================
print("=" * 100)
print("8. RAW XML SAMPLES")
print("=" * 100)

print("\n--- First 3 body paragraphs ---")
for idx, p_elem in enumerate(all_para_elements[:3]):
    xml_str = etree.tostring(p_elem, pretty_print=True).decode()
    print(f"\n--- Para[{idx}] XML ---")
    print(xml_str[:2000])

print()

# ============================================================
# 9. SUMMARY OF FORMATTING BY CONTENT TYPE
# ============================================================
print("=" * 100)
print("9. SUMMARY: STYLE USAGE COUNTS AND EXAMPLES")
print("=" * 100)

style_usage = {}
for p_elem in all_para_elements:
    para = Paragraph(p_elem, doc)
    sn = para.style.name
    if sn not in style_usage:
        style_usage[sn] = {'count': 0, 'examples': []}
    style_usage[sn]['count'] += 1
    if len(style_usage[sn]['examples']) < 3 and para.text.strip():
        style_usage[sn]['examples'].append(para.text[:80])

for sn, info in sorted(style_usage.items(), key=lambda x: -x[1]['count']):
    print(f"\n  '{sn}': used {info['count']} times")
    for ex in info['examples']:
        print(f"    Example: '{ex}'")

print()
print("=" * 100)
print("ANALYSIS COMPLETE")
print("=" * 100)
